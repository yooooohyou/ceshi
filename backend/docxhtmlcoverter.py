import logging
logger = logging.getLogger(__name__)

from spire.doc import *
from spire.doc.common import *
import os
import zipfile
import shutil
import re
import base64
import tempfile
import uuid
from docx import Document as PythonDocx          # python-docx：用于切分
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree



class DocxHtmlConverter:
    """
    DOCX与HTML互转工具类
    核心特性：
    1. 所有路径强制使用绝对路径，脱离工作目录依赖
    2. 临时目录在每次 docx_to_single_html 调用时生成唯一ID，避免多线程/多调用冲突
    3. 图片顺序解析覆盖正文、页眉、页脚、脚注、尾注等所有XML区域
    4. 降级兜底：Spire生成图片路径（find_actual_img_dir递归查找）
    5. 增强路径校验和异常处理
    6. 超限文档自动分片处理（段落>450或表格>20时触发）
    7. 大表格按行拆分，合并HTML时自动还原
    8. 合并时去除页眉页脚
    9. 分片流程：chunk不内嵌图片 → 流式合并 → 统一内嵌，避免内存溢出

    【修复图片顺序错位问题】
    - _embed_images_to_html 改为二进制精确匹配策略，彻底脱离对 Spire 输出顺序的依赖
    - 非分片路径的 docx_to_single_html 同步使用二进制匹配替换图片
    - 非分片路径补充调用 _clean_header_footer，与分片路径行为一致
    """

    def __init__(self):
        self.default_image_format = 0  # 0=PNG，1=JPG，2=BMP，3=GIF
        self.html_validation_type = XHTMLValidationType.none
        self.MAX_PARAGRAPHS = 450
        self.MAX_TABLES = 20

    # ------------------------------------------------------------------ #
    #  路径工具                                                             #
    # ------------------------------------------------------------------ #

    def _apply_style_and_preserve_format(self, para, target_builtin_style):
        """
        [私有方法] 自定义的“无损 ApplyStyle”函数。
        将段落设置为标准标题，同时保留原有文字的字体、字号。
        """
        # 1. 备份该段落中每个文本块 (TextRange) 的格式
        format_backup = []

        # 遍历段落里的所有子对象
        for i in range(para.ChildObjects.Count):
            obj = para.ChildObjects.get_Item(i)
            # 如果是文字块，提取它的字体、字号、颜色等
            if isinstance(obj, TextRange):
                backup = {
                    "font_name": obj.CharacterFormat.FontName,
                    "font_size": obj.CharacterFormat.FontSize,
                    "text_color": obj.CharacterFormat.TextColor,
                    "is_bold": obj.CharacterFormat.Bold
                }
                format_backup.append((obj, backup))

        # 2. 调用原生的 ApplyStyle
        para.ApplyStyle(target_builtin_style)

        # 3. 将备份的格式重新应用回每个文本块
        for obj, backup in format_backup:
            if backup["font_name"]:
                obj.CharacterFormat.FontName = backup["font_name"]
            if backup["font_size"] and backup["font_size"] > 0:
                obj.CharacterFormat.FontSize = backup["font_size"]
            # obj.CharacterFormat.TextColor = backup["text_color"]
            # obj.CharacterFormat.Bold = backup["is_bold"]

    def _fix_underline_span_width(self, html_content: str) -> str:
        """
        [私有方法] 修复下划线变长的问题：
        清理带有固定 width 和 display: inline-block 的下划线 span，
        让其仅依靠内部的 &nbsp; 或空格自然撑开宽度，避免渲染引擎叠加计算。
        """

        def replacer(m):
            span_tag = m.group(1)
            content = m.group(2)
            span_close = m.group(3)

            # 特征检查：
            # 1. 包含下划线
            # 2. 包含 width 宽度限制
            # 3. 内部确实有实体内容（如 &nbsp; 或文本），如果是纯空 span 则不能删 width
            is_underline = 'underline' in span_tag.lower()
            has_width = 'width' in span_tag.lower()
            has_content = bool(content.strip().replace('&nbsp;', '')) or '&nbsp;' in content

            if is_underline and has_width and has_content:
                # 同时处理 style 和 data-mce-style (兼容 TinyMCE)
                for attr_name in ['style', 'data-mce-style']:
                    attr_m = re.search(rf'{attr_name}="([^"]*)"', span_tag, re.IGNORECASE)
                    if attr_m:
                        s = attr_m.group(1)

                        # 剔除导致变长的物理维度属性
                        s = re.sub(r'\bwidth\s*:\s*[\d.]+[a-zA-Z%]+\s*;?', '', s, flags=re.IGNORECASE)
                        s = re.sub(r'\bdisplay\s*:\s*inline-block\s*;?', '', s, flags=re.IGNORECASE)
                        s = re.sub(r'\btext-indent\s*:\s*[\d.]+[a-zA-Z%]+\s*;?', '', s, flags=re.IGNORECASE)

                        # 拼接回标签中
                        new_attr = f'{attr_name}="{s.strip()}"'
                        span_tag = span_tag[:attr_m.start()] + new_attr + span_tag[attr_m.end():]

            return span_tag + content + span_close

        # 使用 DOTALL 匹配 HTML 中的 span 块
        return re.sub(r'(<span\b[^>]*>)(.*?)(</span\s*>)', replacer, html_content, flags=re.IGNORECASE | re.DOTALL)

    def _clean_docx_headings_before_convert(self, document):
        """
        [私有方法] 在导出 HTML 前，遍历 DOCX 文档。
        将变体样式强制重置为标准内置标题样式（全面支持 1-9 级）。
        """
        heading_mapping = [
            (OutlineLevel.Level1, BuiltinStyle.Heading1, 1),
            (OutlineLevel.Level2, BuiltinStyle.Heading2, 2),
            (OutlineLevel.Level3, BuiltinStyle.Heading3, 3),
            (OutlineLevel.Level4, BuiltinStyle.Heading4, 4),
            (OutlineLevel.Level5, BuiltinStyle.Heading5, 5),
            (OutlineLevel.Level6, BuiltinStyle.Heading6, 6),
            (OutlineLevel.Level7, BuiltinStyle.Heading7, 7),
            (OutlineLevel.Level8, BuiltinStyle.Heading8, 8),
            (OutlineLevel.Level9, BuiltinStyle.Heading9, 9),
        ]

        # 遍历文档的所有节 (Section)
        for i in range(document.Sections.Count):
            section = document.Sections.get_Item(i)

            # 遍历节内的所有段落
            for j in range(section.Body.Paragraphs.Count):
                para = section.Body.Paragraphs.get_Item(j)
                style_name = para.StyleName if para.StyleName else ""

                # 检查当前段落是否属于 1-9 级标题
                for outline_lvl, builtin_style, level_num in heading_mapping:
                    pattern = fr'(标题|Heading)\s*{level_num}'

                    if (para.Format.OutlineLevel == outline_lvl or
                            re.search(pattern, style_name, re.IGNORECASE)):
                        # 调用类内部的私有方法：注意这里加了 self.
                        self._apply_style_and_preserve_format(para, builtin_style)
                        break

    def _normalize_path(self, path):
        """【内部方法】统一路径格式并转为绝对路径"""
        if not path:
            return ""
        abs_path = os.path.abspath(path)
        return abs_path.replace('\\', '/').replace('//', '/')

    def _make_temp_dir_prefix(self):
        """
        每次调用生成新的唯一前缀，避免多次调用或多线程复用同一前缀。
        """
        return f"spire_temp_{uuid.uuid4().hex[:8]}"

    # ------------------------------------------------------------------ #
    #  图片工具                                                             #
    # ------------------------------------------------------------------ #

    def _get_image_order_from_docx(self, docx_path):
        """
        【内部方法】解析DOCX，提取图片在文档中的显示顺序
        覆盖范围：正文、页眉、页脚、脚注、尾注等所有XML区域

        返回值：image_order_list - 按文档顺序排列的原始图片文件名列表（去重）
        """
        image_order = []
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                all_files = zip_file.namelist()

                def _xml_sort_key(f):
                    name = os.path.basename(f)
                    if name == 'document.xml':
                        return (0, f)
                    if name in ('footnotes.xml', 'endnotes.xml'):
                        return (1, f)
                    return (2, f)  # header*.xml / footer*.xml

                target_xml_files = sorted(
                    [
                        f for f in all_files
                        if re.match(
                            r'word/(document|header\d*|footer\d*|footnotes|endnotes)\.xml$', f
                        )
                    ],
                    key=_xml_sort_key,
                )

                all_rels = {}
                for xml_file in target_xml_files:
                    xml_basename = os.path.basename(xml_file)
                    rels_path = f'word/_rels/{xml_basename}.rels'
                    if rels_path in all_files:
                        rels_content = zip_file.read(rels_path).decode('utf-8')
                        rel_pattern = re.compile(
                            r'<Relationship\s+Id="(rId\d+)"\s+Type="[^"]*image[^"]*"\s+Target="([^"]+)"'
                        )
                        rels_for_this_xml = {}
                        for match in rel_pattern.finditer(rels_content):
                            r_id, target = match.group(1), match.group(2)
                            rels_for_this_xml[r_id] = os.path.basename(target)
                        all_rels[xml_file] = rels_for_this_xml

                seen = set()
                id_pattern = re.compile(r'r:(?:embed|link|id)="(rId\d+)"')
                for xml_file in target_xml_files:
                    if xml_file not in all_rels:
                        continue
                    xml_content = zip_file.read(xml_file).decode('utf-8')
                    rels_map = all_rels[xml_file]
                    for match in id_pattern.finditer(xml_content):
                        r_id = match.group(1)
                        if r_id in rels_map:
                            img_name = rels_map[r_id]
                            if img_name not in seen:
                                seen.add(img_name)
                                image_order.append(img_name)

        except Exception as e:
            logger.warning(f"⚠️ 解析图片顺序失败：{e}，将使用文件名排序")
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                image_order = sorted(
                    os.path.basename(f.filename)
                    for f in zip_file.infolist()
                    if f.filename.startswith('word/media/') and not f.is_dir()
                )

        logger.debug(f"✅ 解析到图片显示顺序：{image_order}")
        return image_order

    def _extract_original_images(self, docx_path, output_img_dir):
        """【内部方法】从DOCX中提取原始无压缩图片（强制绝对路径）"""
        output_img_dir = self._normalize_path(output_img_dir)

        if os.path.exists(output_img_dir):
            shutil.rmtree(output_img_dir, ignore_errors=True)
        os.makedirs(output_img_dir, exist_ok=True)

        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            for file_info in zip_file.infolist():
                if file_info.filename.startswith('word/media/') and not file_info.is_dir():
                    img_filename = os.path.basename(file_info.filename)
                    save_path = self._normalize_path(
                        os.path.join(output_img_dir, img_filename)
                    )
                    with open(save_path, 'wb') as f:
                        f.write(zip_file.read(file_info.filename))

        return [
            f for f in os.listdir(output_img_dir)
            if os.path.isfile(os.path.join(output_img_dir, f))
        ]

    def _find_actual_img_dir(self, base_dir):
        """
        【内部方法】递归查找第一个实际包含图片文件的目录
        用于兜底：Spire有时将图片输出到嵌套子目录中
        """
        img_exts = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
        for root, dirs, files in os.walk(base_dir):
            if any(f.lower().endswith(img_exts) for f in files):
                return root
        return base_dir

    def _convert_emf_to_png(self, emf_path):
        """
        【内部方法】将 EMF/WMF 文件转换为 PNG，返回 PNG 文件路径。
        优先调用 LibreOffice 命令行（soffice）光栅化，失败则降级到
        spire.doc.common.Image。转换结果缓存在原文件同目录下。
        转换失败返回 None。
        """
        import subprocess

        emf_path = self._normalize_path(emf_path)
        emf_dir  = os.path.dirname(emf_path)
        emf_stem = os.path.splitext(os.path.basename(emf_path))[0]
        png_path = self._normalize_path(os.path.join(emf_dir, emf_stem + '_converted.png'))

        if os.path.exists(png_path):
            return png_path

        soffice_candidates = [
            'libreoffice',
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            '/usr/bin/soffice',
            '/usr/lib/libreoffice/program/soffice',
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        ]

        for soffice in soffice_candidates:
            try:
                result = subprocess.run(
                    [soffice, '--headless', '--convert-to', 'png', '--outdir', emf_dir, emf_path],
                    capture_output=True,
                    timeout=30,
                )
                lo_output = self._normalize_path(os.path.join(emf_dir, emf_stem + '.png'))
                if result.returncode == 0 and os.path.exists(lo_output):
                    if lo_output != png_path:
                        os.rename(lo_output, png_path)
                    logger.debug(f"   🔄 EMF→PNG（LibreOffice）：{os.path.basename(emf_path)}")
                    return png_path
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue
            except Exception as e:
                logger.debug(f"   ⚠️ LibreOffice 转换异常：{e}")
                continue

        try:
            from spire.doc.common import Image as SpireImage
            img = SpireImage.FromFile(emf_path)
            img.Save(png_path)
            if os.path.exists(png_path):
                logger.debug(f"   🔄 EMF→PNG（Spire.Image）：{os.path.basename(emf_path)}")
                return png_path
        except Exception as e:
            logger.debug(f"   ⚠️ Spire.Image 转换失败：{e}")
        logger.debug(f"   ⚠️ EMF→PNG 所有方案均失败，跳过：{os.path.basename(emf_path)}")
        return None

    def _image_to_base64(self, img_path):
        """【内部方法】将图片文件转为Base64编码（带MIME前缀）。
        对 EMF/WMF 格式自动先转为 PNG，再做 base64。
        """
        img_path = self._normalize_path(img_path)
        try:
            if not os.path.exists(img_path):
                logger.warning(f"⚠️ 图片文件不存在（绝对路径）：{img_path}")
                return ""

            img_ext = os.path.splitext(img_path)[1].lower()
            if img_ext in ('.emf', '.wmf'):
                converted = self._convert_emf_to_png(img_path)
                if converted and os.path.exists(converted):
                    img_path = converted
                    img_ext  = '.png'
                else:
                    return ""

            with open(img_path, 'rb') as f:
                img_data = f.read()

            mime_map = {
                '.png':  'image/png',
                '.jpg':  'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif':  'image/gif',
                '.bmp':  'image/bmp',
                '.webp': 'image/webp',
                '.svg':  'image/svg+xml',
            }
            mime_type = mime_map.get(img_ext, 'image/png')
            return f"data:{mime_type};base64,{base64.b64encode(img_data).decode('utf-8')}"

        except Exception as e:
            logger.warning(f"⚠️ 图片 {img_path} 转Base64失败：{e}")
            return ""

    def _build_spire_to_original_map(self, spire_img_names, image_display_order,
                                     original_img_dir, fallback_img_dir):
        """
        构建 Spire生成图片名 → 原始图片绝对路径 的映射。
        优先精确文件名匹配，失败时位置索引降级，再失败时 fallback 目录查找。
        """
        result = {}

        orig_stem_map = {}
        for orig_name in image_display_order:
            stem = os.path.splitext(orig_name)[0].lower()
            abs_path = self._normalize_path(os.path.join(original_img_dir, orig_name))
            if os.path.exists(abs_path):
                orig_stem_map[stem] = abs_path

        for idx, spire_name in enumerate(spire_img_names):
            spire_stem = os.path.splitext(spire_name)[0].lower()

            if spire_stem in orig_stem_map:
                result[spire_name] = orig_stem_map[spire_stem]
                continue

            if idx < len(image_display_order):
                candidate = self._normalize_path(
                    os.path.join(original_img_dir, image_display_order[idx])
                )
                if os.path.exists(candidate):
                    result[spire_name] = candidate
                    logger.debug(f"   ⚠️ {spire_name} 精确匹配失败，位置索引降级 → {image_display_order[idx]}")
                    continue

            fallback = self._normalize_path(os.path.join(fallback_img_dir, spire_name))
            if os.path.exists(fallback):
                result[spire_name] = fallback
                logger.debug(f"   ⚠️ {spire_name} 降级到 Spire 生成目录")
            else:
                logger.debug(f"   ⚠️ {spire_name} 找不到对应原始图片，跳过")
        return result

    # ------------------------------------------------------------------ #
    #  DPI / 尺寸修正工具                                                   #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _extract_image_display_sizes(docx_path):
        """
        从 DOCX 的 document.xml 中提取每张图片的"显示尺寸"（即 Word 排版时的实际渲染宽高），
        返回 {原始图片文件名: (width_px, height_px)} 的映射，分辨率基准为 96 DPI。
        """
        EMU_PER_INCH  = 914400
        PX_PER_INCH   = 96
        EMU_TO_PX     = PX_PER_INCH / EMU_PER_INCH

        WP_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        A_NS   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        R_NS   = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        VML_NS = 'urn:schemas-microsoft-com:vml'
        W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        size_map = {}

        try:
            with zipfile.ZipFile(docx_path, 'r') as zf:
                rels_xml = zf.read('word/_rels/document.xml.rels').decode('utf-8')
                rid_to_img = {}
                for m in re.finditer(
                    r'<Relationship\s+Id="(rId\d+)"\s+Type="[^"]*image[^"]*"\s+Target="([^"]+)"',
                    rels_xml
                ):
                    rid_to_img[m.group(1)] = os.path.basename(m.group(2))

                doc_xml = zf.read('word/document.xml').decode('utf-8')

            root = etree.fromstring(doc_xml.encode('utf-8'))

            for container_tag in (f'{{{WP_NS}}}inline', f'{{{WP_NS}}}anchor'):
                for container in root.iter(container_tag):
                    extent = container.find(f'{{{WP_NS}}}extent')
                    if extent is None:
                        continue
                    try:
                        cx = int(extent.get('cx', 0))
                        cy = int(extent.get('cy', 0))
                    except ValueError:
                        continue

                    w_px = round(cx * EMU_TO_PX)
                    h_px = round(cy * EMU_TO_PX)

                    blip = container.find(f'.//{{{A_NS}}}blip')
                    if blip is not None:
                        r_embed = blip.get(f'{{{R_NS}}}embed')
                        if r_embed and r_embed in rid_to_img:
                            fname = rid_to_img[r_embed]
                            if fname not in size_map:
                                size_map[fname] = (w_px, h_px)

            def _css_dim_to_px(val_str):
                if not val_str:
                    return None
                val_str = val_str.strip().lower()
                try:
                    if val_str.endswith('pt'):
                        return round(float(val_str[:-2]) * PX_PER_INCH / 72)
                    if val_str.endswith('in'):
                        return round(float(val_str[:-2]) * PX_PER_INCH)
                    if val_str.endswith('cm'):
                        return round(float(val_str[:-2]) / 2.54 * PX_PER_INCH)
                    if val_str.endswith('mm'):
                        return round(float(val_str[:-2]) / 25.4 * PX_PER_INCH)
                    if val_str.endswith('px'):
                        return round(float(val_str[:-2]))
                except (ValueError, AttributeError):
                    pass
                return None

            for shape in root.iter(f'{{{VML_NS}}}shape'):
                imagedata = shape.find(f'{{{VML_NS}}}imagedata')
                if imagedata is None:
                    for child in shape:
                        if child.tag.endswith('}imagedata') or child.tag == 'imagedata':
                            imagedata = child
                            break
                if imagedata is None:
                    continue

                r_id = imagedata.get(f'{{{R_NS}}}id')
                if not r_id or r_id not in rid_to_img:
                    continue

                fname = rid_to_img[r_id]
                if fname in size_map:
                    continue

                style = shape.get('style', '')
                style_dict = {}
                for part in style.split(';'):
                    if ':' in part:
                        k, _, v = part.partition(':')
                        style_dict[k.strip().lower()] = v.strip()

                w_px = _css_dim_to_px(style_dict.get('width'))
                h_px = _css_dim_to_px(style_dict.get('height'))

                if w_px and h_px:
                    size_map[fname] = (w_px, h_px)

        except Exception as e:
            logger.warning(f"⚠️ 提取图片显示尺寸失败：{e}")
        logger.debug(f"📐 从 DOCX 提取到 {len(size_map)} 张图片的显示尺寸")
        return size_map

    @staticmethod
    def _extract_content_width_pt(docx_path: str) -> float:
        """
        从 DOCX 的 sectPr 中读取页面宽度和左右边距，计算实际内容区宽度（单位 pt）。
        单位换算：OOXML 使用 twips（1 twip = 1/20 pt）。
        失败时返回 451.0 pt（A4 + 2.54cm 双边距的典型值）。
        """
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        DEFAULT_PT = 451.0
        try:
            with zipfile.ZipFile(docx_path, 'r') as zf:
                doc_xml = zf.read('word/document.xml').decode('utf-8')
            root = etree.fromstring(doc_xml.encode('utf-8'))
            sectPr = root.find(f'.//{{{W_NS}}}sectPr')
            if sectPr is None:
                return DEFAULT_PT
            pgSz  = sectPr.find(f'{{{W_NS}}}pgSz')
            pgMar = sectPr.find(f'{{{W_NS}}}pgMar')
            if pgSz is None:
                return DEFAULT_PT
            page_w  = int(pgSz.get(f'{{{W_NS}}}w', 11906))
            left    = int(pgMar.get(f'{{{W_NS}}}left',  1800)) if pgMar is not None else 1800
            right   = int(pgMar.get(f'{{{W_NS}}}right', 1800)) if pgMar is not None else 1800
            content = (page_w - left - right) / 20.0
            logger.debug(f"📐 页面内容宽度：{content:.1f}pt（页宽{page_w/20:.1f}pt - 左{left/20:.1f}pt - 右{right/20:.1f}pt）")
            return max(content, 200.0)
        except Exception as e:
            logger.warning(f"⚠️ 提取页面内容宽度失败：{e}，使用默认值 {DEFAULT_PT}pt")
            return DEFAULT_PT

    def _make_tables_responsive(self, html_content: str) -> str:
        """
        终极表格修正方案：完美适配 TinyMCE 独立列拖拽，互不影响。
        1. 移除 <table> 总宽度，解除锁定。
        2. 剔除 colspan 合并单元格的宽度干扰。
        3. 真实列宽转为绝对像素 px。
        """
        # 1. 移除 <table> 上的任何 width (包括 style 和 属性)，并注入固定布局
        html = re.sub(
            r'(<table\b[^>]*?(?:style|data-mce-style)="[^"]*?)(?<![-a-zA-Z])width\s*:\s*[\d.]+[a-zA-Z%]+;?',
            r'\1 width=100%;table-layout: fixed; word-break: break-all;',
            html_content,
            flags=re.IGNORECASE
        )
        html = re.sub(
            r'(<table\b[^>]*?)\s*\bwidth="\d+(?:\.\d+)?%?"',
            r'\1 width=100%;',
            html,
            flags=re.IGNORECASE
        )

        # 2. 剔除所有包含 colspan 的 <td> 的 width 属性
        # （这是防止 TinyMCE 把带有 colspan 的首行当做计算网格基准导致崩溃重置的核心）
        for _ in range(2):
            html = re.sub(
                r'(<t[dh]\b[^>]*?colspan="\d+"[^>]*?(?:style|data-mce-style)="[^"]*?)(?<![-a-zA-Z])width\s*:\s*[\d.]+[a-zA-Z%]+;?',
                r'\1',
                html,
                flags=re.IGNORECASE
            )

        # 3. 将正常的 <td> / <th> 中的 pt 宽度转换为绝对像素 px
        def pt_to_px(match):
            prefix = match.group(1)
            pt_val = float(match.group(2))
            px_val = round(pt_val * 1.3333) # 1pt = 1.3333px
            return f"{prefix}width: {px_val}px;"

        for _ in range(2):
            html = re.sub(
                r'(<t[dh]\b[^>]*?(?:style|data-mce-style)="[^"]*?)(?<![-a-zA-Z])width\s*:\s*([\d.]+)pt;?',
                pt_to_px,
                html,
                flags=re.IGNORECASE
            )

        return html

    def _fix_html_img_sizes(self, html_content: str, size_map: dict,
                             spire_img_names: list, image_display_order: list) -> str:
        """
        将 HTML 中每个 <img> 的 width/height 属性强制设置为从 DOCX <wp:extent> 读取的显示尺寸。
        """
        if not size_map:
            return html_content

        spire_to_orig = {}
        orig_stems = {
            os.path.splitext(n)[0].lower(): n
            for n in image_display_order
        }
        for idx, spire_name in enumerate(spire_img_names):
            spire_stem = os.path.splitext(spire_name)[0].lower()
            if spire_stem in orig_stems:
                spire_to_orig[spire_name] = orig_stems[spire_stem]
            elif idx < len(image_display_order):
                spire_to_orig[spire_name] = image_display_order[idx]

        def _replace_img_tag(m):
            tag = m.group(0)
            src_m = re.search(r'src="([^"]+)"', tag)
            if not src_m:
                return tag

            src_basename = os.path.basename(src_m.group(1))
            orig_name = spire_to_orig.get(src_basename)
            if not orig_name:
                return tag

            sizes = size_map.get(orig_name)
            if not sizes:
                orig_stem = os.path.splitext(orig_name)[0].lower()
                for k, v in size_map.items():
                    if os.path.splitext(k)[0].lower() == orig_stem:
                        sizes = v
                        break
            if not sizes:
                return tag

            w_px, h_px = sizes

            # 解析 Spire 当前写入的宽度（px），只做缩减不做放大，防止表格内图片尺寸漂移
            existing_w_px = None
            style_m_pre = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            if style_m_pre:
                sw_m = re.search(r'\bwidth\s*:\s*([\d.]+)(pt|px)', style_m_pre.group(1), re.IGNORECASE)
                if sw_m:
                    val = float(sw_m.group(1))
                    existing_w_px = round(val * 96 / 72) if sw_m.group(2).lower() == 'pt' else round(val)
            if existing_w_px is None:
                attr_w = re.search(r'\bwidth="(\d+)"', tag, re.IGNORECASE)
                if attr_w:
                    existing_w_px = int(attr_w.group(1))
            if existing_w_px is not None and w_px > existing_w_px:
                return tag

            tag = re.sub(r'\s+width="[^"]*"', '', tag)
            tag = re.sub(r'\s+height="[^"]*"', '', tag)
            tag = re.sub(r'(<img\b)', rf'\1 width="{w_px}" height="{h_px}"', tag)

            style_m = re.search(r'style="([^"]*)"', tag)
            if style_m:
                style_str = style_m.group(1)
                style_str = re.sub(r'\bwidth\s*:[^;]+;?', '', style_str, flags=re.IGNORECASE)
                style_str = re.sub(r'\bheight\s*:[^;]+;?', '', style_str, flags=re.IGNORECASE)
                style_str = style_str.rstrip('; ')
                new_style  = f"{style_str}; width:{w_px}px; height:{h_px}px".lstrip('; ')
                tag = tag[:style_m.start()] + f'style="{new_style}"' + tag[style_m.end():]
            else:
                tag = re.sub(r'(<img\b)', rf'\1 style="width:{w_px}px; height:{h_px}px"', tag)
                tag = re.sub(r'(<img\b)(.*?)(<img\b)', r'\1\2', tag)

            logger.debug(f"   📐 {src_basename} → {orig_name} 锁定尺寸 {w_px}×{h_px}px")
            return tag

        return re.sub(r'<img\b[^>]*>', _replace_img_tag, html_content, flags=re.IGNORECASE)

    # ------------------------------------------------------------------ #
    #  表格宽度修正（DOCX→HTML 方向）                                      #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _fix_html_table_widths(html_content: str,
                                content_width_pt: float = 400.0) -> str:
        """
        将 Spire 导出 HTML 中超出版心宽度的表格等比缩放至版心宽度以内。
        """

        def _to_pt(val_str: str) -> float | None:
            if not val_str:
                return None
            s = val_str.strip().lower()
            try:
                if s.endswith('pt'):  return float(s[:-2])
                if s.endswith('px'):  return float(s[:-2]) * 72 / 96
                if s.endswith('in'):  return float(s[:-2]) * 72
                if s.endswith('cm'):  return float(s[:-2]) / 2.54 * 72
                if s.endswith('mm'):  return float(s[:-2]) / 25.4 * 72
                return float(s)
            except ValueError:
                return None

        def _fmt(val_pt: float) -> str:
            return f"{val_pt:.2f}pt"

        def _style_width_pt(style_str: str) -> float | None:
            m = re.search(r'(?<![a-zA-Z-])width\s*:\s*([^;]+)', style_str, re.IGNORECASE)
            if m:
                return _to_pt(m.group(1).strip())
            return None

        def _replace_style_width(style_str: str, new_pt: float) -> str:
            return re.sub(
                r'(?<![a-zA-Z-])width\s*:\s*[^;]+',
                f'width:{_fmt(new_pt)}',
                style_str,
                flags=re.IGNORECASE
            )

        def _scale_tag_width(tag: str, ratio: float) -> str:
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            if style_m:
                style_str = style_m.group(1)
                w_pt = _style_width_pt(style_str)
                if w_pt is not None:
                    new_style = _replace_style_width(style_str, w_pt * ratio)
                    tag = tag[:style_m.start()] + f'style="{new_style}"' + tag[style_m.end():]
                    return tag

            attr_m = re.search(r'\bwidth="([^"]*)"', tag, re.IGNORECASE)
            if attr_m:
                w_pt = _to_pt(attr_m.group(1))
                if w_pt is not None:
                    new_val = _fmt(w_pt * ratio)
                    tag = tag[:attr_m.start()] + f'width="{new_val}"' + tag[attr_m.end():]

            return tag

        result_parts = []
        cursor = 0
        table_open_re  = re.compile(r'<table\b[^>]*>', re.IGNORECASE | re.DOTALL)
        table_close_re = re.compile(r'</table\s*>', re.IGNORECASE)
        cell_tag_re    = re.compile(r'<(?:col|td|th)\b[^>]*>', re.IGNORECASE | re.DOTALL)

        for tbl_open_m in table_open_re.finditer(html_content):
            result_parts.append(html_content[cursor:tbl_open_m.start()])

            tbl_open_tag = tbl_open_m.group(0)
            search_from  = tbl_open_m.end()

            outer_close_m = None
            depth2 = 1
            for m in re.finditer(r'<(/?)table\b[^>]*>', html_content[search_from:], re.IGNORECASE):
                if m.group(1) == '':
                    depth2 += 1
                else:
                    depth2 -= 1
                    if depth2 == 0:
                        abs_start = search_from + m.start()
                        abs_end   = search_from + m.end()
                        outer_close_m = (abs_start, abs_end)
                        break

            if outer_close_m:
                table_inner   = html_content[search_from : outer_close_m[0]]
                cursor_next   = outer_close_m[1]
            else:
                table_inner   = html_content[search_from:]
                cursor_next   = len(html_content)

            tbl_width_pt = None
            style_m = re.search(r'style="([^"]*)"', tbl_open_tag, re.IGNORECASE)
            if style_m:
                tbl_width_pt = _style_width_pt(style_m.group(1))
            if tbl_width_pt is None:
                attr_m = re.search(r'\bwidth="([^"]*)"', tbl_open_tag, re.IGNORECASE)
                if attr_m:
                    tbl_width_pt = _to_pt(attr_m.group(1))

            if tbl_width_pt is None or tbl_width_pt <= content_width_pt:
                result_parts.append(tbl_open_tag)
                result_parts.append(table_inner)
                result_parts.append('</table>')
                cursor = cursor_next
                continue

            ratio = content_width_pt / tbl_width_pt
            logger.debug(f"   📏 表格宽度 {tbl_width_pt:.1f}pt → {content_width_pt:.1f}pt（ratio={ratio:.4f}）")
            tbl_open_tag = _scale_tag_width(tbl_open_tag, ratio)

            def _scale_cell(cm):
                return _scale_tag_width(cm.group(0), ratio)

            table_inner_scaled = cell_tag_re.sub(_scale_cell, table_inner)

            result_parts.append(tbl_open_tag)
            result_parts.append(table_inner_scaled)
            result_parts.append('</table>')
            cursor = cursor_next

        result_parts.append(html_content[cursor:])
        return ''.join(result_parts)

    def _fix_html_img_sizes_for_import(self, html_text: str,
                                        page_width_px: int = 794,
                                        content_width_px: int = 400) -> str:
        """
        HTML→DOCX 方向的图片尺寸修正 。
        """
        import struct
        import urllib.request
        import urllib.error

        MM_PER_INCH  = 25.4
        PX_PER_INCH  = 96.0

        def _css_val_to_px(val_str):
            if not val_str:
                return None
            s = val_str.strip().lower()
            try:
                if s.endswith('px'):  return float(s[:-2])
                if s.endswith('pt'):  return float(s[:-2]) * PX_PER_INCH / 72
                if s.endswith('in'):  return float(s[:-2]) * PX_PER_INCH
                if s.endswith('cm'):  return float(s[:-2]) / 2.54 * PX_PER_INCH
                if s.endswith('mm'):  return float(s[:-2]) / MM_PER_INCH * PX_PER_INCH
                if s.endswith('%'):   return None
                return float(s)
            except ValueError:
                return None

        def _read_image_size_from_bytes(data: bytes):
            try:
                if data[:8] == b'\x89PNG\r\n\x1a\n':
                    w, h = struct.unpack('>II', data[16:24])
                    return w, h
                if data[:2] == b'\xff\xd8':
                    i = 2
                    while i < len(data) - 8:
                        if data[i] != 0xff:
                            break
                        marker = data[i+1]
                        length = struct.unpack('>H', data[i+2:i+4])[0]
                        if marker in (0xc0, 0xc1, 0xc2, 0xc3,
                                      0xc5, 0xc6, 0xc7, 0xc9, 0xca, 0xcb):
                            h, w = struct.unpack('>HH', data[i+5:i+9])
                            return w, h
                        i += 2 + length
                if data[:6] in (b'GIF87a', b'GIF89a'):
                    w, h = struct.unpack('<HH', data[6:10])
                    return w, h
                if data[:2] == b'BM':
                    w, h = struct.unpack('<II', data[18:26])
                    return w, abs(h)
                if data[:4] == b'RIFF' and data[8:12] == b'WEBP':
                    chunk = data[12:16]
                    if chunk == b'VP8 ':
                        w = struct.unpack('<H', data[26:28])[0] & 0x3fff
                        h = struct.unpack('<H', data[28:30])[0] & 0x3fff
                        return w, h
                    if chunk == b'VP8L':
                        bits = struct.unpack('<I', data[21:25])[0]
                        w = (bits & 0x3fff) + 1
                        h = ((bits >> 14) & 0x3fff) + 1
                        return w, h
                    if chunk == b'VP8X':
                        w = (struct.unpack('<I', data[24:27] + b'\x00')[0] & 0xffffff) + 1
                        h = (struct.unpack('<I', data[27:30] + b'\x00')[0] & 0xffffff) + 1
                        return w, h
            except Exception:
                pass
            return None, None

        def _read_image_size_from_file(path: str):
            try:
                with open(path, 'rb') as f:
                    data = f.read(512)
                return _read_image_size_from_bytes(data)
            except Exception:
                return None, None

        _url_cache: dict[str, tuple] = {}

        def _fetch_image_size_from_url(url: str):
            if url in _url_cache:
                return _url_cache[url]
            result = (None, None)
            for attempt, timeout in enumerate((30, 60), start=1):
                try:
                    req = urllib.request.Request(
                        url,
                        headers={'User-Agent': 'Mozilla/5.0', 'Range': 'bytes=0-4095'}
                    )
                    with urllib.request.urlopen(req, timeout=timeout) as resp:
                        data = resp.read(4096)
                    result = _read_image_size_from_bytes(data)
                    if result == (None, None):
                        req2 = urllib.request.Request(
                            url, headers={'User-Agent': 'Mozilla/5.0'}
                        )
                        with urllib.request.urlopen(req2, timeout=timeout) as resp2:
                            full = resp2.read()
                        result = _read_image_size_from_bytes(full)
                    if result != (None, None):
                        break
                except Exception as e:
                    logger.debug(f"   ⚠️ 获取图片尺寸失败（第{attempt}次，timeout={timeout}s）{url[:60]}：{e}")
            _url_cache[url] = result
            return result

        def _constrain(w_raw: float, h_raw: float):
            if w_raw > content_width_px:
                scale = content_width_px / w_raw
                return round(content_width_px), round(h_raw * scale)
            return round(w_raw), round(h_raw)

        def _apply_sizes(tag: str, w_px: int, h_px: int) -> str:
            tag = re.sub(r'\s+width="[^"]*"',  '', tag, flags=re.IGNORECASE)
            tag = re.sub(r'\s+height="[^"]*"', '', tag, flags=re.IGNORECASE)
            tag = re.sub(r'(<img\b)', rf'\1 width="{w_px}" height="{h_px}"', tag)

            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            if style_m:
                s = style_m.group(1)
                s = re.sub(r'\bmax-width\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                s = re.sub(r'\bmax-height\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                s = re.sub(r'\bwidth\s*:[^;]+;?',  '', s, flags=re.IGNORECASE)
                s = re.sub(r'\bheight\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                s = s.rstrip('; ')
                new_style = f"{s}; width:{w_px}px; height:{h_px}px".lstrip('; ')
                tag = tag[:style_m.start()] + f'style="{new_style}"' + tag[style_m.end():]
            else:
                tag = re.sub(r'(<img\b)',
                             rf'\1 style="width:{w_px}px; height:{h_px}px"', tag)
            return tag

        def _get_phys_size(src: str, tag: str):
            if src.startswith('data:'):
                try:
                    b64_part = src.split(',', 1)[1] if ',' in src else ''
                    import base64 as _b64
                    raw_bytes = _b64.b64decode(b64_part[:4096] + '==')
                    return _read_image_size_from_bytes(raw_bytes)
                except Exception:
                    return None, None
            elif os.path.exists(src):
                return _read_image_size_from_file(src)
            elif src.startswith('http://') or src.startswith('https://'):
                return _fetch_image_size_from_url(src)
            return None, None

        def _process_img(tag: str, max_w_px: float) -> str:
            src_m = re.search(r'src="([^"]+)"', tag, re.IGNORECASE)
            src   = src_m.group(1) if src_m else ''

            style_dict = {}
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE | re.DOTALL)
            if style_m:
                style_val = re.sub(r'[\r\n]+\s*', ' ', style_m.group(1))
                for part in style_val.split(';'):
                    if ':' in part:
                        k, _, v = part.partition(':')
                        style_dict[k.strip().lower()] = v.strip()

            raw_w = _css_val_to_px(style_dict.get('width'))
            raw_h = _css_val_to_px(style_dict.get('height'))

            attr_w_m = re.search(r'\bwidth="([^"]*)"',  tag, re.IGNORECASE)
            attr_h_m = re.search(r'\bheight="([^"]*)"', tag, re.IGNORECASE)
            if raw_w is None and attr_w_m:
                raw_w = _css_val_to_px(attr_w_m.group(1))
            if raw_h is None and attr_h_m:
                raw_h = _css_val_to_px(attr_h_m.group(1))

            h_val_str = style_dict.get('height', '').strip().lower()
            height_is_auto = (h_val_str in ('auto', '') or raw_h is None)

            if raw_w and raw_h and raw_w > 0 and raw_h > 0:
                if raw_w > max_w_px:
                    scale = max_w_px / raw_w
                    w_px, h_px = round(max_w_px), round(raw_h * scale)
                else:
                    w_px, h_px = round(raw_w), round(raw_h)
                logger.debug(f"   📐 [A] {round(raw_w)}×{round(raw_h)}px → {w_px}×{h_px}px（上限{round(max_w_px)}）")
                return _apply_sizes(tag, w_px, h_px)

            if raw_w and raw_w > 0 and height_is_auto:
                phys_w, phys_h = _get_phys_size(src, tag)
                if phys_w and phys_h and phys_w > 0:
                    computed_h = raw_w * phys_h / phys_w
                    if raw_w > max_w_px:
                        scale = max_w_px / raw_w
                        w_px, h_px = round(max_w_px), round(computed_h * scale)
                    else:
                        w_px, h_px = round(raw_w), round(computed_h)
                    logger.debug(f"   📐 [A½] width={round(raw_w)}px + height:auto → 物理{phys_w}×{phys_h} → {w_px}×{h_px}px")
                    return _apply_sizes(tag, w_px, h_px)
                else:
                    final_w = round(min(raw_w, max_w_px))
                    tag2 = re.sub(r'\s+height="[^"]*"', '', tag, flags=re.IGNORECASE)
                    style_m2 = re.search(r'style="([^"]*)"', tag2, re.IGNORECASE)
                    if style_m2:
                        s = style_m2.group(1)
                        s = re.sub(r'\bheight\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                        s = re.sub(r'\bmax-width\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                        s = re.sub(r'\bwidth\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                        s = s.rstrip('; ')
                        s = f"{s}; width:{final_w}px".lstrip('; ')
                        tag2 = tag2[:style_m2.start()] + f'style="{s}"' + tag2[style_m2.end():]
                    tag2 = re.sub(r'\s+width="[^"]*"', '', tag2, flags=re.IGNORECASE)
                    tag2 = re.sub(r'(<img\b)', rf'\1 width="{final_w}"', tag2)
                    logger.debug(f"   ⚠️ [A½-fallback] 无法获取物理像素，width={final_w}px，height交由Spire决定")
                    return tag2

            phys_w, phys_h = _get_phys_size(src, tag)

            if not phys_w or not phys_h:
                if 'max-width' in style_dict or 'max-height' in style_dict:
                    tag = _apply_sizes(tag, round(max_w_px), round(max_w_px))
                    logger.debug(f"   ⚠️ 无法获取物理像素，回退到上下文宽度：{src[:60]}")
                return tag

            if phys_w:
                logger.debug(f"   📐 [B/C] 物理 {phys_w}×{phys_h}：{src[:60]}")
            if float(phys_w) > max_w_px:
                scale = max_w_px / phys_w
                w_px, h_px = round(max_w_px), round(phys_h * scale)
            else:
                w_px, h_px = round(phys_w), round(phys_h)
            return _apply_sizes(tag, w_px, h_px)

        PT_TO_PX = PX_PER_INCH / 72

        def _td_max_w_px(td_tag: str) -> float:
            style_dict_td = {}
            sm = re.search(r'style="([^"]*)"', td_tag, re.IGNORECASE)
            if sm:
                for part in sm.group(1).split(';'):
                    if ':' in part:
                        k, _, v = part.partition(':')
                        style_dict_td[k.strip().lower()] = v.strip()

            td_w_px = _css_val_to_px(style_dict_td.get('width'))
            if td_w_px is None:
                attr_m = re.search(r'\bwidth="([^"]*)"', td_tag, re.IGNORECASE)
                if attr_m:
                    td_w_px = _css_val_to_px(attr_m.group(1))

            if not td_w_px:
                return float(content_width_px)

            padding_str = style_dict_td.get('padding', '').strip()
            padding_px  = 0.0
            if padding_str:
                parts = padding_str.split()
                if len(parts) == 1:
                    p = _css_val_to_px(parts[0]) or 0
                    padding_px = p * 2
                elif len(parts) == 2:
                    p = _css_val_to_px(parts[1]) or 0
                    padding_px = p * 2
                elif len(parts) >= 4:
                    pl = _css_val_to_px(parts[3]) or 0
                    pr = _css_val_to_px(parts[1]) or 0
                    padding_px = pl + pr
                elif len(parts) == 3:
                    p = _css_val_to_px(parts[1]) or 0
                    padding_px = p * 2

            avail = td_w_px - padding_px
            return max(avail, 40.0)

        def _replace_img_global(m):
            return _process_img(m.group(0), float(content_width_px))

        IMG_RE    = re.compile(r'<img\b[^>]*>', re.IGNORECASE | re.DOTALL)
        TABLE_RE  = re.compile(r'<table\b[^>]*>.*?</table\s*>', re.IGNORECASE | re.DOTALL)
        TD_RE     = re.compile(r'(<t[dh]\b[^>]*>)(.*?)(?=<t[dh]\b|</tr|</table)', re.IGNORECASE | re.DOTALL)

        result_parts = []
        cursor = 0

        for tbl_m in TABLE_RE.finditer(html_text):
            before = html_text[cursor:tbl_m.start()]
            result_parts.append(IMG_RE.sub(_replace_img_global, before))

            tbl_html   = tbl_m.group(0)
            tbl_result = []
            td_cursor  = 0

            for td_m in TD_RE.finditer(tbl_html):
                tbl_result.append(tbl_html[td_cursor:td_m.start()])
                td_open    = td_m.group(1)
                td_content = td_m.group(2)
                max_w      = _td_max_w_px(td_open)

                def _make_td_replacer(mw):
                    def _r(im):
                        return _process_img(im.group(0), mw)
                    return _r

                td_content_fixed = IMG_RE.sub(_make_td_replacer(max_w), td_content)
                tbl_result.append(td_open)
                tbl_result.append(td_content_fixed)
                td_cursor = td_m.end()

            tbl_result.append(tbl_html[td_cursor:])
            result_parts.append(''.join(tbl_result))
            cursor = tbl_m.end()

        result_parts.append(IMG_RE.sub(_replace_img_global, html_text[cursor:]))
        return ''.join(result_parts)

    # ------------------------------------------------------------------ #
    #  【修复核心】图片内嵌：二进制精确匹配策略                             #
    # ------------------------------------------------------------------ #

    def _embed_images_to_html(self, html_path, image_display_order, original_img_dir):
        """
        对已生成的 HTML 文件做图片 base64 内嵌（in-place）。

        【修复图片顺序错位问题】
        旧方案依赖 spire文件名→原始文件名 的顺序/名称映射，当 Spire 对同一张图
        在页眉、页脚、正文中生成多个 <img> 标签时，映射表只有一条记录，
        导致第二次以后的引用无法命中，最终图片内容全部错位。

        新方案采用三级匹配策略，按优先级依次尝试：
          1. 二进制精确匹配：读取 Spire 生成的图片文件内容，与 original_img_dir
             中所有原始图片做逐字节比对，命中则使用原始无压缩图片的 base64。
             此策略完全不依赖文件名或顺序，对同一张图被引用 N 次的情况也正确。
          2. 文件名 stem 匹配（去扩展名不区分大小写）：
             Spire 修改了扩展名但文件名主体未变时使用。
          3. 兜底：直接使用 Spire 生成的图片文件做 base64（保留原有兜底行为）。
        """
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        html_content = html_content.replace('\\', '/')

        # ── 1. 预加载原始图片（bytes + base64）──────────────────────────
        orig_bytes_map = {}   # orig_filename → bytes
        orig_b64_map   = {}   # orig_filename → base64 data URI
        for fname in os.listdir(original_img_dir):
            fpath = os.path.join(original_img_dir, fname)
            if not os.path.isfile(fpath):
                continue
            try:
                with open(fpath, 'rb') as f:
                    data = f.read()
                b64 = self._image_to_base64(fpath)
                if b64:
                    orig_bytes_map[fname] = data
                    orig_b64_map[fname]   = b64
            except Exception as e:
                logger.debug(f"   ⚠️ 预加载原始图片失败 {fname}：{e}")
        if not orig_b64_map:
            logger.warning("⚠️ 原始图片目录为空，跳过内嵌")
            return

        # ── 2. 构建文件名 stem 查找表（用于第二级匹配）──────────────────
        orig_stem_map = {}   # stem(小写) → orig_filename
        for fname in orig_b64_map:
            stem = os.path.splitext(fname)[0].lower()
            orig_stem_map[stem] = fname

        # ── 3. 逐个 <img> 标签做替换 ─────────────────────────────────────
        img_src_re = re.compile(r'<img\b[^>]*>', re.IGNORECASE | re.DOTALL)

        def _replace_img(m):
            tag   = m.group(0)
            src_m = re.search(r'src="([^"]+)"', tag)
            if not src_m:
                return tag

            src = src_m.group(1)
            if src.startswith('data:'):
                return tag  # 已内嵌，跳过

            spire_fname = os.path.basename(src.replace('\\', '/'))

            # 构建 Spire 图片绝对路径
            abs_src = self._normalize_path(src)
            if not os.path.isabs(abs_src):
                abs_src = self._normalize_path(
                    os.path.join(os.path.dirname(html_path), src)
                )

            matched_b64 = None

            # ── 第一级：二进制精确匹配 ────────────────────────────────────
            if os.path.exists(abs_src):
                try:
                    with open(abs_src, 'rb') as f:
                        spire_bytes = f.read()
                    for orig_fname, orig_bytes in orig_bytes_map.items():
                        if spire_bytes == orig_bytes:
                            matched_b64 = orig_b64_map[orig_fname]
                            logger.debug(f"🔄 {spire_fname} 二进制匹配 → {orig_fname}（原始无压缩）")
                            break
                except Exception as e:
                    logger.debug(f"   ⚠️ 读取 Spire 图片失败 {spire_fname}：{e}")
            # ── 第二级：文件名 stem 匹配 ──────────────────────────────────
            if matched_b64 is None:
                spire_stem = os.path.splitext(spire_fname)[0].lower()
                if spire_stem in orig_stem_map:
                    orig_fname  = orig_stem_map[spire_stem]
                    matched_b64 = orig_b64_map[orig_fname]
                    logger.debug(f"🔄 {spire_fname} 文件名匹配 → {orig_fname}")
            # ── 第三级：兜底，直接用 Spire 生成的图片 ────────────────────
            if matched_b64 is None and os.path.exists(abs_src):
                matched_b64 = self._image_to_base64(abs_src)
                if matched_b64:
                    logger.warning(f"⚠️ {spire_fname} 无原始匹配，使用 Spire 生成图片（兜底）")
            if matched_b64:
                new_tag = tag[:src_m.start()] + f'src="{matched_b64}"' + tag[src_m.end():]
                return new_tag

            logger.debug(f"   ⚠️ {spire_fname} 所有匹配均失败，保留原 src")
            return tag

        html_content = img_src_re.sub(_replace_img, html_content)

        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        logger.debug("✅ 图片内嵌完成（二进制精确匹配）")
    # ------------------------------------------------------------------ #
    #  分片工具                                                             #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _pydocx_collect_body_elements(docx_path):
        """
        用 python-docx 读取文档，返回 body 下所有顶层元素列表。
        """
        doc = PythonDocx(docx_path)
        body = doc.element.body
        elements = [
            el for el in body
            if el.tag in (qn('w:p'), qn('w:tbl'))
        ]
        logger.debug(f"[python-docx 诊断] body 顶层元素总数: {len(elements)}")
        para_cnt  = sum(1 for e in elements if e.tag == qn('w:p'))
        table_cnt = sum(1 for e in elements if e.tag == qn('w:tbl'))
        logger.debug(f"[python-docx 诊断] 段落: {para_cnt}, 表格: {table_cnt}")
        return doc, elements

    @staticmethod
    def _pydocx_count_table_paras(tbl_el):
        """统计 w:tbl 元素内所有 w:p 的数量"""
        return len(tbl_el.findall('.//' + qn('w:p')))

    @staticmethod
    def _pydocx_copy_sectPr(src_doc, dst_doc):
        """
        将源文档的页面设置（w:sectPr）复制到目标文档，
        同时彻底移除页眉页脚相关内容。
        """
        HF_TAGS = {
            qn('w:headerReference'),
            qn('w:footerReference'),
            qn('w:titlePg'),
        }

        src_body = src_doc.element.body
        src_sectPr = src_body.find(qn('w:sectPr'))
        if src_sectPr is None:
            last_p = src_body.findall(qn('w:p'))
            if last_p:
                pPr = last_p[-1].find(qn('w:pPr'))
                if pPr is not None:
                    src_sectPr = pPr.find(qn('w:sectPr'))

        if src_sectPr is not None:
            sectPr_copy = copy.deepcopy(src_sectPr)

            for tag in HF_TAGS:
                for node in sectPr_copy.findall(tag):
                    sectPr_copy.remove(node)

            pgMar = sectPr_copy.find(qn('w:pgMar'))
            if pgMar is not None:
                pgMar.attrib.pop(qn('w:header'), None)
                pgMar.attrib.pop(qn('w:footer'), None)

            dst_body = dst_doc.element.body
            old = dst_body.find(qn('w:sectPr'))
            if old is not None:
                dst_body.remove(old)
            dst_body.append(sectPr_copy)

    @staticmethod
    def _pydocx_new_doc(src_doc):
        """
        创建一个继承源文档样式、编号和页面设置的空白 python-docx Document。
        """
        dst = PythonDocx()
        dst.element.body.clear()

        try:
            src_styles_xml = src_doc.part.styles._element
            dst.part.styles._element.clear()
            for child in src_styles_xml:
                dst.part.styles._element.append(copy.deepcopy(child))
        except Exception as e:
            logger.warning(f"⚠️ 样式复制失败（使用默认样式）：{e}")
        try:
            src_numbering_part = src_doc.part.numbering_part
            if src_numbering_part is not None:
                rel_type = (
                    'http://schemas.openxmlformats.org/officeDocument/2006/'
                    'relationships/numbering'
                )
                dst_part = dst.part
                existing_num_part = None
                for rel in dst_part.rels.values():
                    if rel.reltype == rel_type:
                        existing_num_part = rel.target_part
                        break
                if existing_num_part is not None:
                    existing_num_part._element.clear()
                    for child in src_numbering_part._element:
                        existing_num_part._element.append(copy.deepcopy(child))
        except Exception as e:
            logger.warning(f"⚠️ 编号定义复制失败（忽略）：{e}")
        DocxHtmlConverter._pydocx_copy_sectPr(src_doc, dst)
        return dst

    @staticmethod
    def _inject_resources_into_chunk(src_docx_path, chunk_docx_path,
                                     referenced_images=None):
        """
        通过 zipfile 将源文档的资源精确注入已保存的 chunk docx。
        """
        RELS_PATH = 'word/_rels/document.xml.rels'

        with zipfile.ZipFile(chunk_docx_path, 'r') as zf:
            chunk_doc_xml  = zf.read('word/document.xml').decode('utf-8')
            chunk_rels_xml = zf.read(RELS_PATH).decode('utf-8') if RELS_PATH in zf.namelist() else ''
            chunk_files    = set(zf.namelist())

        doc_rids = set(
            re.findall(r'r:(?:id|embed|link)="(rId\d+)"', chunk_doc_xml)
        ) | set(
            re.findall(r'relationships}(?:id|embed|link)="(rId\d+)"', chunk_doc_xml)
        )
        existing_rids = set(re.findall(r'Id="(rId\d+)"', chunk_rels_xml))
        missing_rids  = doc_rids - existing_rids

        if not missing_rids:
            return

        with zipfile.ZipFile(src_docx_path, 'r') as zf:
            src_rels_xml = zf.read(RELS_PATH).decode('utf-8') if RELS_PATH in zf.namelist() else ''
            src_namelist = zf.namelist()

            vml_image_rids = set()
            for fname in src_namelist:
                if re.match(r'word/vmlDrawing\d*\.vml', fname):
                    try:
                        vml_content = zf.read(fname).decode('utf-8', errors='replace')
                        vml_image_rids.update(
                            re.findall(r'r:(?:id|href)="(rId\d+)"', vml_content)
                        )
                    except Exception:
                        pass

            src_rels_map = {}
            for m in re.finditer(r'<Relationship[^>]*/>', src_rels_xml):
                entry  = m.group(0)
                rid_m  = re.search(r'Id="(rId\d+)"', entry)
                tgt_m  = re.search(r'Target="([^"]+)"', entry)
                if rid_m and tgt_m:
                    rid    = rid_m.group(1)
                    target = tgt_m.group(1)
                    zip_path = 'word/' + target if not target.startswith('/') else target.lstrip('/')
                    src_rels_map[rid] = (entry, zip_path)

            entries_to_add = []
            files_to_copy  = {}

            HF_REL_TYPES = ('header', 'footer')

            for rid in sorted(missing_rids):
                if rid not in src_rels_map:
                    logger.debug(f"   ⚠️ rId={rid} 在源文档 rels 中也找不到，跳过")
                    continue

                entry, zip_path = src_rels_map[rid]

                if any(hf in entry.lower() for hf in HF_REL_TYPES):
                    continue

                if zip_path.startswith('word/media/'):
                    fname = os.path.basename(zip_path)
                    entry_lower  = entry.lower()
                    target_lower = zip_path.lower()
                    is_ole_or_vml = (
                        'oleobject'   in entry_lower or
                        'vmldrawing'  in entry_lower or
                        'ole'         in target_lower or
                        'vml'         in target_lower or
                        rid in vml_image_rids
                    )
                    if (referenced_images is not None
                            and fname not in referenced_images
                            and not is_ole_or_vml):
                        continue

                entries_to_add.append(entry)

                if zip_path in src_namelist and zip_path not in chunk_files:
                    files_to_copy[zip_path] = zf.read(zip_path)

        if not entries_to_add and not files_to_copy:
            return

        tmp_path = chunk_docx_path + '.tmp'
        with zipfile.ZipFile(chunk_docx_path, 'r') as src_zip, \
             zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as dst_zip:

            for item in src_zip.infolist():
                data = src_zip.read(item.filename)

                if item.filename == RELS_PATH and entries_to_add:
                    updated = data.decode('utf-8').replace(
                        '</Relationships>',
                        '\n'.join(entries_to_add) + '\n</Relationships>'
                    )
                    dst_zip.writestr(item, updated.encode('utf-8'))

                elif item.filename == 'word/settings.xml':
                    settings = data.decode('utf-8')
                    settings = re.sub(r'<w:evenAndOddHeaders\s*/>', '', settings)
                    dst_zip.writestr(item, settings.encode('utf-8'))

                else:
                    dst_zip.writestr(item, data)

            for zip_path, file_data in files_to_copy.items():
                dst_zip.writestr(zip_path, file_data)

        os.replace(tmp_path, chunk_docx_path)
        logger.debug(f"   💉 注入 {len(entries_to_add)} 条 rels + {len(files_to_copy)} 个资源文件")
    @staticmethod
    def _pydocx_append_element(dst_doc, el, src_doc=None):
        """将元素深拷贝追加到目标文档 body（sectPr 前）。"""
        dst_body = dst_doc.element.body
        sect_pr  = dst_body.find(qn('w:sectPr'))
        cloned   = copy.deepcopy(el)
        if sect_pr is not None:
            sect_pr.addprevious(cloned)
        else:
            dst_body.append(cloned)

    @staticmethod
    def _pydocx_make_marker_para(group_id):
        """
        生成一个携带 TABLE_SPLIT_MARKER 的段落 lxml Element。
        字体 1pt、白色，尽量不影响视觉排版。
        """
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'),  '0')
        pPr.append(spacing)
        p.append(pPr)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '2')
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '2')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FFFFFF')
        rPr.extend([sz, szCs, color])
        r.append(rPr)

        t = OxmlElement('w:t')
        t.text = f'TABLE_SPLIT_MARKER::{group_id}'
        r.append(t)
        p.append(r)
        return p

    def _sanitize_html_styles_for_import(self, html: str) -> str:
        """
        [针对 64% 缩放问题的专项修复]
        1. 表格：强制 100% 宽度，消除 440pt 导致的挤压。
        2. 图片：如果只有单维度，通过 style="width:auto; height:auto" 配合属性清除，
           强制 Spire 走等比缩放逻辑，不触发 96->72 的二次缩放。
        """
        # --- 修复表格：从 439.4pt 固宽改为 100% 自适应 ---
        html = re.sub(r'(<table[^>]*style="[^"]*)\bwidth\s*:\s*[\d.]+pt;?', r'\1 width:100%;', html, flags=re.IGNORECASE)
        html = re.sub(r'(<table[^>]*)\bwidth="\d+(?:\.\d+)?"', r'\1 width="100%"', html, flags=re.IGNORECASE)

        # --- 修复图片：防止 64% 比例偏移 ---
        def _fix_img_ratio(m):
            tag = m.group(0)
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            if not style_m:
                return tag

            style_str = style_m.group(1)
            # 提取 style 中的宽高（pt 或 px）
            sw_m = re.search(r'\bwidth\s*:\s*([\d.]+)(pt|px)', style_str, re.IGNORECASE)
            sh_m = re.search(r'\bheight\s*:\s*([\d.]+)(pt|px)', style_str, re.IGNORECASE)

            # 策略：如果只有宽度（如 146pt），必须显式把 height 设为 auto 并移除 HTML 属性
            # 这样 Spire 才会根据图片原始比例计算，而不会去读那个 height="113" 的旧像素值
            if sw_m and not sh_m:
                tag = re.sub(r'\s+height="[^"]*"', '', tag, flags=re.IGNORECASE)  # 移除属性
                style_str = re.sub(r'\bheight\s*:[^;]+;?', '', style_str, flags=re.IGNORECASE)
                style_str = style_str.rstrip('; ') + "; height:auto !important;"

            elif sh_m and not sw_m:
                tag = re.sub(r'\s+width="[^"]*"', '', tag, flags=re.IGNORECASE)  # 移除属性
                style_str = re.sub(r'\bwidth\s*:[^;]+;?', '', style_str, flags=re.IGNORECASE)
                style_str = style_str.rstrip('; ') + "; width:auto !important;"

            return re.sub(r'style="([^"]*)"', f'style="{style_str}"', tag, count=1, flags=re.IGNORECASE)

        html = re.sub(r'<img\b[^>]*>', _fix_img_ratio, html, flags=re.IGNORECASE | re.DOTALL)
        return html

    def _needs_chunking(self, docx_path):
        """
        【内部方法】用 python-docx 统计段落/表格数，判断是否需要分片。
        """
        _, elements = self._pydocx_collect_body_elements(docx_path)
        total_paras  = sum(1 for e in elements if e.tag == qn('w:p'))
        total_tables = sum(1 for e in elements if e.tag == qn('w:tbl'))
        needs = total_paras > self.MAX_PARAGRAPHS or total_tables > self.MAX_TABLES
        logger.debug(f"📊 文档规模：{total_paras} 段落，{total_tables} 表格，{'需要' if needs else '无需'}分片")
        return needs

    @staticmethod
    def _pydocx_sanitize_element(el, src_rels_rids):
        """
        处理元素内的 w:object（OLE 嵌入对象）：
        保留 v:shape（含预览图），移除 o:OLEObject 本体节点。
        """
        O_NS     = 'urn:schemas-microsoft-com:office:office'
        W_NS     = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        VML_NS   = 'urn:schemas-microsoft-com:vml'
        R_NS     = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        RID_ATTR = f'{{{R_NS}}}id'

        def _css_dim_to_twip(val_str):
            if not val_str:
                return None
            val_str = val_str.strip().lower()
            try:
                if val_str.endswith('pt'):
                    return round(float(val_str[:-2]) * 20)
                if val_str.endswith('in'):
                    return round(float(val_str[:-2]) * 72 * 20)
                if val_str.endswith('cm'):
                    return round(float(val_str[:-2]) * 28.3465 * 20)
                if val_str.endswith('mm'):
                    return round(float(val_str[:-2]) * 2.83465 * 20)
                if val_str.endswith('px'):
                    return round(float(val_str[:-2]) * 0.75 * 20)
            except (ValueError, AttributeError):
                pass
            return None

        def _twip_to_pt_str(twip):
            return f"{twip / 20:.1f}pt"

        def _parse_style(style_str):
            order, d = [], {}
            for part in style_str.split(';'):
                part = part.strip()
                if not part:
                    continue
                if ':' in part:
                    k, _, v = part.partition(':')
                    k, v = k.strip(), v.strip()
                    d[k] = v
                    order.append(k)
                else:
                    d[part] = None
                    order.append(part)
            return order, d

        def _serialize_style(order, d):
            parts = []
            for k in order:
                parts.append(f"{k}:{d[k]}" if d[k] is not None else k)
            return ';'.join(parts)

        cloned = copy.deepcopy(el)

        for obj in cloned.findall('.//' + qn('w:object')):
            ole_node = obj.find(f'{{{O_NS}}}OLEObject')
            if ole_node is None:
                continue

            ole_rid = ole_node.get(RID_ATTR)

            v_shape = obj.find(f'{{{VML_NS}}}shape')
            shape_w_twip = None
            shape_h_twip = None
            style_order, style_dict = [], {}

            if v_shape is not None:
                style_order, style_dict = _parse_style(v_shape.get('style', ''))
                shape_w_twip = _css_dim_to_twip(style_dict.get('width'))
                shape_h_twip = _css_dim_to_twip(style_dict.get('height'))

            dxa_attr = f'{{{W_NS}}}dxaOrig'
            dya_attr = f'{{{W_NS}}}dyaOrig'
            dxa_orig = obj.get(dxa_attr)
            dya_orig = obj.get(dya_attr)

            try:
                dxa_twip = int(dxa_orig) if dxa_orig else None
            except ValueError:
                dxa_twip = None
            try:
                dya_twip = int(dya_orig) if dya_orig else None
            except ValueError:
                dya_twip = None

            if shape_w_twip is None:
                shape_w_twip = dxa_twip
            if shape_h_twip is None:
                shape_h_twip = dya_twip

            log_parts = []

            if v_shape is not None and (shape_w_twip or shape_h_twip):
                modified = False

                if shape_w_twip and (not style_dict.get('width') or
                                     style_dict['width'] in ('0', '0pt', 'auto')):
                    style_dict['width'] = _twip_to_pt_str(shape_w_twip)
                    if 'width' not in style_order:
                        style_order.insert(0, 'width')
                    modified = True

                if shape_h_twip and (not style_dict.get('height') or
                                     style_dict['height'] in ('0', '0pt', 'auto')):
                    style_dict['height'] = _twip_to_pt_str(shape_h_twip)
                    if 'height' not in style_order:
                        idx = (style_order.index('width') + 1
                               if 'width' in style_order else 0)
                        style_order.insert(idx, 'height')
                    modified = True

                if modified:
                    v_shape.set('style', _serialize_style(style_order, style_dict))
                    log_parts.append(
                        f"style补全 {style_dict.get('width')}×{style_dict.get('height')}"
                    )

            if shape_w_twip and str(shape_w_twip) != (dxa_orig or ''):
                obj.set(dxa_attr, str(shape_w_twip))
                log_parts.append(f"dxaOrig→{shape_w_twip}")
            if shape_h_twip and str(shape_h_twip) != (dya_orig or ''):
                obj.set(dya_attr, str(shape_h_twip))
                log_parts.append(f"dyaOrig→{shape_h_twip}")

            para_el = obj.getparent()
            while para_el is not None and para_el.tag != f'{{{W_NS}}}p':
                para_el = para_el.getparent()

            if para_el is not None:
                pPr = para_el.find(f'{{{W_NS}}}pPr')
                if pPr is not None:
                    frame_pr = pPr.find(f'{{{W_NS}}}framePr')
                    if frame_pr is not None:
                        fw_attr = f'{{{W_NS}}}w'
                        fh_attr = f'{{{W_NS}}}h'

                        if shape_w_twip:
                            old_fw = frame_pr.get(fw_attr)
                            if old_fw != str(shape_w_twip):
                                frame_pr.set(fw_attr, str(shape_w_twip))
                                log_parts.append(f"framePr.w {old_fw}→{shape_w_twip}")

                        if shape_h_twip:
                            old_fh = frame_pr.get(fh_attr)
                            if old_fh != str(shape_h_twip):
                                frame_pr.set(fh_attr, str(shape_h_twip))
                                log_parts.append(f"framePr.h {old_fh}→{shape_h_twip}")

            obj.remove(ole_node)

            summary = f"（{', '.join(log_parts)}）" if log_parts else "（尺寸无变化）"
            logger.debug(f"   ✂️ 移除 OLEObject 本体（r:id={ole_rid}）{summary}")
        return cloned

    def _split_docx_to_chunks(self, docx_path, chunk_dir, image_display_order=None):
        """
        【内部方法】用 python-docx + lxml 将大文档拆分为多个子 DOCX。
        """
        os.makedirs(chunk_dir, exist_ok=True)

        src_doc, elements = self._pydocx_collect_body_elements(docx_path)
        total_elements = len(elements)

        src_rels_rids = set()
        try:
            with zipfile.ZipFile(docx_path, 'r') as zf:
                if 'word/_rels/document.xml.rels' in zf.namelist():
                    rels_xml = zf.read('word/_rels/document.xml.rels').decode('utf-8')
                    for m in re.finditer(r'<Relationship[^>]*/>', rels_xml):
                        entry = m.group(0)
                        if 'oleObject' in entry:
                            rid_m = re.search(r'Id="(rId\d+)"', entry)
                            if rid_m:
                                src_rels_rids.add(rid_m.group(1))
        except Exception as e:
            logger.warning(f"⚠️ 读取源文档 rels 失败：{e}")
        chunk_paths = []
        chunk_idx   = 0
        para_count  = 0
        table_count = 0
        dst_doc     = self._pydocx_new_doc(src_doc)

        def _save_chunk(doc, idx, p_cnt, t_cnt):
            path = self._normalize_path(
                os.path.join(chunk_dir, f"chunk_{idx:04d}.docx")
            )
            doc.save(path)
            ref_imgs = set(image_display_order) if image_display_order else None
            self._inject_resources_into_chunk(docx_path, path, ref_imgs)
            logger.debug(f"✅ 切片 chunk_{idx:04d}：{p_cnt} 段落，{t_cnt} 表格")
            return path

        def _flush_current_chunk():
            nonlocal chunk_idx, para_count, table_count, dst_doc
            if para_count > 0 or table_count > 0:
                chunk_paths.append(_save_chunk(dst_doc, chunk_idx, para_count, table_count))
                chunk_idx  += 1
            dst_doc     = self._pydocx_new_doc(src_doc)
            para_count  = 0
            table_count = 0

        i = 0
        while i < total_elements:
            el = elements[i]

            if el.tag == qn('w:tbl'):
                inner_paras = self._pydocx_count_table_paras(el)
                rows        = el.findall(qn('w:tr'))
                total_rows  = len(rows)

                if (para_count + inner_paras <= self.MAX_PARAGRAPHS and
                        table_count + 1 <= self.MAX_TABLES):
                    clean_tbl = self._pydocx_sanitize_element(el, src_rels_rids)
                    self._pydocx_append_element(dst_doc, clean_tbl)
                    para_count  += inner_paras
                    table_count += 1
                    i += 1
                else:
                    _flush_current_chunk()

                    split_group_id = uuid.uuid4().hex[:8]
                    row_cursor     = 0

                    tbl_pr   = el.find(qn('w:tblPr'))
                    tbl_grid = el.find(qn('w:tblGrid'))

                    while row_cursor < total_rows:
                        split_doc = self._pydocx_new_doc(src_doc)

                        split_doc.element.body.insert(
                            0, self._pydocx_make_marker_para(split_group_id)
                        )

                        sub_tbl = OxmlElement('w:tbl')
                        if tbl_pr is not None:
                            sub_tbl.append(copy.deepcopy(tbl_pr))
                        if tbl_grid is not None:
                            sub_tbl.append(copy.deepcopy(tbl_grid))

                        rows_in_chunk     = 0
                        chunk_inner_paras = 0

                        while row_cursor < total_rows:
                            row = rows[row_cursor]
                            row_paras = len(row.findall('.//' + qn('w:p')))
                            if (chunk_inner_paras + row_paras > self.MAX_PARAGRAPHS
                                    and rows_in_chunk > 0):
                                break
                            clean_row = self._pydocx_sanitize_element(row, src_rels_rids)
                            sub_tbl.append(clean_row)
                            chunk_inner_paras += row_paras
                            rows_in_chunk     += 1
                            row_cursor        += 1

                        self._pydocx_append_element(split_doc, sub_tbl, src_doc)

                        split_path = self._normalize_path(
                            os.path.join(chunk_dir, f"chunk_{chunk_idx:04d}.docx")
                        )
                        split_doc.save(split_path)
                        ref_imgs = set(image_display_order) if image_display_order else None
                        self._inject_resources_into_chunk(docx_path, split_path, ref_imgs)
                        logger.debug(f"✅ 切片 chunk_{chunk_idx:04d}：（表格分片 {split_group_id}，{rows_in_chunk} 行）")
                        chunk_paths.append(split_path)
                        chunk_idx += 1

                    i += 1

            else:
                if para_count + 1 > self.MAX_PARAGRAPHS and (para_count > 0 or table_count > 0):
                    _flush_current_chunk()

                clean_el = self._pydocx_sanitize_element(el, src_rels_rids)
                self._pydocx_append_element(dst_doc, clean_el)
                para_count += 1
                i += 1

        if para_count > 0 or table_count > 0:
            chunk_paths.append(_save_chunk(dst_doc, chunk_idx, para_count, table_count))

        logger.debug(f"📦 共切分为 {len(chunk_paths)} 个子文档")
        return chunk_paths

    # ------------------------------------------------------------------ #
    #  HTML 合并工具                                                        #
    # ------------------------------------------------------------------ #

    def _clean_header_footer(self, html_content):
        """
        去除页眉页脚相关元素。
        使用基于嵌套深度计数的方式精确提取匹配的 </div>，
        避免正则贪婪匹配误删正文内容。
        """
        spire_hf_pattern = re.compile(
            r'<div[^>]*-spr-headerfooter-type[^>]*>',
            re.IGNORECASE
        )
        generic_hf_pattern = re.compile(
            r'<div[^>]*(?:class|id)\s*=\s*["\'][^"\']*(?:header|footer)[^"\']*["\'][^>]*>',
            re.IGNORECASE
        )

        def _remove_div_block(content, pattern):
            result = []
            pos = 0
            for m in pattern.finditer(content):
                result.append(content[pos:m.start()])
                depth  = 1
                cursor = m.end()
                while cursor < len(content) and depth > 0:
                    open_m  = re.search(r'<div[^>]*>', content[cursor:], re.IGNORECASE)
                    close_m = re.search(r'</div\s*>', content[cursor:], re.IGNORECASE)
                    if close_m is None:
                        break
                    if open_m and open_m.start() < close_m.start():
                        depth  += 1
                        cursor += open_m.end()
                    else:
                        depth  -= 1
                        cursor += close_m.end()
                pos = cursor
            result.append(content[pos:])
            return ''.join(result)

        html_content = _remove_div_block(html_content, spire_hf_pattern)
        html_content = _remove_div_block(html_content, generic_hf_pattern)
        return html_content

    def _docx_to_html_no_embed(self, docx_path, html_path):
        """
        【内部方法】DOCX转HTML，图片保留为文件引用，不做base64内嵌。
        专供分片流程使用。
        """
        docx_path = self._normalize_path(docx_path)
        html_path = self._normalize_path(html_path)
        html_dir = os.path.dirname(html_path)

        spire_img_dir = self._normalize_path(
            os.path.join(html_dir, f"img_{uuid.uuid4().hex[:8]}")
        )

        document = Document()
        try:
            document.LoadFromFile(docx_path)
            document.HtmlExportOptions.ImageEmbedded = False
            document.HtmlExportOptions.ImagesPath = spire_img_dir
            document.HtmlExportOptions.ImageFormat = self.default_image_format
            document.SaveToFile(html_path, FileFormat.Html)
        except Exception as e:
            logger.error(f"❌ Spire转换HTML失败：{e}")
            return False
        finally:
            document.Close()
            del document

        css_file_path = self._normalize_path(os.path.splitext(html_path)[0] + '_styles.css')
        if os.path.exists(css_file_path):
            with open(css_file_path, 'r', encoding='utf-8') as f:
                css_content = f.read()
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            html_content = re.sub(r'<link[^>]*href="[^"]+\.css"[^>]*>', '', html_content)
            html_content = html_content.replace(
                '</head>',
                f'<style type="text/css">\n{css_content}\n</style>\n</head>'
            )
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            os.remove(css_file_path)

        logger.debug(f"✅ chunk转换完成（图片未内嵌）：{os.path.basename(html_path)}")
        return True

    def _merge_html_files_to_disk(self, chunk_html_paths, output_path):
        """
        【内部方法】流式合并多个chunk HTML为一个完整HTML文件。
        """
        if not chunk_html_paths:
            return

        # 预收集 chunk 1+ 的 CSS，合并到第一个 chunk 的 <head> 中
        extra_css_blocks = []
        for extra_path in chunk_html_paths[1:]:
            try:
                with open(extra_path, 'r', encoding='utf-8') as ef:
                    extra_content = ef.read()
                for style_m in re.finditer(r'<style[^>]*>(.*?)</style>', extra_content,
                                           re.DOTALL | re.IGNORECASE):
                    css = style_m.group(1).strip()
                    if css:
                        extra_css_blocks.append(css)
            except Exception:
                pass

        marker_re = re.compile(
            r'<p[^>]*>\s*<span[^>]*>TABLE_SPLIT_MARKER::([a-f0-9]{8})</span>\s*</p>',
            re.IGNORECASE | re.DOTALL
        )

        pending_table_group = None
        pending_table_open  = None
        pending_trs         = []

        def _flush_pending_table(out_f):
            nonlocal pending_table_group, pending_table_open, pending_trs
            if pending_table_group:
                out_f.write(f"{pending_table_open}\n")
                for trs in pending_trs:
                    out_f.write(trs)
                out_f.write("</table>\n")
                pending_table_group = None
                pending_table_open  = None
                pending_trs         = []

        def _extract_trs(text, search_start=0):
            text_stripped = re.sub(r'^\s*<div[^>]*>\s*', '', text, flags=re.IGNORECASE)
            tbl_open  = re.search(r'<table[^>]*>', text_stripped, re.IGNORECASE)
            tbl_close = re.search(r'</table>', text_stripped, re.IGNORECASE)
            if tbl_open and tbl_close:
                trs   = text_stripped[tbl_open.end():tbl_close.start()]
                after = text_stripped[tbl_close.end():]
                after = re.sub(r'^\s*</div>\s*', '', after, flags=re.IGNORECASE)
                after = re.sub(r'\s*</div>\s*$', '', after, flags=re.IGNORECASE)
                return tbl_open.group(0), trs, after
            return None, None, text

        def _clean_div_wrapper(text):
            text = re.sub(r'^\s*<div[^>]*>\s*', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\s*</div>\s*$', '', text, flags=re.IGNORECASE)
            return text

        def _fix_border_top(trs_content):
            tr_end = re.search(r'</tr>', trs_content, re.IGNORECASE)
            if not tr_end:
                return trs_content
            first_row = trs_content[:tr_end.end()]
            rest      = trs_content[tr_end.end():]

            def fix_cell(m):
                tag = m.group(0)
                style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
                if style_m:
                    s = re.sub(r'border-top-style\s*:[^;]+;?', '', style_m.group(1), flags=re.IGNORECASE)
                    s = re.sub(r'border-top-width\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = re.sub(r'border-top-color\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = s.rstrip('; ') + '; border-top-style:none;'
                    tag = tag[:style_m.start()] + f'style="{s}"' + tag[style_m.end():]
                return tag

            return re.sub(r'<t[dh][^>]*>', fix_cell, first_row, flags=re.IGNORECASE) + rest

        with open(output_path, 'w', encoding='utf-8') as out_f:
            for file_idx, chunk_path in enumerate(chunk_html_paths):
                logger.debug(f"🔗 合并 chunk {file_idx}: {os.path.basename(chunk_path)}")
                with open(chunk_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                if file_idx == 0:
                    head_m = re.search(r'^(.*?<body[^>]*>)', content, re.DOTALL | re.IGNORECASE)
                    if head_m:
                        head_part = head_m.group(1)
                        if extra_css_blocks:
                            merged = '<style type="text/css">\n' + '\n'.join(extra_css_blocks) + '\n</style>\n'
                            head_part = head_part.replace('</head>', merged + '</head>')
                        out_f.write(head_part + '\n')

                body_m = re.search(r'<body[^>]*>(.*?)</body>', content, re.DOTALL | re.IGNORECASE)
                if not body_m:
                    logger.debug(f"   ⚠️ chunk {file_idx} 未找到 body，跳过")
                    continue

                body = body_m.group(1)
                body = self._clean_header_footer(body)

                cursor = 0
                while cursor <= len(body):
                    marker_m = marker_re.search(body, cursor)

                    if not marker_m:
                        remaining = _clean_div_wrapper(body[cursor:])
                        if remaining.strip():
                            if pending_table_group:
                                tbl_open, trs, after = _extract_trs(remaining)
                                if trs is not None:
                                    pending_trs.append(_fix_border_top(trs))
                                    if after.strip():
                                        _flush_pending_table(out_f)
                                        out_f.write(after)
                            else:
                                out_f.write(remaining)
                        break

                    before = _clean_div_wrapper(body[cursor:marker_m.start()])
                    group  = marker_m.group(1)

                    if before.strip():
                        if pending_table_group and pending_table_group != group:
                            _flush_pending_table(out_f)
                        out_f.write(before)

                    search_from = marker_m.end()
                    tbl_open_m  = re.search(r'<table[^>]*>',  body[search_from:], re.IGNORECASE)
                    tbl_close_m = re.search(r'</table>',       body[search_from:], re.IGNORECASE)

                    if tbl_open_m and tbl_close_m:
                        abs_open_end  = search_from + tbl_open_m.end()
                        abs_close_end = search_from + tbl_close_m.end()
                        trs_content   = body[abs_open_end : search_from + tbl_close_m.start()]
                        after_content = body[abs_close_end:]
                        after_content = re.sub(r'^\s*</div>\s*', '', after_content, flags=re.IGNORECASE)
                        after_content = re.sub(r'\s*</div>\s*$', '', after_content, flags=re.IGNORECASE)

                        if pending_table_group == group:
                            pending_trs.append(_fix_border_top(trs_content))
                        else:
                            _flush_pending_table(out_f)
                            pending_table_group = group
                            pending_table_open  = tbl_open_m.group(0)
                            pending_trs         = [trs_content]

                        cursor = abs_close_end
                    else:
                        remaining = _clean_div_wrapper(body[search_from:])
                        if remaining.strip():
                            out_f.write(remaining)
                        break

                out_f.write('\n')
                logger.debug(f"   ✅ chunk {file_idx} 合并完成")
            _flush_pending_table(out_f)
            out_f.write("</body>\n</html>\n")

        logger.debug(f"✅ 流式合并完成：{output_path}")
    # ------------------------------------------------------------------ #
    #  分片转换主流程                                                        #
    # ------------------------------------------------------------------ #

    def _chunked_docx_to_html(self, docx_path, html_path, temp_dir_prefix):
        """
        【内部方法】分片转换主流程
        步骤：拆分文档 → chunk各自转HTML（不内嵌图片）→ 流式合并 → 统一内嵌图片
        """
        html_dir = os.path.dirname(html_path)
        chunk_dir = self._normalize_path(
            os.path.join(html_dir, f"{temp_dir_prefix}_chunks")
        )
        original_img_dir = self._normalize_path(
            os.path.join(chunk_dir, "original_images")
        )

        image_display_order = self._get_image_order_from_docx(docx_path)
        self._extract_original_images(docx_path, original_img_dir)

        img_size_map = self._extract_image_display_sizes(docx_path)

        try:
            chunk_paths = self._split_docx_to_chunks(
                docx_path, chunk_dir, image_display_order
            )

            chunk_html_paths = []
            for idx, chunk_path in enumerate(chunk_paths):
                chunk_html_path = self._normalize_path(
                    os.path.join(chunk_dir, f"chunk_{idx:04d}.html")
                )
                ok = self._docx_to_html_no_embed(chunk_path, chunk_html_path)
                if ok:
                    chunk_html_paths.append(chunk_html_path)
                else:
                    logger.warning(f"⚠️ chunk_{idx:04d} 转换失败，跳过")
            logger.debug(f"📋 收集到 {len(chunk_html_paths)} 个chunk HTML文件，开始流式合并...")
            self._merge_html_files_to_disk(chunk_html_paths, html_path)

            # 【修复核心】使用二进制精确匹配内嵌图片
            logger.debug("🖼️ 开始统一内嵌图片（二进制精确匹配）...")
            self._embed_images_to_html(html_path, image_display_order, original_img_dir)

            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            if img_size_map:
                logger.debug("📐 修正图片显示尺寸...")
                html_content = self._fix_html_img_sizes(
                    html_content, img_size_map,
                    image_display_order,
                    image_display_order
                )

            logger.debug("📏 将表格固定宽度改为 100% 自适应...")
            html_content = self._make_tables_responsive(html_content)
            html_content = self._fix_underline_span_width(html_content)

            # 还原内部路径 href（去除占位域名 / 清理 file:// 前缀）
            html_content = self._decode_internal_hrefs(html_content)

            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            logger.debug(f"🎉 分片转换完成：{html_path}")
        except Exception as e:
            logger.error(f"❌ 分片转换异常：{e}")
            import traceback
            traceback.print_exc()

        finally:
            if os.path.exists(chunk_dir):
                shutil.rmtree(chunk_dir, ignore_errors=True)
                logger.debug(f"🗑️ 清理chunk目录：{chunk_dir}")
        if os.path.exists(html_path):
            with open(html_path, 'r', encoding='utf-8') as f:
                return f.read()
        return ""

    # ------------------------------------------------------------------ #
    #  公开方法                                                             #
    # ------------------------------------------------------------------ #

    def docx_to_single_html(self, docx_path, html_path):
        """
        公开方法：DOCX转单文件HTML
        特性：图片Base64内嵌 | CSS内嵌 | 图片无压缩 | 顺序对齐
              超限自动分片 | 大表格按行拆分后还原 | 去除页眉页脚

        【修复图片顺序错位问题】
        非分片路径同步改用二进制精确匹配策略替换图片，
        并补充调用 _clean_header_footer 去除页眉页脚（与分片路径行为一致）。

        :param docx_path: 输入DOCX文件路径（支持相对/绝对）
        :param html_path: 输出HTML文件路径（支持相对/绝对）
        :return: 生成的HTML文本内容，失败返回空字符串
        """
        # 1. 路径校验与标准化
        docx_path = self._normalize_path(docx_path)
        html_path = self._normalize_path(html_path)

        if not os.path.exists(docx_path):
            logger.error(f"❌ 输入DOCX文件不存在（绝对路径）：{docx_path}")
            return ""

        html_dir = os.path.dirname(html_path)
        os.makedirs(html_dir, exist_ok=True)

        temp_dir_prefix = self._make_temp_dir_prefix()

        # 2. 提取分页符/分节符，注入占位符（生成临时DOCX供Spire使用）
        breaks_map = {}
        docx_path_for_spire = docx_path
        try:
            docx_path_for_spire, breaks_map = self._inject_break_placeholders(docx_path, html_dir)
        except Exception as _be:
            logger.warning(f"⚠️ 分隔符预处理失败，跳过：{_be}")

        # 3. 检测文档规模，超限走分片流程
        if self._needs_chunking(docx_path_for_spire):
            html_result = self._chunked_docx_to_html(docx_path_for_spire, html_path, temp_dir_prefix)
            if breaks_map:
                html_result = self._restore_break_markers_in_html(html_result, breaks_map)
                with open(html_path, 'w', encoding='utf-8') as _f:
                    _f.write(html_result)
            if docx_path_for_spire != docx_path and os.path.exists(docx_path_for_spire):
                try:
                    os.remove(docx_path_for_spire)
                except Exception:
                    pass
            return html_result

        # 4. 创建唯一临时目录
        spire_temp_dir   = self._normalize_path(os.path.join(html_dir, temp_dir_prefix))
        original_img_dir = self._normalize_path(os.path.join(spire_temp_dir, "original_images"))
        spire_img_dir    = self._normalize_path(os.path.join(spire_temp_dir, "images"))

        # 5. 解析图片顺序 + 提取原始图片 + 提取显示尺寸
        image_display_order = self._get_image_order_from_docx(docx_path_for_spire)
        extracted_imgs = self._extract_original_images(docx_path_for_spire, original_img_dir)
        img_size_map   = self._extract_image_display_sizes(docx_path_for_spire)

        if not image_display_order and extracted_imgs:
            image_display_order = sorted(extracted_imgs)
            logger.warning(f"⚠️ 顺序解析为空，兜底使用：{image_display_order}")
        # 6. Spire转换生成临时HTML（图片不内嵌，保留文件引用）
        document = Document()
        try:
            document.LoadFromFile(docx_path_for_spire)
            document.HtmlExportOptions.ImageEmbedded = False
            document.HtmlExportOptions.ImagesPath = spire_img_dir
            document.HtmlExportOptions.ImageFormat = self.default_image_format
            # 标题清理格式
            self._clean_docx_headings_before_convert(document)
            document.SaveToFile(html_path, FileFormat.Html)
        except Exception as e:
            logger.error(f"❌ Spire转换HTML失败：{e}")
            return ""
        finally:
            document.Close()
            del document

        # 6. 读取HTML，统一路径分隔符
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        html_content = html_content.replace('\\', '/')

        # 7. 内嵌CSS
        css_file_path = self._normalize_path(os.path.splitext(html_path)[0] + '_styles.css')
        if os.path.exists(css_file_path):
            with open(css_file_path, 'r', encoding='utf-8') as f:
                css_content = f.read()
            html_content = re.sub(r'<link[^>]*href="[^"]+\.css"[^>]*>', '', html_content)
            html_content = html_content.replace(
                '</head>',
                f'<style type="text/css">\n{css_content}\n</style>\n</head>'
            )
            logger.debug("✅ 已内嵌CSS样式")
        # 8. 【修复】去除页眉页脚（非分片路径补充，与分片路径行为一致）
        logger.debug("🧹 去除页眉页脚...")
        html_content = self._clean_header_footer(html_content)

        # 9. 将处理后的HTML写回文件，供 _embed_images_to_html 读取
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        # 10. 【修复核心】二进制精确匹配内嵌图片
        logger.debug("🖼️ 内嵌图片（二进制精确匹配）...")
        self._embed_images_to_html(html_path, image_display_order, original_img_dir)

        # 11. 重新读取（_embed_images_to_html 已写回文件）
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # 12. 修正图片显示尺寸（DPI 2 倍问题）
        if img_size_map:
            logger.debug("📐 修正图片显示尺寸...")
            html_content = self._fix_html_img_sizes(
                html_content, img_size_map, image_display_order, image_display_order
            )

        # ====== 新增：修正表格固定宽度，使其网页自适应 ======
        logger.debug("📏 将表格固定宽度改为 100% 自适应...")
        html_content = self._make_tables_responsive(html_content)
        # ====================================================
        html_content = self._fix_underline_span_width(html_content)

        # 还原内部路径 href（去除占位域名 / 清理 file:// 前缀）
        html_content = self._decode_internal_hrefs(html_content)

        # 13. 恢复分页符/分节符自定义标记
        logger.warning(f"📄 准备恢复分隔符，breaks_map={len(breaks_map)} 个")
        if breaks_map:
            html_content = self._restore_break_markers_in_html(html_content, breaks_map)

        # 14. 保存最终HTML
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        # 15. 清理临时文件
        cleanup_paths = [spire_temp_dir, css_file_path]
        if docx_path_for_spire != docx_path:
            cleanup_paths.append(docx_path_for_spire)
        for temp_path in cleanup_paths:
            if os.path.exists(temp_path):
                try:
                    if os.path.isdir(temp_path):
                        shutil.rmtree(temp_path, ignore_errors=True)
                    else:
                        os.remove(temp_path)
                    logger.debug(f"🗑️ 清理临时文件：{temp_path}")
                except Exception as e:
                    logger.warning(f"⚠️ 清理临时文件失败 {temp_path}：{e}")
        logger.debug(f"\n🎉 DOCX转HTML完成！")
        logger.debug(f"📄 最终文件绝对路径：{html_path}")
        logger.debug(f"✅ 特性：图片Base64内嵌 | CSS内嵌 | 图片无压缩 | 顺序对齐 | 页眉页脚已清理")
        return html_content

    # ------------------------------------------------------------------ #
    #  分页符 / 分节符 处理工具                                              #
    # ------------------------------------------------------------------ #

    # ---------- DOCX → HTML 方向 ----------

    def _inject_break_placeholders(self, docx_path: str, work_dir: str):
        """
        读取 DOCX，将分页符和分节符替换为唯一文本占位符，
        保存为临时文件供 Spire 转换，并返回占位符元数据映射。

        分页符处理两种情形：
          ① 段落内容仅为分页符（正文为空或只有 \\xa0）→ 整段替换为占位符
          ② 分页符与正文在同一段落 → 在该段落前插入新的占位符段落，并从原段落移除 w:br

        Returns:
            (docx_path_for_spire, breaks_map)
            breaks_map = {
                'PB_MARKER_XXXXXXXX': {'type': 'page'},
                'SB_MARKER_XXXXXXXX': {'type': 'section', 'meta': {data-*: value, ...}},
            }
        """

        def _is_real_text(t: str) -> bool:
            """判断是否为实质性文字（\xa0 / 普通空白 不算）"""
            return bool(t.replace('\xa0', '').strip())

        try:
            doc = PythonDocx(docx_path)
        except Exception as e:
            logger.warning(f"⚠️ 无法读取DOCX提取分隔符：{e}")
            return docx_path, {}

        breaks_map = {}
        modified = False

        # 用列表收集（避免在迭代中修改 body 导致跳项）
        paragraphs = list(doc.paragraphs)

        for para in paragraphs:
            p_elem = para._p

            # ── 优先检测段落级分节符（pPr/sectPr）────────────────────────
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    marker_id = uuid.uuid4().hex[:8].upper()
                    marker = f'SB_MARKER_{marker_id}'

                    meta = {}
                    type_el = sectPr.find(qn('w:type'))
                    if type_el is not None:
                        meta['data-section-type'] = type_el.get(qn('w:val'), 'nextPage')

                    pgSz = sectPr.find(qn('w:pgSz'))
                    if pgSz is not None:
                        for attr, key in [(qn('w:w'), 'data-page-width'),
                                          (qn('w:h'), 'data-page-height'),
                                          (qn('w:orient'), 'data-orientation')]:
                            val = pgSz.get(attr)
                            if val:
                                meta[key] = val
                        # 部分 Word 不写 w:orient，仅靠宽>高表示横版，需补充推断
                        if 'data-orientation' not in meta:
                            try:
                                w = int(meta.get('data-page-width', 0))
                                h = int(meta.get('data-page-height', 0))
                                if w and h:
                                    meta['data-orientation'] = 'landscape' if w > h else 'portrait'
                            except ValueError:
                                pass

                    pgMar = sectPr.find(qn('w:pgMar'))
                    if pgMar is not None:
                        for side in ['top', 'bottom', 'left', 'right']:
                            val = pgMar.get(qn(f'w:{side}'))
                            if val:
                                meta[f'data-margin-{side}'] = val

                    breaks_map[marker] = {'type': 'section', 'meta': meta}

                    for r in p_elem.findall(qn('w:r')):
                        p_elem.remove(r)
                    r_el = OxmlElement('w:r')
                    t_el = OxmlElement('w:t')
                    t_el.text = marker
                    r_el.append(t_el)
                    p_elem.append(r_el)
                    pPr.remove(sectPr)
                    modified = True
                    continue

            # ── 检测分页符 ────────────────────────────────────────────────
            # 收集含 w:br type=page 的 run，以及有实质文字的 run
            br_runs = []    # [(run_elem, [br_child, ...])]
            has_real_text = False

            for r in p_elem.findall(qn('w:r')):
                brs_in_run = [
                    child for child in r
                    if child.tag == qn('w:br') and child.get(qn('w:type')) == 'page'
                ]
                if brs_in_run:
                    br_runs.append((r, brs_in_run))
                for child in r:
                    if child.tag == qn('w:t') and _is_real_text(child.text or ''):
                        has_real_text = True

            if not br_runs:
                continue

            marker_id = uuid.uuid4().hex[:8].upper()
            marker = f'PB_MARKER_{marker_id}'
            breaks_map[marker] = {'type': 'page'}

            if not has_real_text:
                # ① 纯分页符段落：整段替换为占位符
                logger.warning(f"[分页符①] 纯分页符段落，注入标记：{marker}")
                for r in p_elem.findall(qn('w:r')):
                    p_elem.remove(r)
                r_el = OxmlElement('w:r')
                t_el = OxmlElement('w:t')
                t_el.text = marker
                r_el.append(t_el)
                p_elem.append(r_el)
            else:
                # ② 分页符与正文共存：在本段前插入占位符段落，然后从 run 中移除 w:br
                logger.warning(f"[分页符②] 混合段落，插入标记：{marker}，正文={repr(para.text[:30])}")
                marker_p = OxmlElement('w:p')
                marker_r = OxmlElement('w:r')
                marker_t = OxmlElement('w:t')
                marker_t.text = marker
                marker_r.append(marker_t)
                marker_p.append(marker_r)
                p_elem.addprevious(marker_p)

                for r, brs in br_runs:
                    for br in brs:
                        r.remove(br)
                    # 若 run 除 rPr 外已无实质内容，则整个 run 删除
                    remaining = [c for c in r if c.tag != qn('w:rPr')]
                    all_empty = all(
                        c.tag == qn('w:t') and not _is_real_text(c.text or '')
                        for c in remaining
                    )
                    if not remaining or all_empty:
                        p_elem.remove(r)

            modified = True

        if not modified:
            return docx_path, {}

        # 为每个分节符填入"下一节"的方向，供标签"此处以下"显示。
        # OOXML 规则：pPr/sectPr 描述当前（结束于此）节；分节符之后那节的属性
        # 来自下一个 pPr/sectPr，最后一节来自 body 级 sectPr。
        body_orient = ''
        body_sectPr_el = doc.element.body.find(qn('w:sectPr'))
        if body_sectPr_el is not None:
            _pg = body_sectPr_el.find(qn('w:pgSz'))
            if _pg is not None:
                body_orient = _pg.get(qn('w:orient'), '')
                if not body_orient:
                    try:
                        _bw = int(_pg.get(qn('w:w')) or 0)
                        _bh = int(_pg.get(qn('w:h')) or 0)
                        if _bw and _bh:
                            body_orient = 'landscape' if _bw > _bh else 'portrait'
                    except (ValueError, TypeError):
                        pass
        sb_markers = [m for m, v in breaks_map.items() if v['type'] == 'section']
        for _i, _marker in enumerate(sb_markers):
            _next_orient = (
                breaks_map[sb_markers[_i + 1]]['meta'].get('data-orientation', '')
                if _i + 1 < len(sb_markers)
                else body_orient
            )
            if _next_orient:
                breaks_map[_marker]['meta']['data-next-orientation'] = _next_orient

        temp_docx_name = f"_breaks_{uuid.uuid4().hex[:8]}.docx"
        temp_docx_path = self._normalize_path(os.path.join(work_dir, temp_docx_name))
        doc.save(temp_docx_path)
        logger.warning(f"📄 分隔符占位符已注入，临时文件：{temp_docx_path}，"
                       f"分页符：{sum(1 for v in breaks_map.values() if v['type']=='page')}，"
                       f"分节符：{sum(1 for v in breaks_map.values() if v['type']=='section')}")
        return temp_docx_path, breaks_map


    def _restore_break_markers_in_html(self, html_content: str, breaks_map: dict) -> str:
        """
        将 Spire 生成的 HTML 中的占位符段落替换为自定义分页符/分节符 HTML。
        占位符文本可能被 Spire 包裹在 span 等子标签中。
        """
        if not breaks_map:
            return html_content

        PAGE_BREAK_HTML = (
            '<p contenteditable="false" class="page-break"'
            ' style="page-break-after: always; height: 3px; border-top: 3px dashed #169179;'
            ' text-align: center; margin:10px 0; clear: both; cursor: default; user-select: none;">'
            '<span style="position: relative; top: -12px; background: #fff;'
            ' padding: 0 10px; color: #169179; font-size: 14px;">此处为分页符</span>'
            '</p><p>&nbsp;</p>'
        )

        # data-* 属性的输出顺序（与前端保持一致）
        DATA_ATTR_ORDER = [
            'data-section-type', 'data-page-width', 'data-page-height',
            'data-orientation', 'data-margin-top', 'data-margin-bottom',
            'data-margin-left', 'data-margin-right',
        ]

        def make_section_break_html(meta: dict) -> str:
            # 按约定顺序拼接 data-* 属性，排除内部字段 data-next-orientation
            display_meta = {k: v for k, v in meta.items() if k != 'data-next-orientation'}
            ordered = {k: display_meta[k] for k in DATA_ATTR_ORDER if k in display_meta}
            ordered.update({k: v for k, v in display_meta.items() if k not in ordered})
            data_attrs = ' '.join(f'{k}="{v}"' for k, v in ordered.items())

            # 标签优先用"下一节"方向（data-next-orientation），表示分节符以下那节的版式
            label_orient = meta.get('data-next-orientation') or meta.get('data-orientation', '')
            if label_orient == 'landscape':
                label = '此处以下为分节符(横版)'
            elif label_orient == 'portrait':
                label = '此处以下为分节符(竖版)'
            else:
                label = '此处为分节符'

            return (
                f'<p contenteditable="false" {data_attrs} class="section-break"'
                f' style="border-top: 3px dashed #ff4040; text-align: center; height: 3px;'
                f' margin:10px 0; clear: both; cursor: default; user-select: none;">'
                f'<span style="position: relative; top: -12px; background: #fff;'
                f' padding: 0 10px; color: #ff4040; font-size: 14px;">{label}</span>'
                f'</p>'
            )

        # 支持 Spire 将标记段落渲染为 <p> 或 <h1>~<h6> 的情况
        BLOCK_OPEN_RE = re.compile(r'<(p|h[1-6])\b', re.IGNORECASE)

        for marker, info in breaks_map.items():
            replacement = PAGE_BREAK_HTML if info['type'] == 'page' \
                else make_section_break_html(info.get('meta', {}))

            idx = html_content.find(marker)
            if idx == -1:
                logger.warning(f"⚠️ 未在HTML中找到占位符：{marker}（类型={info['type']}）")
                continue

            # 向前找最近的块级开始标签（<p 或 <h1>~<h6>）
            p_start = -1
            matched_tag = 'p'
            for m in BLOCK_OPEN_RE.finditer(html_content[:idx]):
                p_start = m.start()
                matched_tag = m.group(1).lower()

            if p_start == -1:
                logger.warning(f"⚠️ 找不到占位符 {marker} 的开始标签")
                continue

            # 向后找对应的关闭标签
            close_tag = f'</{matched_tag}>'
            p_end_tag = html_content.find(close_tag, idx)
            if p_end_tag == -1:
                logger.warning(f"⚠️ 找不到占位符 {marker} 的关闭标签 {close_tag}")
                continue

            p_end = p_end_tag + len(close_tag)
            html_content = html_content[:p_start] + replacement + html_content[p_end:]
            logger.warning(f"✅ 已恢复分隔符：{marker} → {info['type']}（原标签 <{matched_tag}>）")

        return html_content

    # ---------- HTML → DOCX 方向 ----------

    def _extract_break_markers_from_html(self, html_text: str):
        """
        从 HTML 中提取自定义分页符/分节符 <p> 标签，
        替换为唯一文本占位符（Spire 可识别的普通段落）。

        Returns:
            (processed_html, breaks_info)
            breaks_info: [{'marker': str, 'type': 'page'|'section', 'meta': {...}}, ...]
        """
        breaks_info = []

        # 匹配 class 含 page-break 或 section-break 的 <p>，以及紧随的 &nbsp; 段落
        BREAK_RE = re.compile(
            r'<p\b([^>]*\bclass="[^"]*\b(?:page-break|section-break)\b[^"]*"[^>]*)>'
            r'.*?</p\s*>'
            r'(?:\s*<p[^>]*>\s*&nbsp;\s*</p\s*>)?',
            re.IGNORECASE | re.DOTALL
        )

        def replacer(m):
            attrs_str = m.group(1)
            class_m = re.search(r'\bclass="([^"]*)"', attrs_str, re.IGNORECASE)
            class_val = class_m.group(1) if class_m else ''
            is_section = 'section-break' in class_val

            marker_id = uuid.uuid4().hex[:8].upper()

            if is_section:
                meta = {}
                for data_m in re.finditer(r'\b(data-[\w-]+)="([^"]*)"', attrs_str):
                    meta[data_m.group(1)] = data_m.group(2)
                marker = f'SB_MARKER_{marker_id}'
                breaks_info.append({'marker': marker, 'type': 'section', 'meta': meta})
            else:
                marker = f'PB_MARKER_{marker_id}'
                breaks_info.append({'marker': marker, 'type': 'page'})

            return f'<p>{marker}</p>'

        processed_html = BREAK_RE.sub(replacer, html_text)
        logger.debug(f"📄 从HTML提取分隔符：{len(breaks_info)} 个")
        return processed_html, breaks_info

    def _apply_break_markers_to_docx(self, docx_path: str, breaks_info: list):
        """
        在 Spire 生成的 DOCX 中找到占位符段落，替换为真实的分页符/分节符 XML。
        """
        if not breaks_info:
            return

        marker_map = {item['marker']: item for item in breaks_info}
        doc = PythonDocx(docx_path)
        modified = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if text not in marker_map:
                continue

            info = marker_map[text]
            p_elem = para._p

            # 清空段落所有 run
            for r in p_elem.findall(qn('w:r')):
                p_elem.remove(r)

            if info['type'] == 'page':
                r_el = OxmlElement('w:r')
                br_el = OxmlElement('w:br')
                br_el.set(qn('w:type'), 'page')
                r_el.append(br_el)
                p_elem.append(r_el)
                logger.debug(f"✅ 写入分页符：{text}")

            elif info['type'] == 'section':
                meta = info.get('meta', {})

                pPr = p_elem.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    p_elem.insert(0, pPr)

                # 移除旧的 sectPr（若有）
                old = pPr.find(qn('w:sectPr'))
                if old is not None:
                    pPr.remove(old)

                sectPr = OxmlElement('w:sectPr')

                section_type = meta.get('data-section-type', 'nextPage')
                type_el = OxmlElement('w:type')
                type_el.set(qn('w:val'), section_type)
                sectPr.append(type_el)

                if 'data-page-width' in meta or 'data-page-height' in meta:
                    pgSz = OxmlElement('w:pgSz')
                    pgSz.set(qn('w:w'), meta.get('data-page-width', '12240'))
                    pgSz.set(qn('w:h'), meta.get('data-page-height', '15840'))
                    if meta.get('data-orientation') == 'landscape':
                        pgSz.set(qn('w:orient'), 'landscape')
                    sectPr.append(pgSz)

                margin_keys = [f'data-margin-{s}' for s in ['top', 'bottom', 'left', 'right']]
                if any(k in meta for k in margin_keys):
                    pgMar = OxmlElement('w:pgMar')
                    for side in ['top', 'bottom', 'left', 'right']:
                        val = meta.get(f'data-margin-{side}')
                        if val:
                            pgMar.set(qn(f'w:{side}'), val)
                    sectPr.append(pgMar)

                pPr.append(sectPr)
                logger.debug(f"✅ 写入分节符：{text}，类型={section_type}")

            modified = True

        if modified:
            doc.save(docx_path)
            logger.debug(f"💾 分隔符已写入DOCX：{docx_path}")

    # ------------------------------------------------------------------ #
    #  HTML → DOCX 分片工具                                                #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _html_count_paragraphs(html_text: str) -> int:
        """
        快速估算 HTML 中的段落数，用于判断是否需要分片。
        """
        return len(re.findall(
            r'<(?:p|li|h[1-6]|td|th|caption|dt|dd)[\s>]',
            html_text,
            re.IGNORECASE
        ))

    def _split_html_to_chunks(self, html_text: str) -> list[str]:
        """
        将大 HTML 按段落数切分为若干 chunk，每个 chunk 均为完整的 HTML 文档。
        """
        head_match = re.search(r'(<html[^>]*>.*?<body[^>]*>)', html_text,
                               re.DOTALL | re.IGNORECASE)
        if head_match:
            preamble = head_match.group(1)
            body_start = head_match.end()
        else:
            preamble = '<html><body>'
            body_start = 0

        body_end_match = re.search(r'</body\s*>', html_text, re.IGNORECASE)
        body_end = body_end_match.start() if body_end_match else len(html_text)
        body_content = html_text[body_start:body_end]

        TOP_LEVEL_TAGS = re.compile(
            r'<(table|p|ul|ol|dl|h[1-6]|div|blockquote|pre|figure|section|article|header|footer|aside|nav|main)[\s>]',
            re.IGNORECASE
        )
        OPEN_TAG  = re.compile(r'<([a-zA-Z][a-zA-Z0-9]*)[\s/>]')
        CLOSE_TAG = re.compile(r'</([a-zA-Z][a-zA-Z0-9]*)\s*>')
        VOID_TAGS = {'br','hr','img','input','meta','link','area','base',
                     'col','embed','param','source','track','wbr'}

        def _find_top_level_blocks(text):
            blocks  = []
            cursor  = 0
            n       = len(text)

            while cursor < n:
                m = TOP_LEVEL_TAGS.search(text, cursor)
                if not m:
                    if cursor < n:
                        blocks.append((cursor, n, None))
                    break

                if m.start() > cursor:
                    blocks.append((cursor, m.start(), None))

                tag_name = m.group(1).lower()
                open_pos = m.start()

                if tag_name in VOID_TAGS:
                    blocks.append((open_pos, m.end(), tag_name))
                    cursor = m.end()
                    continue

                depth  = 0
                pos    = open_pos
                end    = open_pos

                for tm in re.finditer(r'</?[a-zA-Z][a-zA-Z0-9]*[\s/>]?', text[open_pos:]):
                    raw = tm.group(0)
                    abs_start = open_pos + tm.start()
                    abs_end   = open_pos + tm.end()

                    inner_tag = re.match(r'</?([a-zA-Z][a-zA-Z0-9]*)', raw)
                    if not inner_tag:
                        continue
                    inner_name = inner_tag.group(1).lower()

                    if raw.startswith('</'):
                        if inner_name == tag_name:
                            depth -= 1
                            if depth == 0:
                                close_end = text.find('>', abs_end - 1)
                                end = (close_end + 1) if close_end != -1 else abs_end
                                break
                    elif inner_name not in VOID_TAGS and not raw.endswith('/>'):
                        if inner_name == tag_name or abs_start == open_pos:
                            depth += 1

                if end <= open_pos:
                    end = n

                blocks.append((open_pos, end, tag_name))
                cursor = end

            return blocks

        blocks = _find_top_level_blocks(body_content)

        LEAF_TAGS = {'table', 'p', 'ul', 'ol', 'dl',
                     'h1', 'h2', 'h3', 'h4', 'h5', 'h6', None}
        for _ in range(5):
            if len(blocks) != 1:
                break
            s, e, t = blocks[0]
            if t in LEAF_TAGS:
                break
            frag = body_content[s:e]
            if self._html_count_paragraphs(frag) <= self.MAX_PARAGRAPHS:
                break
            open_end   = frag.find('>') + 1
            close_start = frag.rfind('</')
            if not (0 < open_end < close_start):
                break
            body_content = frag[open_end:close_start]
            blocks = _find_top_level_blocks(body_content)

        chunks_html = []
        current_parts = []
        current_para_count = 0

        def _flush_chunk():
            nonlocal current_parts, current_para_count
            if current_parts:
                body_inner = ''.join(current_parts)
                chunks_html.append(f'{preamble}\n{body_inner}\n</body></html>')
                current_parts = []
                current_para_count = 0

        for start, end, tag in blocks:
            fragment = body_content[start:end]
            frag_paras = self._html_count_paragraphs(fragment)

            if tag == 'table' and frag_paras > self.MAX_PARAGRAPHS:
                _flush_chunk()
                sub_chunks = self._split_html_table_rows(fragment, preamble)
                chunks_html.extend(sub_chunks)
                continue

            if current_para_count + frag_paras > self.MAX_PARAGRAPHS and current_parts:
                _flush_chunk()

            current_parts.append(fragment)
            current_para_count += frag_paras

        _flush_chunk()

        logger.debug(f"📦 HTML 切分为 {len(chunks_html)} 个 chunk（总估算段落：{self._html_count_paragraphs(body_content)}）")
        return chunks_html

    @staticmethod
    def _find_top_level_trs(html: str) -> list[str]:
        """
        从表格 HTML 中提取顶层 <tr>...</tr>，正确跳过嵌套表格内的行。
        使用深度计数而非正则，避免嵌套 <table><tr> 被误匹配。
        """
        rows = []
        tag_re = re.compile(r'<(/?)(?:tr|table)\b[^>]*>', re.IGNORECASE)
        table_depth = 0  # 当前 <tr> 内部嵌套 <table> 的深度
        tr_start = -1   # 顶层 <tr> 的起始位置
        tr_depth = 0    # 顶层 <tr> 的嵌套层数（处理同级 tr 计数）

        for m in tag_re.finditer(html):
            is_close = m.group(1) == '/'
            tag_name_m = re.match(r'</?([a-zA-Z]+)', m.group(0))
            if not tag_name_m:
                continue
            tag_name = tag_name_m.group(1).lower()

            if tag_name == 'table':
                if not is_close:
                    if tr_start >= 0:   # 在顶层 <tr> 内部遇到 <table>
                        table_depth += 1
                else:
                    if tr_start >= 0 and table_depth > 0:
                        table_depth -= 1
            elif tag_name == 'tr':
                if not is_close:
                    if table_depth == 0:   # 顶层 <tr>（不在嵌套 table 内）
                        if tr_depth == 0:
                            tr_start = m.start()
                        tr_depth += 1
                else:
                    if table_depth == 0 and tr_depth > 0:
                        tr_depth -= 1
                        if tr_depth == 0:
                            rows.append(html[tr_start:m.end()])
                            tr_start = -1

        return rows

    def _split_html_table_rows(self, table_html: str, preamble: str) -> list[str]:
        """
        将单张超大 HTML 表格按 <tr> 行粒度切分为多个 chunk。
        """
        tbl_open_m = re.match(r'<table[^>]*>', table_html, re.IGNORECASE)
        tbl_open   = tbl_open_m.group(0) if tbl_open_m else '<table>'

        thead_m = re.search(r'<thead[\s>].*?</thead\s*>', table_html,
                            re.IGNORECASE | re.DOTALL)
        thead_html = thead_m.group(0) if thead_m else ''

        body_area = table_html
        if thead_m:
            body_area = table_html[:thead_m.start()] + table_html[thead_m.end():]

        all_trs = self._find_top_level_trs(body_area)

        chunks_html = []
        current_trs = []
        current_para_count = 0

        def _flush_table_chunk():
            nonlocal current_trs, current_para_count
            if current_trs:
                body_inner = (
                    f'{tbl_open}\n'
                    f'{thead_html}\n'
                    f'<tbody>\n{"".join(current_trs)}\n</tbody>\n'
                    f'</table>'
                )
                chunks_html.append(f'{preamble}\n{body_inner}\n</body></html>')
                current_trs = []
                current_para_count = 0

        for tr in all_trs:
            tr_paras = self._html_count_paragraphs(tr)
            if current_para_count + tr_paras > self.MAX_PARAGRAPHS and current_trs:
                _flush_table_chunk()
            current_trs.append(tr)
            current_para_count += tr_paras

        _flush_table_chunk()

        logger.debug(f"   📊 超大表格按行切分为 {len(chunks_html)} 个 chunk（{len(all_trs)} 行）")
        return chunks_html

    def _html_chunk_to_docx(self, html_chunk: str, output_path: str,
                             temp_img_dir: str | None) -> bool:
        """
        将单个 HTML chunk 转为 DOCX（通过 Spire）。
        """
        document       = None
        temp_html_path = None
        try:
            html_dir = (temp_img_dir if temp_img_dir and os.path.isdir(temp_img_dir)
                        else os.path.dirname(output_path))
            os.makedirs(html_dir, exist_ok=True)

            temp_html_path = self._normalize_path(
                os.path.join(html_dir, f"_chunk_{uuid.uuid4().hex[:8]}.html")
            )

            html_to_write = html_chunk
            if temp_img_dir and os.path.isdir(temp_img_dir):
                def _to_rel_src(m):
                    tag   = m.group(0)
                    src_m = re.search(r'src="([^"]+)"', tag)
                    if not src_m:
                        return tag
                    src = src_m.group(1)
                    if src.startswith('http') or src.startswith('data:'):
                        return tag
                    try:
                        rel = os.path.relpath(src, html_dir).replace('\\', '/')
                        return tag[:src_m.start()] + f'src="{rel}"' + tag[src_m.end():]
                    except ValueError:
                        return tag
                html_to_write = re.sub(
                    r'<img\b[^>]*>', _to_rel_src, html_chunk, flags=re.IGNORECASE
                )

            # 保护内部路径 href，防止 Spire 将其转换为 file:// 路径
            html_to_write = self._encode_internal_hrefs(html_to_write)

            with open(temp_html_path, 'w', encoding='utf-8') as f:
                f.write(html_to_write)

            document = Document()
            document.LoadFromFile(temp_html_path, FileFormat.Html, self.html_validation_type)
            document.SaveToFile(output_path, FileFormat.Docx2016)
            return True

        except Exception as e:
            logger.debug(f"   ❌ chunk 转 DOCX 失败（{os.path.basename(output_path)}）：{e}")
            return False
        finally:
            if document:
                document.Close()
                del document
            if temp_html_path and os.path.exists(temp_html_path):
                try:
                    os.remove(temp_html_path)
                except Exception:
                    pass

    def _merge_docx_chunks(self, chunk_docx_paths: list[str], output_docx_path: str) -> bool:
        """
        将多个 chunk DOCX 文件用 python-docx + lxml 合并为一个完整 DOCX。
        """
        if not chunk_docx_paths:
            return False
        if len(chunk_docx_paths) == 1:
            shutil.copy2(chunk_docx_paths[0], output_docx_path)
            return True

        def _read_zip(path):
            files = {}
            with zipfile.ZipFile(path, 'r') as zf:
                for item in zf.infolist():
                    files[item.filename] = zf.read(item.filename)
            return files

        base_files = _read_zip(chunk_docx_paths[0])

        used_media = {
            os.path.basename(k)
            for k in base_files
            if k.startswith('word/media/')
        }

        RELS_PATH = 'word/_rels/document.xml.rels'
        base_doc_xml  = base_files.get('word/document.xml', b'').decode('utf-8')
        base_rels_xml = base_files.get(RELS_PATH, b'').decode('utf-8')

        existing_rids = re.findall(r'Id="(rId\d+)"', base_rels_xml)
        max_rid = max((int(r[3:]) for r in existing_rids), default=0)

        base_doc_tree = etree.fromstring(base_doc_xml.encode('utf-8'))
        base_body = base_doc_tree.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')

        W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        sectPr_tag = f'{{{W_NS}}}sectPr'
        base_sectPr = base_body.find(sectPr_tag)
        if base_sectPr is not None:
            base_body.remove(base_sectPr)

        extra_rels  = []
        extra_media = {}

        for chunk_idx, chunk_path in enumerate(chunk_docx_paths[1:], start=1):
            chunk_files = _read_zip(chunk_path)
            chunk_doc_xml  = chunk_files.get('word/document.xml', b'').decode('utf-8')
            chunk_rels_xml = chunk_files.get(RELS_PATH, b'').decode('utf-8')

            chunk_rels_map = {}
            try:
                rels_root = etree.fromstring(chunk_rels_xml.encode('utf-8'))
                for rel in rels_root:
                    rid = rel.get('Id') or rel.get('id')
                    rel_type = rel.get('Type') or rel.get('type') or ''
                    target = rel.get('Target') or rel.get('target') or ''
                    if rid:
                        chunk_rels_map[rid] = (rel_type, target)
            except Exception:
                for m in re.finditer(r'<Relationship\b([^>]+)/>', chunk_rels_xml, re.IGNORECASE):
                    attrs = m.group(1)
                    rid_m   = re.search(r'\bId="(rId\d+)"', attrs, re.IGNORECASE)
                    type_m  = re.search(r'\bType="([^"]+)"', attrs, re.IGNORECASE)
                    tgt_m   = re.search(r'\bTarget="([^"]+)"', attrs, re.IGNORECASE)
                    if rid_m and type_m and tgt_m:
                        chunk_rels_map[rid_m.group(1)] = (type_m.group(1), tgt_m.group(1))

            rid_remap = {}

            for old_rid, (rel_type, target) in chunk_rels_map.items():
                rel_type_lower = rel_type.lower()
                if not any(k in rel_type_lower for k in ('image', 'hyperlink', 'oleobject')):
                    continue

                max_rid += 1
                new_rid = f'rId{max_rid}'
                rid_remap[old_rid] = new_rid

                if 'image' in rel_type_lower or 'oleobject' in rel_type_lower:
                    orig_fname  = os.path.basename(target)
                    new_fname   = orig_fname
                    fname_stem  = os.path.splitext(orig_fname)[0]
                    fname_ext   = os.path.splitext(orig_fname)[1]

                    if new_fname in used_media:
                        new_fname = f'{fname_stem}_c{chunk_idx}{fname_ext}'
                        if new_fname in used_media:
                            new_fname = f'{fname_stem}_{uuid.uuid4().hex[:6]}{fname_ext}'

                    used_media.add(new_fname)
                    new_target  = f'media/{new_fname}'
                    src_zip_path = f'word/media/{orig_fname}'

                    if src_zip_path in chunk_files:
                        extra_media[f'word/{new_target}'] = chunk_files[src_zip_path]

                    extra_rels.append(
                        f'<Relationship Id="{new_rid}" Type="{rel_type}" Target="{new_target}"/>'
                    )
                else:
                    extra_rels.append(
                        f'<Relationship Id="{new_rid}" Type="{rel_type}" Target="{target}" TargetMode="External"/>'
                    )

            chunk_doc_patched = chunk_doc_xml
            for old_rid in sorted(rid_remap, key=lambda r: int(r[3:]), reverse=True):
                new_rid = rid_remap[old_rid]
                chunk_doc_patched = chunk_doc_patched.replace(
                    f'"{old_rid}"', f'"{new_rid}"'
                )

            try:
                chunk_tree = etree.fromstring(chunk_doc_patched.encode('utf-8'))
                chunk_body = chunk_tree.find(f'{{{W_NS}}}body')
                if chunk_body is None:
                    logger.debug(f"   ⚠️ chunk {chunk_idx} 无 body，跳过")
                    continue
                for el in list(chunk_body):
                    if el.tag == sectPr_tag:
                        continue
                    base_body.append(copy.deepcopy(el))
            except Exception as e:
                logger.debug(f"   ⚠️ chunk {chunk_idx} XML 解析失败：{e}，跳过")
                continue

            logger.debug(f"   ✅ chunk {chunk_idx} 合并完成（rId 重映射 {len(rid_remap)} 条，媒体 {len(extra_media)} 个）")
        if base_sectPr is not None:
            base_body.append(base_sectPr)

        new_doc_xml = etree.tostring(base_doc_tree, xml_declaration=True,
                                     encoding='UTF-8', standalone=True)

        new_rels_xml = base_rels_xml.replace(
            '</Relationships>',
            '\n'.join(extra_rels) + '\n</Relationships>'
        )

        # 更新 [Content_Types].xml：补充后续 chunk 引入的新媒体文件扩展名
        CONTENT_TYPES_PATH = '[Content_Types].xml'
        KNOWN_EXT_TYPES = {
            'png':  'image/png',
            'jpg':  'image/jpeg',
            'jpeg': 'image/jpeg',
            'gif':  'image/gif',
            'bmp':  'image/bmp',
            'tif':  'image/tiff',
            'tiff': 'image/tiff',
            'webp': 'image/webp',
        }
        content_types_xml = base_files.get(CONTENT_TYPES_PATH, b'').decode('utf-8')
        new_content_types_entries = []
        for zip_path in extra_media:
            ext = os.path.splitext(zip_path)[1].lstrip('.').lower()
            mime = KNOWN_EXT_TYPES.get(ext)
            if mime and f'Extension="{ext}"' not in content_types_xml:
                new_content_types_entries.append(
                    f'<Default Extension="{ext}" ContentType="{mime}"/>'
                )
                content_types_xml += ''  # 标记需写入
        if new_content_types_entries:
            content_types_xml = content_types_xml.replace(
                '</Types>',
                '\n'.join(new_content_types_entries) + '\n</Types>'
            )

        tmp_path = output_docx_path + '.mergetmp'
        with zipfile.ZipFile(chunk_docx_paths[0], 'r') as src_zip, \
             zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as dst_zip:

            for item in src_zip.infolist():
                if item.filename == 'word/document.xml':
                    dst_zip.writestr(item, new_doc_xml)
                elif item.filename == RELS_PATH:
                    dst_zip.writestr(item, new_rels_xml.encode('utf-8'))
                elif item.filename == CONTENT_TYPES_PATH and new_content_types_entries:
                    dst_zip.writestr(item, content_types_xml.encode('utf-8'))
                else:
                    dst_zip.writestr(item, src_zip.read(item.filename))

            for zip_path, data in extra_media.items():
                dst_zip.writestr(zip_path, data)

        os.replace(tmp_path, output_docx_path)
        logger.debug(f"✅ DOCX 合并完成：{output_docx_path}")
        return True

    # ------------------------------------------------------------------ #
    #  公开方法（续）                                                        #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _strip_paragraph_spacing_css(html: str) -> str:
        """
        HTML→DOCX 方向的段落间距预处理。

        Spire 将 DOCX 导出为 HTML 时，会把段落的 before/after 间距编码为
        <p> 标签的 margin-top / margin-bottom 内联样式。当该 HTML 被 Spire
        再次导入为 DOCX 时，这些 CSS 值会被叠加到 Word 段落样式自带的间距上，
        造成段前/段后间距翻倍（表现为行与行之间出现额外空白）。

        解决方案：在导入前把 <p> 标签 style 属性里的
        margin-top / margin-bottom / padding-top / padding-bottom
        全部置零，交由 Word 样式控制间距，避免 CSS 与样式双重叠加。
        """
        BLOCK_RE = re.compile(r'<(p|h[1-6])\b([^>]*)>', re.IGNORECASE)

        def _zero_spacing(m):
            tag_name = m.group(1)
            attrs    = m.group(2)
            style_m  = re.search(r'style="([^"]*)"', attrs, re.IGNORECASE)
            if not style_m:
                return m.group(0)
            style = style_m.group(1)
            for prop in ('margin-top', 'margin-bottom', 'padding-top', 'padding-bottom'):
                style = re.sub(
                    rf'\b{prop}\s*:[^;]+;?', f'{prop}:0pt;',
                    style, flags=re.IGNORECASE
                )
            new_attrs = attrs[:style_m.start()] + f'style="{style}"' + attrs[style_m.end():]
            return f'<{tag_name}{new_attrs}>'

        return BLOCK_RE.sub(_zero_spacing, html)

    @staticmethod
    def _preprocess_tables_for_import(html: str) -> str:
        """HTML→DOCX 前：将所有 <table> 设为 width:100% + table-layout:fixed，
        使 Spire 按页面全宽渲染表格，同时保留各列宽度作为比例参考。"""
        def _fix_table_open(m):
            tag = m.group(0)
            tag = re.sub(r'\s+width="[^"]*"', '', tag, flags=re.IGNORECASE)
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            if style_m:
                style = style_m.group(1)
                style = re.sub(r'(?<![a-zA-Z-])width\s*:[^;]+;?', '', style, flags=re.IGNORECASE)
                style = re.sub(r'table-layout\s*:[^;]+;?', '', style, flags=re.IGNORECASE)
                style = style.strip('; ')
                style = (style + '; ' if style else '') + 'width:100%;table-layout:fixed;'
                tag = tag[:style_m.start()] + f'style="{style}"' + tag[style_m.end():]
            else:
                tag = re.sub(
                    r'(<table\b)', r'\1 style="width:100%;table-layout:fixed;"',
                    tag, flags=re.IGNORECASE, count=1,
                )
            return tag

        return re.sub(r'<table\b[^>]*>', _fix_table_open, html, flags=re.IGNORECASE | re.DOTALL)

    @staticmethod
    @staticmethod
    def _clean_mce_html(html: str) -> str:
        """
        清理 TinyMCE 产生的冗余标记，避免 Spire 转换时因空块级元素报错：
        1. 去除所有 data-mce-* 属性（包括 data-mce-bogus、data-mce-style 等）
        2. 将仅含 <br> 的空块级元素（h1-h6、p、div）替换为 <p>&nbsp;</p>
        """
        # 1. 去除 data-mce-* 属性
        html = re.sub(r'\s+data-mce-[a-zA-Z0-9_-]+(?:="[^"]*"|=\'[^\']*\'|(?=[>\s]))', '', html)

        # 2. 空块级元素（内容为空或只有 <br>）→ <p>&nbsp;</p>
        block_tags = r'(?:h[1-6]|p|div)'
        html = re.sub(
            r'<(' + block_tags + r')(\s[^>]*)?>(\s*<br\s*/?>)*\s*</\1>',
            '<p>&nbsp;</p>',
            html,
            flags=re.IGNORECASE
        )
        return html

    def html_text_to_docx(self, html_text: str, output_docx_path: str):
        """
        公开方法：HTML文本转DOCX

        支持超大 HTML（段落 > MAX_PARAGRAPHS）自动切片转换，绕过 Spire 免费版限制。

        :param html_text: 输入HTML字符串
        :param output_docx_path: 输出DOCX文件路径（支持相对/绝对）
        :return: 成功返回True，失败返回False
        """
        output_docx_path = self._normalize_path(output_docx_path)
        html_text = self._clean_mce_html(html_text)

        if not html_text.strip():
            logger.error("❌ HTML文本为空，无法转换")
            return False

        # 在处理图片 base64 之前，先修正表格宽度和图片比例意图
        logger.debug("🧹 正在优化 HTML 表格宽度与图片比例...")
        # html_text = self._sanitize_html_styles_for_import(html_text)
        output_dir   = os.path.dirname(output_docx_path)
        os.makedirs(output_dir, exist_ok=True)

        temp_img_dir  = None
        chunk_dir     = None

        try:
            html_text, temp_img_dir = self._extract_base64_images(html_text, output_dir)
            html_text = self._fix_centered_images_for_import(html_text)
            html_text = self._normalize_img_units_for_import(html_text)
            html_text = self._preprocess_tables_for_import(html_text)

            # 提取分页符/分节符，替换为唯一占位符
            breaks_info = []
            try:
                html_text, breaks_info = self._extract_break_markers_from_html(html_text)
            except Exception as _be:
                logger.warning(f"⚠️ 分隔符预处理失败，跳过：{_be}")

            para_count = self._html_count_paragraphs(html_text)
            logger.debug(f"📊 HTML 段落估算：{para_count}，阈值：{self.MAX_PARAGRAPHS}")
            if para_count <= self.MAX_PARAGRAPHS:
                logger.debug("✅ 无需分片，直接转换")
                ok = self._html_chunk_to_docx(html_text, output_docx_path, temp_img_dir)
                if ok and breaks_info:
                    try:
                        self._apply_break_markers_to_docx(output_docx_path, breaks_info)
                    except Exception as _be:
                        logger.warning(f"⚠️ 分隔符写入DOCX失败：{_be}")
                return ok

            logger.debug(f"⚡ 触发 HTML 分片转换（段落估算 {para_count} > {self.MAX_PARAGRAPHS}）")
            chunk_dir = self._normalize_path(
                os.path.join(output_dir, f"html2docx_{uuid.uuid4().hex[:8]}")
            )
            os.makedirs(chunk_dir, exist_ok=True)

            html_chunks = self._split_html_to_chunks(html_text)

            chunk_docx_paths = []
            for idx, chunk_html in enumerate(html_chunks):
                chunk_docx_path = self._normalize_path(
                    os.path.join(chunk_dir, f"chunk_{idx:04d}.docx")
                )
                ok = self._html_chunk_to_docx(chunk_html, chunk_docx_path, temp_img_dir)
                if ok:
                    chunk_docx_paths.append(chunk_docx_path)
                    logger.debug(f"   ✅ chunk_{idx:04d} 转换完成")
                else:
                    logger.debug(f"   ⚠️ chunk_{idx:04d} 转换失败，跳过")
            if not chunk_docx_paths:
                logger.error("❌ 所有 chunk 均转换失败")
                return False

            logger.debug(f"🔗 开始合并 {len(chunk_docx_paths)} 个 chunk DOCX...")
            ok = self._merge_docx_chunks(chunk_docx_paths, output_docx_path)
            if ok and breaks_info:
                try:
                    self._apply_break_markers_to_docx(output_docx_path, breaks_info)
                except Exception as _be:
                    logger.warning(f"⚠️ 分隔符写入DOCX失败（分片路径）：{_be}")
            return ok

        except Exception as e:
            logger.error(f"❌ HTML转DOCX失败：{str(e)}")
            import traceback
            traceback.print_exc()
            return False

        finally:
            if temp_img_dir and os.path.exists(temp_img_dir):
                try:
                    shutil.rmtree(temp_img_dir, ignore_errors=True)
                    logger.debug(f"🗑️ 清理图片临时目录：{temp_img_dir}")
                except Exception as e:
                    logger.warning(f"⚠️ 清理图片临时目录失败：{e}")
            if chunk_dir and os.path.exists(chunk_dir):
                try:
                    shutil.rmtree(chunk_dir, ignore_errors=True)
                    logger.debug(f"🗑️ 清理chunk目录：{chunk_dir}")
                except Exception as e:
                    logger.warning(f"⚠️ 清理chunk目录失败：{e}")
    # ------------------------------------------------------------------ #
    #  超链接保护（HTML ↔ DOCX 方向）                                        #
    # ------------------------------------------------------------------ #

    # Spire 在将 HTML 写入 DOCX 时，会把形如 /path 的绝对路径 href
    # 解析为本地文件路径并加上 file:// 前缀，导致回转 HTML 后链接失效。
    # 通过以下占位域名"欺骗" Spire，让它把内部路径当成外部 HTTP 链接处理，
    # 从而完整保留原始 URL。
    _HREF_PLACEHOLDER_DOMAIN = "https://xappinternal.local"

    def _encode_internal_hrefs(self, html_text: str) -> str:
        """
        HTML→DOCX 前调用：将非 http/https/mailto/tel/#/data: 开头的 href
        替换为带占位域名的形式，防止 Spire 将其转换为 file:// 路径。

        仅处理绝对路径（以 / 开头），相对路径暂不处理以避免语义变化。
        示例：href="/doc_editor/embeds/EMB_xxx/go"
              → href="https://xappinternal.local/doc_editor/embeds/EMB_xxx/go"
        """
        _safe_prefixes = ('http://', 'https://', 'mailto:', 'tel:', 'javascript:', '#', 'data:')

        def _replace(m):
            href = m.group(1)
            if any(href.lower().startswith(p) for p in _safe_prefixes):
                return m.group(0)
            if href.startswith('/'):
                return f'href="{self._HREF_PLACEHOLDER_DOMAIN}{href}"'
            return m.group(0)

        return re.sub(r'\bhref="([^"]*)"', _replace, html_text, flags=re.IGNORECASE)

    def _decode_internal_hrefs(self, html_text: str) -> str:
        """
        DOCX→HTML 后调用：还原 _encode_internal_hrefs 写入的占位域名，
        并兜底清理 Spire 遗留的 file:/// 前缀（用于处理旧存量 DOCX）。

        规则：
          1. https://xappinternal.local/path → /path
          2. href="file:///path"             → href="/path"  （兜底）
          3. href="file://path"              → href="/path"  （兜底）
        """
        placeholder = re.escape(self._HREF_PLACEHOLDER_DOMAIN)
        # 还原占位域名
        html_text = re.sub(
            rf'\bhref="{placeholder}(/[^"]*)"',
            r'href="\1"',
            html_text,
            flags=re.IGNORECASE,
        )
        # 兜底1：file:///绝对路径 → /path
        html_text = re.sub(
            r'\bhref="file:///([^"]*)"',
            r'href="/\1"',
            html_text,
            flags=re.IGNORECASE,
        )
        # 兜底2：file://相对路径 → /path（少数情况）
        html_text = re.sub(
            r'\bhref="file://([^/][^"]*)"',
            r'href="/\1"',
            html_text,
            flags=re.IGNORECASE,
        )
        return html_text

    # ------------------------------------------------------------------ #
    #  图片居中适配（HTML→DOCX 方向）                                       #
    # ------------------------------------------------------------------ #

    def _fix_centered_images_for_import(self, html_text: str) -> str:
        """
        将 <img> 标签上的 CSS 居中写法（display:block + margin-left/right:auto）
        转换为 DOCX 可识别的段落居中（text-align:center）。

        处理逻辑：
          1. 若居中 img 在 <p> 内 → 给该 <p> 加 text-align:center，并清除 img 上的居中 CSS
          2. 若居中 img 为独立标签  → 用 <p style="text-align:center;"> 包裹，并清除 img 上的居中 CSS
        """
        IMG_RE = re.compile(r'<img\b[^>]*>', re.IGNORECASE | re.DOTALL)

        def _has_center_style(style_str: str) -> bool:
            s = style_str.lower()
            has_block = bool(re.search(r'\bdisplay\s*:\s*block\b', s))
            has_ml    = bool(re.search(r'\bmargin-left\s*:\s*auto\b', s))
            has_mr    = bool(re.search(r'\bmargin-right\s*:\s*auto\b', s))
            # 支持 margin: X auto 简写形式
            has_short = bool(re.search(r'\bmargin\s*:\s*\S+\s+auto\b', s)) or \
                        bool(re.search(r'\bmargin\s*:\s*auto\b', s))
            return has_block and ((has_ml and has_mr) or has_short)

        def _img_is_centered(tag: str) -> bool:
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            return bool(style_m and _has_center_style(style_m.group(1)))

        def _remove_center_css(tag: str) -> str:
            """从 img 标签 style 中移除居中相关 CSS 属性。"""
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            if not style_m:
                return tag
            s = style_m.group(1)
            s = re.sub(r'\bdisplay\s*:\s*block\s*;?\s*', '', s, flags=re.IGNORECASE)
            s = re.sub(r'\bmargin-left\s*:\s*auto\s*;?\s*', '', s, flags=re.IGNORECASE)
            s = re.sub(r'\bmargin-right\s*:\s*auto\s*;?\s*', '', s, flags=re.IGNORECASE)
            s = s.strip().strip(';').strip()
            return tag[:style_m.start()] + f'style="{s}"' + tag[style_m.end():]

        def _add_center_to_p(p_open: str) -> str:
            """给 <p> 同时加 align="center" 属性和 text-align:center 样式，兼容 Spire 各版本。"""
            # align 属性
            if not re.search(r'\balign\s*=', p_open, re.IGNORECASE):
                p_open = re.sub(r'(<p\b)', r'\1 align="center"', p_open, flags=re.IGNORECASE)
            # text-align:center 样式
            style_m = re.search(r'style="([^"]*)"', p_open, re.IGNORECASE)
            if style_m:
                s = style_m.group(1)
                if not re.search(r'\btext-align\s*:', s, re.IGNORECASE):
                    s = s.rstrip('; ') + '; text-align:center'
                    p_open = p_open[:style_m.start()] + f'style="{s}"' + p_open[style_m.end():]
            else:
                p_open = re.sub(r'(<p\b)', r'\1 style="text-align:center"',
                                p_open, flags=re.IGNORECASE)
            return p_open

        # ── Step 1：对含有居中 img 的 <p>，追加居中属性 ──────────────────────
        def _fix_p(m):
            p_open, p_body, p_close = m.group(1), m.group(2), m.group(3)
            if not any(_img_is_centered(im.group(0)) for im in IMG_RE.finditer(p_body)):
                return m.group(0)
            p_open = _add_center_to_p(p_open)
            # 清理 img 上的居中 CSS（避免 Step 2 重复包裹）
            p_body = IMG_RE.sub(
                lambda im: _remove_center_css(im.group(0)) if _img_is_centered(im.group(0))
                           else im.group(0),
                p_body
            )
            return p_open + p_body + p_close

        html_text = re.sub(
            r'(<p\b[^>]*>)(.*?)(</p\s*>)',
            _fix_p,
            html_text,
            flags=re.IGNORECASE | re.DOTALL
        )

        # ── Step 2：剩余独立居中 img（不在 <p> 内）包裹为居中段落 ──────────
        def _wrap_img(m):
            tag = m.group(0)
            if not _img_is_centered(tag):
                return tag
            logger.debug(f"   🖼️ 居中 img 包裹为居中段落")
            return f'<p align="center" style="text-align:center;">{_remove_center_css(tag)}</p>'

        html_text = IMG_RE.sub(_wrap_img, html_text)
        return html_text

    # ------------------------------------------------------------------ #
    #  图片预处理工具（HTML→DOCX 方向）                                     #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _decode_b64_safe(b64_raw: str) -> bytes | None:
        """容忍换行/空白/残缺 padding 的 base64 解码，失败返回 None。"""
        try:
            clean   = re.sub(r'\s', '', b64_raw)
            padding = (4 - len(clean) % 4) % 4
            return base64.b64decode(clean + '=' * padding, validate=False)
        except Exception:
            return None

    @staticmethod
    def _read_image_wh(data: bytes):
        """
        从图片二进制头解析物理像素（w, h），支持 PNG/JPEG/GIF/BMP/WEBP。
        失败返回 (None, None)。
        """
        import struct
        try:
            if data[:8] == b'\x89PNG\r\n\x1a\n':
                return struct.unpack('>II', data[16:24])
            if data[:2] == b'\xff\xd8':
                i = 2
                while i < len(data) - 8:
                    if data[i] != 0xff: break
                    marker = data[i + 1]
                    if i + 3 >= len(data): break
                    length = struct.unpack('>H', data[i + 2:i + 4])[0]
                    if marker in (0xc0, 0xc1, 0xc2, 0xc3, 0xc5, 0xc6, 0xc7, 0xc9, 0xca, 0xcb):
                        h, w = struct.unpack('>HH', data[i + 5:i + 9])
                        return w, h
                    i += 2 + length
            if data[:6] in (b'GIF87a', b'GIF89a'):
                return struct.unpack('<HH', data[6:10])
            if data[:2] == b'BM':
                w, h = struct.unpack('<II', data[18:26])
                return w, abs(h)
            if data[:4] == b'RIFF' and len(data) >= 30 and data[8:12] == b'WEBP':
                chunk = data[12:16]
                if chunk == b'VP8 ':
                    w = struct.unpack('<H', data[26:28])[0] & 0x3fff
                    h = struct.unpack('<H', data[28:30])[0] & 0x3fff
                    return w, h
                if chunk == b'VP8L' and len(data) >= 25:
                    bits = struct.unpack('<I', data[21:25])[0]
                    return (bits & 0x3fff) + 1, ((bits >> 14) & 0x3fff) + 1
                if chunk == b'VP8X' and len(data) >= 30:
                    w = (int.from_bytes(data[24:27], 'little') & 0xffffff) + 1
                    h = (int.from_bytes(data[27:30], 'little') & 0xffffff) + 1
                    return w, h
        except Exception:
            pass
        return None, None

    @staticmethod
    def _normalize_img_units_for_import(html: str) -> str:
        """
        HTML→DOCX 方向图片尺寸单位归一化。

        核心规则：
          CSS style 中显式设置的维度 = 用户意图的显示尺寸（以 pt 为准）。
          HTML 属性 height/width（无单位=px）只是浏览器预布局提示，当
          style 未设置对应维度时，不应将其当作显示尺寸强制写入。

        典型问题场景（TinyMCE 生成）：
          style="width: 146pt;"   ← 用户设定的显示宽度，单位 pt
          height="113"             ← 原始图片像素高度，非显示高度
          width="146pt"            ← 非标准 pt 单位属性

        错误处理：把 height="113" 当作显示高度 113px=84.75pt 写入 →
          最终 width=146pt, height=84.75pt，但图片原始高度是 2.99cm
          (≈84.75pt)，等于按 100% 高度 + 64% 宽度显示 → 长宽比失真。

        正确处理：
          - 只处理 CSS style 中明确声明的维度；
          - style 有 width 无 height → 只设 width，删 height 属性（高度自动等比）；
          - style 有 height 无 width → 只设 height，删 width 属性；
          - style 同时有两者 → 都设置，统一换算为 pt；
          - style 什么都没有 → 不修改。
          写入时：style 用 pt，属性用无单位整数 px（HTML 规范），两者换算一致。
        """
        PT_PER_PX = 72.0 / 96.0
        PX_PER_PT = 96.0 / 72.0

        def _parse_to_pt(val_str):
            if not val_str:
                return None
            s = val_str.strip().lower()
            try:
                if s.endswith('pt'):   return float(s[:-2])
                if s.endswith('px'):   return float(s[:-2]) * PT_PER_PX
                if s.endswith('in'):   return float(s[:-2]) * 72.0
                if s.endswith('cm'):   return float(s[:-2]) / 2.54 * 72.0
                if s.endswith('mm'):   return float(s[:-2]) / 25.4 * 72.0
                if s.endswith('%'):    return None
                return float(s) * PT_PER_PX   # 无单位 → px → pt
            except ValueError:
                return None

        def _process_img(m):
            tag = m.group(0)

            # ── 1. 仅从 CSS style 中提取显示尺寸（属性不参与显示尺寸决策）
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            style_str = style_m.group(1) if style_m else ''

            sw_m = re.search(r'\bwidth\s*:\s*([^;]+)',  style_str, re.IGNORECASE)
            sh_m = re.search(r'\bheight\s*:\s*([^;]+)', style_str, re.IGNORECASE)
            sw_pt = _parse_to_pt(sw_m.group(1).strip() if sw_m else None)
            sh_pt = _parse_to_pt(sh_m.group(1).strip() if sh_m else None)

            # style 里没有任何尺寸信息 → 尝试从 HTML 属性读取（纯整数=px）
            if sw_pt is None and sh_pt is None:
                attr_w = re.search(r'\bwidth="(\d+)"',  tag, re.IGNORECASE)
                attr_h = re.search(r'\bheight="(\d+)"', tag, re.IGNORECASE)
                if not attr_w and not attr_h:
                    return tag
                if attr_w:
                    sw_pt = int(attr_w.group(1)) * PT_PER_PX
                if attr_h:
                    sh_pt = int(attr_h.group(1)) * PT_PER_PX

            # style 中的单位已全为 pt → 无需处理
            if (sw_pt is not None and sh_pt is not None
                    and sw_m and 'pt' in sw_m.group(1).lower()
                    and sh_m and 'pt' in sh_m.group(1).lower()):
                return tag

            # ── 2. 重建 CSS style（只保留 style 中本来就有的维度）
            new_style = style_str
            new_style = re.sub(r'\bwidth\s*:[^;]+;?\s*',  '', new_style, flags=re.IGNORECASE)
            new_style = re.sub(r'\bheight\s*:[^;]+;?\s*', '', new_style, flags=re.IGNORECASE)
            new_style = new_style.strip().rstrip(';').strip()

            size_parts = []
            if sw_pt is not None:
                size_parts.append(f'width:{sw_pt:.2f}pt')
            if sh_pt is not None:
                size_parts.append(f'height:{sh_pt:.2f}pt')
            size_css = '; '.join(size_parts)
            new_style = (size_css + '; ' + new_style).rstrip('; ') if new_style else size_css

            # ── 3. 重建 tag（先替换 style，再删旧属性，最后追加新属性）
            if style_m:
                tag = re.sub(r'style="[^"]*"', f'style="{new_style}"',
                             tag, count=1, flags=re.IGNORECASE)
            else:
                tag = re.sub(r'(<img\b)', rf'\1 style="{new_style}"', tag, flags=re.IGNORECASE)

            tag = re.sub(r'\s+width="[^"]*"',  '', tag, flags=re.IGNORECASE)
            tag = re.sub(r'\s+height="[^"]*"', '', tag, flags=re.IGNORECASE)

            # 属性写为无单位整数 px，且只写 style 中有的维度
            # （style 没设 height → 不写 height 属性 → Spire 自动等比缩放高度）
            if sw_pt is not None:
                tag = re.sub(r'(<img\b)', rf'\1 width="{round(sw_pt * PX_PER_PT)}"',
                             tag, flags=re.IGNORECASE)
            if sh_pt is not None:
                tag = re.sub(r'(<img\b)', rf'\1 height="{round(sh_pt * PX_PER_PT)}"',
                             tag, flags=re.IGNORECASE)

            return tag

        return re.sub(r'<img\b[^>]*>', _process_img, html, flags=re.IGNORECASE | re.DOTALL)

    @staticmethod
    def _fix_exif_orientation(img_bytes: bytes) -> bytes:
        """
        将 JPEG 图片的 EXIF Orientation 信息烘焙到像素中并清除该标签。
        浏览器会自动应用 EXIF 方向，但 Spire 不处理，导致 DOCX 中图片旋转。
        非 JPEG、方向正常（orientation==1）或 Pillow 不可用时原样返回，不重新编码。
        """
        if img_bytes[:2] != b'\xff\xd8':   # 非 JPEG，跳过
            return img_bytes
        try:
            from PIL import Image as _PILImage, ImageOps as _ImageOps
            import io as _io
            img = _PILImage.open(_io.BytesIO(img_bytes))
            # 明确读取 Orientation 值，避免依赖 `rotated is img` 的 Pillow 版本差异
            try:
                orientation = img.getexif().get(0x0112, 1)
            except Exception:
                orientation = 1
            if orientation in (None, 1):
                return img_bytes   # 方向正常，无需重新编码
            rotated = _ImageOps.exif_transpose(img)
            buf = _io.BytesIO()
            rotated.save(buf, format='JPEG', quality=92)   # 不写 EXIF，清除方向标签
            return buf.getvalue()
        except Exception:
            return img_bytes

    @staticmethod
    def _guess_mime(data: bytes) -> str:
        """从文件头推断真实 MIME 类型，默认返回 image/png。"""
        if data[:8]  == b'\x89PNG\r\n\x1a\n': return 'image/png'
        if data[:2]  == b'\xff\xd8':            return 'image/jpeg'
        if data[:6]  in (b'GIF87a', b'GIF89a'):  return 'image/gif'
        if data[:2]  == b'BM':                    return 'image/bmp'
        if data[:4]  == b'RIFF' and data[8:12] == b'WEBP': return 'image/webp'
        return 'image/png'

    def _extract_base64_images(self, html_text: str, base_dir: str):
        """
        HTML→DOCX 方向的图片预处理，返回 (modified_html, temp_img_dir)。
        """
        import urllib.request

        MIME_TO_EXT = {
            'image/png':     '.png',
            'image/jpeg':    '.jpg',
            'image/gif':     '.gif',
            'image/bmp':     '.bmp',
            'image/webp':    '.webp',
            'image/svg+xml': '.svg',
            'image/tiff':    '.tif',
        }

        b64_re  = re.compile(
            r'data:(image/[a-zA-Z0-9+\-]+);base64,([A-Za-z0-9+/=\s]+)',
            re.DOTALL
        )
        img_tag_re = re.compile(r'<img\b[^>]*>', re.IGNORECASE | re.DOTALL)
        src_re     = re.compile(r'src="([^"]*)"', re.IGNORECASE)

        if '\\"' in html_text:
            html_text = html_text.replace('\\"', '"')
            logger.debug("   🔧 检测到 JSON 转义引号，已还原 \\\" → \"")
        matches = list(img_tag_re.finditer(html_text))
        if not matches:
            return html_text, None

        temp_img_dir = self._normalize_path(
            os.path.join(base_dir, f"b64tmp_{uuid.uuid4().hex[:8]}")
        )
        os.makedirs(temp_img_dir, exist_ok=True)

        url_cache  = {}
        patch_list = []
        has_any    = False

        for m in matches:
            tag    = m.group(0)
            src_m  = src_re.search(tag)
            src    = src_m.group(1).strip() if src_m else ''

            img_bytes = None

            b64_m = b64_re.match(src) if src else None
            if b64_m:
                img_bytes = self._decode_b64_safe(b64_m.group(2))

            elif src.startswith('http://') or src.startswith('https://'):
                if src in url_cache:
                    img_bytes = url_cache[src]
                else:
                    for attempt, timeout in enumerate((30, 60, 90), start=1):
                        try:
                            req = urllib.request.Request(
                                src, headers={'User-Agent': 'Mozilla/5.0'}
                            )
                            with urllib.request.urlopen(req, timeout=timeout) as resp:
                                img_bytes = resp.read()
                            url_cache[src] = img_bytes
                            logger.debug(f"   🌐 下载图片（第{attempt}次）：{src[:60]}  {len(img_bytes):,}B")
                            break
                        except Exception as e:
                            logger.debug(f"   ⚠️ 下载失败（第{attempt}次，timeout={timeout}s）{src[:60]}：{e}")
                            if attempt == 3:
                                url_cache[src] = None
                    if not img_bytes:
                        img_bytes = url_cache.get(src)

            elif src and os.path.exists(src):
                try:
                    with open(src, 'rb') as f:
                        img_bytes = f.read()
                except Exception as e:
                    logger.debug(f"   ⚠️ 读取本地图片失败：{e}")
            if not img_bytes:
                continue

            has_any = True

            real_mime = self._guess_mime(img_bytes[:16])

            SPIRE_UNSUPPORTED = {'image/webp', 'image/svg+xml', 'image/tiff',
                                 'image/bmp', 'image/gif'}
            if real_mime in SPIRE_UNSUPPORTED:
                try:
                    from PIL import Image as _PILImage, ImageOps as _ImageOps
                    import io as _io
                    pil_img = _PILImage.open(_io.BytesIO(img_bytes))
                    pil_img = _ImageOps.exif_transpose(pil_img)  # 先修正 EXIF 方向，再转 RGB（否则 convert 会丢弃 EXIF）
                    pil_img = pil_img.convert('RGB')
                    buf = _io.BytesIO()
                    pil_img.save(buf, format='JPEG', quality=92)
                    img_bytes = buf.getvalue()
                    real_mime = 'image/jpeg'
                    logger.debug(f"   🔄 {real_mime} 不受 Spire 支持，已用 Pillow 转换为 JPEG")
                except Exception as e:
                    logger.debug(f"   ⚠️ Pillow 转换失败（{real_mime}→JPEG）：{e}，保留原格式")
            # 修正 EXIF 方向（Spire 不处理 Orientation，浏览器会自动应用；无条件调用，函数内部自行判断）
            fixed = self._fix_exif_orientation(img_bytes)
            if fixed is not img_bytes:
                logger.debug(f"   🔄 已修正 EXIF Orientation")
                img_bytes = fixed

            ext   = MIME_TO_EXT.get(real_mime, '.png')
            fname = f"img_{len(patch_list):04d}{ext}"
            fpath = self._normalize_path(os.path.join(temp_img_dir, fname))
            with open(fpath, 'wb') as f:
                f.write(img_bytes)

            declared_mime = b64_m.group(1).lower() if b64_m else ''
            if declared_mime and declared_mime != real_mime:
                logger.debug(f"   🔧 MIME 校正：{declared_mime} → {real_mime}（{fname}）")
            phys_w, phys_h = self._read_image_wh(img_bytes)

            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE | re.DOTALL)
            style_w_px = None
            style_h_px = None
            style_val  = ''
            if style_m:
                style_val = re.sub(r'[\r\n]+\s*', ' ', style_m.group(1))
                w_m = re.search(
                    r'(?<![a-zA-Z\-])width\s*:\s*([\d.]+)\s*(px|pt|in|cm|mm)',
                    style_val, re.IGNORECASE
                )
                if w_m:
                    val  = float(w_m.group(1))
                    unit = w_m.group(2).lower()
                    if unit == 'px':  style_w_px = round(val)
                    elif unit == 'pt': style_w_px = round(val * 96 / 72)
                    elif unit == 'in': style_w_px = round(val * 96)
                    elif unit == 'cm': style_w_px = round(val / 2.54 * 96)
                    elif unit == 'mm': style_w_px = round(val / 25.4 * 96)
                h_m = re.search(
                    r'(?<![a-zA-Z\-])height\s*:\s*([\d.]+)\s*(px|pt|in|cm|mm)',
                    style_val, re.IGNORECASE
                )
                if h_m:
                    val  = float(h_m.group(1))
                    unit = h_m.group(2).lower()
                    if unit == 'px':  style_h_px = round(val)
                    elif unit == 'pt': style_h_px = round(val * 96 / 72)
                    elif unit == 'in': style_h_px = round(val * 96)
                    elif unit == 'cm': style_h_px = round(val / 2.54 * 96)
                    elif unit == 'mm': style_h_px = round(val / 25.4 * 96)

            # 读取 HTML width/height 属性（优先级低于 style，高于物理尺寸）
            attr_w_m2 = re.search(r'\bwidth="(\d+(?:\.\d+)?)"', tag, re.IGNORECASE)
            attr_w_px = round(float(attr_w_m2.group(1))) if attr_w_m2 else None
            attr_h_m2 = re.search(r'\bheight="(\d+(?:\.\d+)?)"', tag, re.IGNORECASE)
            attr_h_px = round(float(attr_h_m2.group(1))) if attr_h_m2 else None

            w_px = style_w_px or attr_w_px or phys_w
            declared_h = style_h_px or attr_h_px
            if declared_h:
                h_px = declared_h
            elif w_px and phys_w and phys_h and phys_w > 0:
                h_px = round(w_px * phys_h / phys_w)
            else:
                h_px = phys_h

            new_tag = tag
            if src_m:
                new_tag = (new_tag[:src_m.start()]
                           + f'src="{fpath}"'
                           + new_tag[src_m.end():])

            if w_px and h_px:
                new_tag = re.sub(r'\s+width="[^"]*"', '', new_tag, flags=re.IGNORECASE)
                new_tag = re.sub(r'\s+height="[^"]*"', '', new_tag, flags=re.IGNORECASE)
                new_tag = re.sub(r'(<img\b)', rf'\1 width="{w_px}" height="{h_px}"', new_tag)

                style_m2 = re.search(r'style="([^"]*)"', new_tag, re.IGNORECASE)
                if style_m2:
                    s = style_m2.group(1)
                    s = re.sub(r'\bmax-width\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = re.sub(r'\bmax-height\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = re.sub(r'\bwidth\s*:[^;]+;?',  '', s, flags=re.IGNORECASE)
                    s = re.sub(r'\bheight\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = s.rstrip('; ')
                    new_style = f"{s}; width:{w_px}px; height:{h_px}px".lstrip('; ')
                    new_tag = (new_tag[:style_m2.start()]
                               + f'style="{new_style}"'
                               + new_tag[style_m2.end():])

            patch_list.append((m.start(), m.end(), new_tag))
            logger.debug(f"   📤 {fname}（{real_mime}，{w_px}×{h_px}px）")
        if not has_any:
            shutil.rmtree(temp_img_dir, ignore_errors=True)
            return html_text, None

        result = list(html_text)
        for start, end, new_tag in reversed(patch_list):
            result[start:end] = list(new_tag)
        html_text = ''.join(result)

        logger.debug(f"   📦 共处理 {len(patch_list)} 张图片 → {temp_img_dir}")
        return html_text, temp_img_dir


# ------------------------------ 调用示例 ------------------------------
if __name__ == "__main__":
    converter = DocxHtmlConverter()

    # 示例1：DOCX转单文件HTML（自动判断是否需要分片）
    input_docx = r"C:\Users\you62\Desktop\测试.docx"
    output_html = r"C:\Users\you62\Desktop\index.html"
    html_content = converter.docx_to_single_html(input_docx, output_html)

    # 示例2：HTML文本转DOCX
    # if os.path.exists(output_html):
    #     with open(output_html, 'r', encoding='utf-8') as f:
    #         sample_html = f.read()
    #     converter.html_text_to_docx(sample_html, "output.docx")
    # else:
    #     logger.debug(f"错误：未找到HTML文件 {output_html}")