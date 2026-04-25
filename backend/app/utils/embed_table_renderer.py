"""
嵌入组件 —— 表格渲染器

支持的 payload 格式
-------------------
{
    "headers": ["姓名", "部门", "职位"],        # 表头列（为空则无表头行）
    "rows": [                                   # 数据行（字符串 or 数字）
        ["张三", "研发部", "工程师"],
        ["李四", "市场部", "经理"]
    ],
    "caption": "表1 人员信息",                  # 可选：表格标题（插入在表格上方）
    "has_header": true,                         # 默认 true；false 时 headers 作普通行
    "col_widths": [3.0, 4.0, 3.0],             # 可选：各列宽度（cm），null = 自适应
    "merge_cells": [                            # 可选：合并单元格 [行, 起始列, 结束列]
        [0, 0, 2]                               # 第 0 行 col0~col2 合并
    ],
    "style": {
        "header_bg":        "4472C4",           # 表头背景色（hex，不含 #）
        "header_color":     "FFFFFF",           # 表头字体颜色
        "border_color":     "BFBFBF",           # 边框颜色
        "font_name":        "仿宋",
        "header_font_size": 10,                 # pt
        "body_font_size":   9,                  # pt
        "alignment":        "center",           # left / center / right（表格整体居中）
        "cell_alignment":   "center"            # 单元格内文字对齐
    }
}
"""

from __future__ import annotations

import html
import logging
import time
from copy import deepcopy
from typing import Any, Dict, List, Optional

from app.utils.embed_marker import (
    EmbedSpec,
    TYPE_TABLE,
    register_docx_renderer,
    register_html_renderer,
)

logger = logging.getLogger(__name__)

# 大表格按 _PROGRESS_ROW_STEP 行打印一次进度
_PROGRESS_ROW_STEP = 200

# ─── 默认样式 ──────────────────────────────────────────────────────────────────

_DEFAULT_STYLE: Dict[str, Any] = {
    "header_bg":        "4472C4",
    "header_color":     "FFFFFF",
    "border_color":     "BFBFBF",
    "font_name":        "仿宋",
    "header_font_size": 10,
    "body_font_size":   9,
    "alignment":        "center",
    "cell_alignment":   "center",
}


def _merge_style(user: Optional[Dict]) -> Dict[str, Any]:
    if not user:
        return dict(_DEFAULT_STYLE)
    return {**_DEFAULT_STYLE, **user}


# ─── DOCX 通用工具（延迟导入 docx，运行时才加载） ─────────────────────────────

def _hex_to_rgb(hex_color: str):
    """延迟导入 RGBColor，避免模块加载时依赖 docx。"""
    from docx.shared import RGBColor
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _align_enum(alignment: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    return {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(str(alignment).lower(), WD_ALIGN_PARAGRAPH.CENTER)


def _set_cell_border(cell, color: str) -> None:
    """给单元格四边设置单线边框。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color.upper().lstrip("#"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill_color: str) -> None:
    """设置单元格背景色。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_color.upper().lstrip("#"))
    tcPr.append(shd)


def _write_cell(
    cell,
    text: str,
    font_name: str,
    font_size: int,
    color: str,
    bold: bool,
    align,
    bg: Optional[str] = None,
    border_color: str = "BFBFBF",
) -> None:
    """写入单元格文字 + 样式。（兼容老调用方；内部走慢路径）"""
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt

    _set_cell_border(cell, border_color)
    if bg:
        _set_cell_shading(cell, bg)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    para = cell.paragraphs[0]
    para.alignment = align
    para.clear()
    run = para.add_run(str(text))
    run.font.name   = font_name
    run.font.size   = Pt(font_size)
    run.font.bold   = bold
    run.font.color.rgb = _hex_to_rgb(color)
    # 中文字体需同时设置 eastAsia
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), font_name
    )


def _build_tc_borders_template(border_color: str):
    """预先构造一个 w:tcBorders 元素，循环里 deepcopy 即可，避免每个单元格重复创建 5 个 OxmlElement。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    color = border_color.upper().lstrip("#")
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    return tcBorders


def _build_shd_template(fill_color: str):
    """预先构造一个 w:shd 背景元素。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_color.upper().lstrip("#"))
    return shd


def _fast_write_cell(
    cell,
    text: str,
    font_name: str,
    pt_size,
    rgb,
    bold: bool,
    align,
    tc_borders_tpl,
    shd_tpl=None,
    east_asia_qn=None,
    vcenter=None,
) -> None:
    """
    高频路径下使用的快速单元格写入：
      - 不重复创建 OxmlElement，统一用 deepcopy(模板)
      - 不重复解析 hex / 构造 Pt / RGBColor，全部由调用方预计算
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcPr.append(deepcopy(tc_borders_tpl))
    if shd_tpl is not None:
        tcPr.append(deepcopy(shd_tpl))
    if vcenter is not None:
        cell.vertical_alignment = vcenter

    para = cell.paragraphs[0]
    para.alignment = align
    # 新建 cell 时段落里只有一个空 run，省去 clear() 调用（其内部会遍历删除）
    run = para.add_run(text)
    font = run.font
    font.name = font_name
    font.size = pt_size
    font.bold = bold
    font.color.rgb = rgb
    # 中文字体：eastAsia
    if east_asia_qn is not None:
        run._element.get_or_add_rPr().get_or_add_rFonts().set(east_asia_qn, font_name)


# ─── 表格页面宽度工具 ──────────────────────────────────────────────────────────

def _sectPr_content_width(sectPr) -> int:
    """从 sectPr XML 元素计算正文区域宽度（twips）。失败返回 0。"""
    from docx.oxml.ns import qn
    pgSz  = sectPr.find(qn("w:pgSz"))
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgSz is None:
        return 0
    page_w = int(pgSz.get(qn("w:w"), 0))
    left   = int(pgMar.get(qn("w:left"),  1440) if pgMar is not None else 1440)
    right  = int(pgMar.get(qn("w:right"), 1440) if pgMar is not None else 1440)
    return max(0, page_w - left - right)


def _get_section_content_width_for_tbl(tbl_elem, document) -> int:
    """获取 tbl_elem **所在节**的正文区域宽度（twips）。

    OOXML 规则：段落 pPr 内的 sectPr 是该节的最后一段的标记；
    表格所属的节，由 tbl_elem 之后第一个含 sectPr 的段落决定。
    找不到时退回到文档级 sectPr，再找不到则调用 _get_page_content_width_twips。
    """
    from docx.oxml.ns import qn

    try:
        body = tbl_elem.getparent()
        if body is None:
            raise ValueError("tbl_elem 无父节点")

        # 扫描 tbl_elem 之后的兄弟节点，找最近的段落级 sectPr
        passed_tbl = False
        for child in body:
            if child is tbl_elem:
                passed_tbl = True
                continue
            if not passed_tbl:
                continue
            if child.tag != qn("w:p"):
                continue
            pPr = child.find(qn("w:pPr"))
            if pPr is None:
                continue
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is None:
                continue
            w = _sectPr_content_width(sectPr)
            if w > 0:
                logger.debug(
                    "_get_section_content_width_for_tbl: 段落级 sectPr → content=%d twips", w
                )
                return w

        # 回退：文档级 sectPr（body 的直接子元素）
        body_sectPr = body.find(qn("w:sectPr"))
        if body_sectPr is not None:
            w = _sectPr_content_width(body_sectPr)
            if w > 0:
                logger.debug(
                    "_get_section_content_width_for_tbl: 文档级 sectPr → content=%d twips", w
                )
                return w
    except Exception as _e:
        logger.debug("_get_section_content_width_for_tbl 异常: %s", _e)

    # 最终兜底
    return _get_page_content_width_twips(document)


def _get_page_content_width_twips(document) -> int:
    """获取文档正文区域宽度，单位 twips。

    优先直接读取 XML 中的 w:pgSz / w:pgMar（本身已是 twips，无需换算）；
    其次回退到 python-docx sections API（EMU → twips）；
    均失败时兜底返回 9072（A4 标准正文宽度，页宽 11906 - 左右各 1440）。
    """
    from docx.oxml.ns import qn

    # 方法一：直接解析 body 末尾的文档级 sectPr（twips，精度最高）
    try:
        body = document.element.body
        sectPr = body.find(qn("w:sectPr"))
        if sectPr is not None:
            pgSz  = sectPr.find(qn("w:pgSz"))
            pgMar = sectPr.find(qn("w:pgMar"))
            if pgSz is not None:
                page_w = int(pgSz.get(qn("w:w"), 0))
                if pgMar is not None:
                    left  = int(pgMar.get(qn("w:left"),  1440))
                    right = int(pgMar.get(qn("w:right"), 1440))
                else:
                    left = right = 1440
                content = page_w - left - right
                if content > 0:
                    logger.debug(
                        "_get_page_content_width_twips: sectPr → page=%d left=%d right=%d content=%d",
                        page_w, left, right, content,
                    )
                    return content
    except Exception as _e:
        logger.debug("_get_page_content_width_twips sectPr 解析失败: %s", _e)

    # 方法二：python-docx sections API（EMU ÷ 635 → twips）
    try:
        section = document.sections[0]
        pw = section.page_width
        lm = section.left_margin
        rm = section.right_margin
        if pw is not None and lm is not None and rm is not None:
            emu = int(pw) - int(lm) - int(rm)
            if emu > 0:
                content = max(100, emu // 635)
                logger.debug("_get_page_content_width_twips: sections API → content=%d twips", content)
                return content
    except Exception as _e:
        logger.debug("_get_page_content_width_twips sections API 失败: %s", _e)

    logger.debug("_get_page_content_width_twips: 回退兜底 9072")
    return 9072


def _apply_table_page_width(tbl_elem, col_count: int, col_twips: List[int]) -> None:
    """将 Word 表格宽度信息（tblW / tblGrid / 每格 tcW）统一设置为给定列宽列表。

    所有宽度单位均为 twips (dxa)。调用方负责：
    - col_twips 长度 == col_count
    - sum(col_twips) == 目标总宽（通常为页面正文宽度）
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    total_twips = sum(col_twips)

    # 1. 更新 w:tblW（表格总宽）
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    old_tblW = tblPr.find(qn("w:tblW"))
    if old_tblW is not None:
        tblPr.remove(old_tblW)
    tblW_el = OxmlElement("w:tblW")
    tblW_el.set(qn("w:w"), str(total_twips))
    tblW_el.set(qn("w:type"), "dxa")
    # OOXML 规定 tblW 须在 tblStyle 之后，插到 tblStyle 后面；没有 tblStyle 则插首位
    tblStyle = tblPr.find(qn("w:tblStyle"))
    _tblW_pos = (list(tblPr).index(tblStyle) + 1) if tblStyle is not None else 0
    tblPr.insert(_tblW_pos, tblW_el)

    # 2. 更新 w:tblGrid 各列宽（影响列定义）
    tblGrid = tbl_elem.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for i, gc in enumerate(tblGrid.findall(qn("w:gridCol"))):
            if i < col_count:
                gc.set(qn("w:w"), str(col_twips[i]))

    # 3. 更新每个单元格 tcW（合并单元格保持原合并宽不动，只改非合并格）
    for tr in tbl_elem.findall(qn("w:tr")):
        tcs = tr.findall(qn("w:tc"))
        for ci, tc in enumerate(tcs):
            if ci >= col_count:
                break
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            old_tcW = tcPr.find(qn("w:tcW"))
            if old_tcW is not None:
                tcPr.remove(old_tcW)
            tcW_el = OxmlElement("w:tcW")
            tcW_el.set(qn("w:w"), str(col_twips[ci]))
            tcW_el.set(qn("w:type"), "dxa")
            tcPr.insert(0, tcW_el)


# ─── DOCX 渲染器 ───────────────────────────────────────────────────────────────

def render_table_to_docx(
    spec: EmbedSpec,
    paragraph,
    document,
    *,
    page_twips_override: Optional[int] = None,
) -> None:
    """
    用 spec.payload 中的表格数据，在 paragraph 所在位置插入真实 Word 表格。
    插入完成后移除占位符段落。

    Args:
        page_twips_override: 显式指定正文宽度（twips）。
            并行渲染场景下由主进程预计算传入，避免在临时 Document 中
            通过 _get_section_content_width_for_tbl 取到默认页宽导致尺寸错位。

    性能优化要点：
        - 所有样式相关对象（RGBColor/Pt/border tcPr 模板/shd 模板）在循环外
          预计算一次，循环内仅做 deepcopy + 赋值
        - 整行取 cells 后顺序遍历，避免反复 table.cell(ri, ci) 索引
        - 使用 lxml getnext/getparent，定位占位符位置不再 O(N)
        - 大表格按 _PROGRESS_ROW_STEP 行打印进度日志，便于观察
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    t_total = time.perf_counter()

    payload = spec.payload or {}
    style   = _merge_style(payload.get("style"))

    headers: List[str]       = [str(h) for h in (payload.get("headers") or [])]
    rows:    List[List[str]] = [[str(c) for c in r] for r in (payload.get("rows") or [])]
    caption: str             = payload.get("caption") or ""
    has_header: bool         = bool(payload.get("has_header", True)) and bool(headers)
    col_widths: Optional[List[float]] = payload.get("col_widths")
    merge_instructions: List = payload.get("merge_cells") or []

    # ── 列数 ─────────────────────────────────────────────────────────────────
    col_count = len(headers) if headers else (max((len(r) for r in rows), default=1))
    if col_count == 0:
        logger.warning("render_table_to_docx: 列数为 0，跳过 embed_id=%s", spec.embed_id)
        return

    # ── 构建全部行 ────────────────────────────────────────────────────────────
    all_rows: List[List[str]] = []
    is_header_row: List[bool] = []
    if has_header:
        all_rows.append(headers)
        is_header_row.append(True)
    for r in rows:
        padded = r + [""] * max(0, col_count - len(r))
        all_rows.append(padded[:col_count])
        is_header_row.append(False)

    if not all_rows:
        logger.warning("render_table_to_docx: 无任何行数据，跳过 embed_id=%s", spec.embed_id)
        return

    total_rows = len(all_rows)
    align_para = _align_enum(style["cell_alignment"])
    font_name  = style["font_name"]

    # ── 预计算样式对象（循环内零分配） ────────────────────────────────────────
    header_pt   = Pt(int(style["header_font_size"]))
    body_pt     = Pt(int(style["body_font_size"]))
    header_rgb  = _hex_to_rgb(style["header_color"])
    body_rgb    = _hex_to_rgb("000000")
    border_tpl  = _build_tc_borders_template(style["border_color"])
    header_shd  = _build_shd_template(style["header_bg"]) if style.get("header_bg") else None
    east_asia   = qn("w:eastAsia")
    vcenter     = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # ── 占位符父节点定位（lxml O(1)）─────────────────────────────────────────
    p_elem = paragraph._element
    parent = p_elem.getparent()
    if parent is None:
        logger.warning("render_table_to_docx: 占位段落已脱离父节点 embed_id=%s", spec.embed_id)
        return
    insert_idx = parent.index(p_elem)  # lxml 原生 index，比 list(parent).index 快

    # ── 可选：在占位段落前插入标题段落 ───────────────────────────────────────
    if caption:
        cap_p = OxmlElement("w:p")
        cap_r = OxmlElement("w:r")
        cap_rPr = OxmlElement("w:rPr")
        cap_rPr.append(OxmlElement("w:b"))
        cap_r.append(cap_rPr)
        cap_t = OxmlElement("w:t")
        cap_t.text = caption
        cap_r.append(cap_t)
        cap_p.append(cap_r)
        parent.insert(insert_idx, cap_p)
        insert_idx += 1   # 占位符已后移一位

    # ── 创建 Word 表格 ────────────────────────────────────────────────────────
    t_create = time.perf_counter()
    table = document.add_table(rows=total_rows, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_elem = table._tbl
    tbl_elem.getparent().remove(tbl_elem)    # 从 body 末尾摘除
    parent.insert(insert_idx, tbl_elem)       # 插到占位符所在位置
    create_cost = time.perf_counter() - t_create

    # ── 列宽（自适应页面正文宽度） ───────────────────────────────────────────
    # 列宽设置推迟到单元格填充之后执行，确保 tcPr 已全部建立

    # ── 填充单元格（核心循环，热路径） ───────────────────────────────────────
    t_fill = time.perf_counter()
    table_rows = table.rows  # 缓存属性引用
    for ri in range(total_rows):
        row_data = all_rows[ri]
        is_hdr   = is_header_row[ri]
        if is_hdr:
            pt_size, rgb, shd = header_pt, header_rgb, header_shd
        else:
            pt_size, rgb, shd = body_pt, body_rgb, None
        row_cells = table_rows[ri].cells   # 整行一次性取出
        for ci in range(col_count):
            _fast_write_cell(
                cell           = row_cells[ci],
                text           = row_data[ci],
                font_name      = font_name,
                pt_size        = pt_size,
                rgb            = rgb,
                bold           = is_hdr,
                align          = align_para,
                tc_borders_tpl = border_tpl,
                shd_tpl        = shd,
                east_asia_qn   = east_asia,
                vcenter        = vcenter,
            )
        # 大表格分阶段进度日志
        if total_rows >= _PROGRESS_ROW_STEP and (
            (ri + 1) % _PROGRESS_ROW_STEP == 0 or ri == total_rows - 1
        ):
            elapsed = time.perf_counter() - t_fill
            done = ri + 1
            speed = done / elapsed if elapsed > 0 else 0.0
            eta = (total_rows - done) / speed if speed > 0 else 0.0
            logger.info(
                "render_table_to_docx 写入进度 embed_id=%s %d/%d 行"
                " 已耗时=%.2fs 速度=%.0f行/s ETA≈%.2fs",
                spec.embed_id, done, total_rows, elapsed, speed, eta,
            )
    fill_cost = time.perf_counter() - t_fill

    # ── 自适应页面正文宽度 ────────────────────────────────────────────────────
    # 单元格填充完成后统一设置列宽，确保所有 tcPr 已建立。
    # 无论是否指定 col_widths，表格总宽均与页面正文宽度对齐；
    # 有 col_widths 时按比例缩放，无 col_widths 时等分。
    # 两种情况均使用 autofit=False（固定布局），防止 Word 按内容收缩列宽。
    if page_twips_override is not None:
        page_twips = page_twips_override
    else:
        page_twips = _get_section_content_width_for_tbl(tbl_elem, document)
    if col_widths:
        raw = [max(0.01, float(w)) for w in col_widths[:col_count]]
        while len(raw) < col_count:
            raw.append(raw[-1] if raw else 1.0)
        total_raw = sum(raw)
        scaled = [int(page_twips * c / total_raw) for c in raw]
        # 修正整数截断累计误差（末列吸收）
        scaled[-1] = max(1, page_twips - sum(scaled[:-1]))
    else:
        base = page_twips // col_count
        scaled = [base] * col_count
        scaled[-1] = max(1, page_twips - base * (col_count - 1))
    _apply_table_page_width(tbl_elem, col_count, scaled)
    # 固定布局：Word 严格按 tcW 渲染，不会按内容收缩，保证表格撑满页面正文宽度
    table.autofit = False
    logger.debug(
        "render_table_to_docx 宽度设置: page_twips=%d col_count=%d scaled=%s",
        page_twips, col_count, scaled,
    )

    # ── 合并单元格（[行, 起始列, 结束列]） ────────────────────────────────────
    for instr in merge_instructions:
        if len(instr) != 3:
            continue
        ri, c_start, c_end = int(instr[0]), int(instr[1]), int(instr[2])
        if ri >= total_rows or c_end >= col_count:
            continue
        try:
            table.cell(ri, c_start).merge(table.cell(ri, c_end))
        except Exception as e:
            logger.warning(f"合并单元格失败 [{ri},{c_start},{c_end}]：{e}")

    # ── 移除占位符段落（此时它已被推到 insert_idx+1 位置） ────────────────────
    parent.remove(p_elem)

    total_cost = time.perf_counter() - t_total
    logger.info(
        "render_table_to_docx 完成 embed_id=%s rows=%d cols=%d"
        " 创建=%.2fs 填充=%.2fs 总计=%.2fs",
        spec.embed_id, total_rows, col_count, create_cost, fill_cost, total_cost,
    )


# ─── HTML 预览渲染器（无外部依赖，可单独测试） ─────────────────────────────────

def render_table_to_html(spec: EmbedSpec) -> str:
    """生成在富文本编辑器中可见的 HTML 预览表格。"""
    payload = spec.payload or {}
    style   = _merge_style(payload.get("style"))

    headers: List[str] = [str(h) for h in (payload.get("headers") or [])]
    rows: List[List]   = payload.get("rows") or []
    caption: str       = payload.get("caption") or ""
    has_header: bool   = bool(payload.get("has_header", True)) and bool(headers)

    header_bg    = f"#{style['header_bg']}"
    header_color = f"#{style['header_color']}"
    border_color = f"#{style['border_color']}"
    font_name    = style["font_name"]
    h_size       = style["header_font_size"]
    b_size       = style["body_font_size"]
    cell_align   = style["cell_alignment"]

    td_base = (
        f"padding:4px 8px; border:1px solid {border_color}; "
        f"font-family:{html.escape(font_name)},serif; text-align:{cell_align};"
    )

    lines: List[str] = []
    if caption:
        lines.append(
            f'<p class="excel-swk-start" style="font-weight:bold; margin:4px 0; '
            f'font-family:{html.escape(font_name)},serif;">'
            f'{html.escape(caption)}</p>'
        )

    lines.append(
        f'<table style="border-collapse:collapse; width:100%; '
        f'font-size:{b_size}pt; margin:4px 0;" data-embed-preview="table">'
    )

    if has_header:
        lines.append("<thead><tr>")
        for h in headers:
            lines.append(
                f'<th style="{td_base} background:{header_bg}; '
                f'color:{header_color}; font-size:{h_size}pt; '
                f'font-weight:bold;">{html.escape(str(h))}</th>'
            )
        lines.append("</tr></thead>")

    if rows:
        col_count = len(headers) if headers else (len(rows[0]) if rows else 0)
        lines.append("<tbody>")
        for row in rows:
            lines.append("<tr>")
            for ci in range(col_count):
                val = row[ci] if ci < len(row) else ""
                lines.append(f'<td style="{td_base}">{html.escape(str(val))}</td>')
            lines.append("</tr>")
        lines.append("</tbody>")

    lines.append("</table>")
    return "\n".join(lines)


# ─── 并行渲染入口（模块级顶层函数，可被 ProcessPoolExecutor pickle） ──────────

def build_table_xml(spec_dict: Dict[str, Any], page_twips: int) -> List[bytes]:
    """
    【子进程入口】在临时隔离 Document 中独立渲染单个表格。

    返回该表格渲染过程中新增的全部元素（caption 段落 + w:tbl）的 XML 字节序列，
    主进程按顺序 splice 到占位段落位置即可。

    主进程必须预先计算正文宽度并传入 page_twips（twips），
    因为临时 Document 的默认页宽与主文档可能不同。

    - 必须是模块级顶层函数（ProcessPoolExecutor 的 pickle 要求）
    - 此函数不修改主 Document，worker 进程间完全隔离
    """
    import os
    import threading
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    spec = EmbedSpec.from_dict(spec_dict)
    t_start = time.perf_counter()
    pid = os.getpid()

    # 每 10 秒打印一次心跳，表明该 worker 进程仍在运行
    done_event = threading.Event()

    def _heartbeat():
        while not done_event.wait(10):
            elapsed = time.perf_counter() - t_start
            logger.info(
                "build_table_xml [pid=%d] 渲染中 embed_id=%s 已耗时=%.1fs",
                pid, spec.embed_id, elapsed,
            )

    monitor = threading.Thread(target=_heartbeat, daemon=True)
    monitor.start()

    try:
        temp_doc = Document()
        temp_para = temp_doc.add_paragraph("")
        body = temp_doc.element.body

        # 记录渲染前 body 的已有子元素，用于渲染后差分出新增元素
        before_ids = {id(c) for c in body}

        render_table_to_docx(spec, temp_para, temp_doc, page_twips_override=page_twips)

        sectPr_tag = qn("w:sectPr")
        new_elems = [c for c in body if id(c) not in before_ids and c.tag != sectPr_tag]
        result = [etree.tostring(c, xml_declaration=False) for c in new_elems]
    finally:
        done_event.set()
        monitor.join(timeout=1)

    elapsed = time.perf_counter() - t_start
    logger.info(
        "build_table_xml [pid=%d] 完成 embed_id=%s 耗时=%.2fs 元素数=%d",
        pid, spec.embed_id, elapsed, len(result),
    )
    return result


# ─── 注册 ─────────────────────────────────────────────────────────────────────

register_html_renderer(TYPE_TABLE, render_table_to_html)
register_docx_renderer(TYPE_TABLE, render_table_to_docx)

logger.debug("embed_table_renderer: HTML 和 DOCX 渲染器已注册")