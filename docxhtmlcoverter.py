from spire.doc import *
from spire.doc.common import *
import os
import zipfile
import shutil
import re
import base64
import tempfile
import uuid
from docx import Document as PythonDocx          # python-docx：用于切分
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree


class DocxHtmlConverter:
    """
    DOCX与HTML互转工具类
    核心特性：
    1. 所有路径强制使用绝对路径，脱离工作目录依赖
    2. 临时目录在每次 docx_to_single_html 调用时生成唯一ID，避免多线程/多调用冲突
    3. 图片顺序解析覆盖正文、页眉、页脚、脚注、尾注等所有XML区域
    4. 降级兜底：Spire生成图片路径（find_actual_img_dir递归查找）
    5. 增强路径校验和异常处理
    6. 超限文档自动分片处理（段落>450或表格>20时触发）
    7. 大表格按行拆分，合并HTML时自动还原
    8. 合并时去除页眉页脚
    9. 分片流程：chunk不内嵌图片 → 流式合并 → 统一内嵌，避免内存溢出
    """

    def __init__(self):
        self.default_image_format = 0  # 0=PNG，1=JPG，2=BMP，3=GIF
        self.html_validation_type = XHTMLValidationType.none
        self.MAX_PARAGRAPHS = 450
        self.MAX_TABLES = 20

    # ------------------------------------------------------------------ #
    #  路径工具                                                             #
    # ------------------------------------------------------------------ #

    def _normalize_path(self, path):
        """【内部方法】统一路径格式并转为绝对路径"""
        if not path:
            return ""
        abs_path = os.path.abspath(path)
        return abs_path.replace('\\', '/').replace('//', '/')

    def _make_temp_dir_prefix(self):
        """
        【修复问题5】每次调用生成新的唯一前缀，避免多次调用或多线程复用同一前缀。
        原设计在 __init__ 中只生成一次，导致并发或重复调用时临时目录冲突。
        """
        return f"spire_temp_{uuid.uuid4().hex[:8]}"

    # ------------------------------------------------------------------ #
    #  图片工具                                                             #
    # ------------------------------------------------------------------ #

    def _get_image_order_from_docx(self, docx_path):
        """
        【内部方法】解析DOCX，提取图片在文档中的显示顺序
        覆盖范围：正文、页眉、页脚、脚注、尾注等所有XML区域

        修复：将 id= 宽泛正则收窄为 r:embed / r:link，避免误匹配
              非图片关系节点（bookmark、style等）的 id 属性。

        返回值：(image_order_list, rids_to_imgname_map)
            image_order_list  - 按文档顺序排列的原始图片文件名列表（去重）
            rids_to_imgname_map - {xml_file: {rId: img_filename}} 供调用方使用
        """
        image_order = []
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                all_files = zip_file.namelist()

                target_xml_files = [
                    f for f in all_files
                    if re.match(
                        r'word/(document|header\d*|footer\d*|footnotes|endnotes)\.xml$', f
                    )
                ]

                all_rels = {}
                for xml_file in target_xml_files:
                    xml_basename = os.path.basename(xml_file)
                    rels_path = f'word/_rels/{xml_basename}.rels'
                    if rels_path in all_files:
                        rels_content = zip_file.read(rels_path).decode('utf-8')
                        rel_pattern = re.compile(
                            r'<Relationship\s+Id="(rId\d+)"\s+Type="[^"]*image[^"]*"\s+Target="([^"]+)"'
                        )
                        rels_for_this_xml = {}
                        for match in rel_pattern.finditer(rels_content):
                            r_id, target = match.group(1), match.group(2)
                            rels_for_this_xml[r_id] = os.path.basename(target)
                        all_rels[xml_file] = rels_for_this_xml

                seen = set()
                id_pattern = re.compile(r'r:(?:embed|link|id)="(rId\d+)"')
                for xml_file in target_xml_files:
                    if xml_file not in all_rels:
                        continue
                    xml_content = zip_file.read(xml_file).decode('utf-8')
                    rels_map = all_rels[xml_file]
                    for match in id_pattern.finditer(xml_content):
                        r_id = match.group(1)
                        if r_id in rels_map:
                            img_name = rels_map[r_id]
                            if img_name not in seen:
                                seen.add(img_name)
                                image_order.append(img_name)

        except Exception as e:
            print(f"⚠️ 解析图片顺序失败：{e}，将使用文件名排序")
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                image_order = sorted(
                    os.path.basename(f.filename)
                    for f in zip_file.infolist()
                    if f.filename.startswith('word/media/') and not f.is_dir()
                )

        print(f"✅ 解析到图片显示顺序：{image_order}")
        return image_order

    def _extract_original_images(self, docx_path, output_img_dir):
        """【内部方法】从DOCX中提取原始无压缩图片（强制绝对路径）"""
        output_img_dir = self._normalize_path(output_img_dir)

        if os.path.exists(output_img_dir):
            shutil.rmtree(output_img_dir, ignore_errors=True)
        os.makedirs(output_img_dir, exist_ok=True)

        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            for file_info in zip_file.infolist():
                if file_info.filename.startswith('word/media/') and not file_info.is_dir():
                    img_filename = os.path.basename(file_info.filename)
                    save_path = self._normalize_path(
                        os.path.join(output_img_dir, img_filename)
                    )
                    with open(save_path, 'wb') as f:
                        f.write(zip_file.read(file_info.filename))

        return [
            f for f in os.listdir(output_img_dir)
            if os.path.isfile(os.path.join(output_img_dir, f))
        ]

    def _find_actual_img_dir(self, base_dir):
        """
        【内部方法】递归查找第一个实际包含图片文件的目录
        用于兜底：Spire有时将图片输出到嵌套子目录中
        """
        img_exts = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
        for root, dirs, files in os.walk(base_dir):
            if any(f.lower().endswith(img_exts) for f in files):
                return root
        return base_dir

    def _convert_emf_to_png(self, emf_path):
        """
        【内部方法】将 EMF/WMF 文件转换为 PNG，返回 PNG 文件路径。
        优先调用 LibreOffice 命令行（soffice）光栅化，失败则降级到
        spire.doc.common.Image。转换结果缓存在原文件同目录下。
        转换失败返回 None。
        """
        import subprocess

        emf_path = self._normalize_path(emf_path)
        emf_dir  = os.path.dirname(emf_path)
        emf_stem = os.path.splitext(os.path.basename(emf_path))[0]
        png_path = self._normalize_path(os.path.join(emf_dir, emf_stem + '_converted.png'))

        if os.path.exists(png_path):
            return png_path  # 已转换过，直接复用

        soffice_candidates = [
            'libreoffice',
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            '/usr/bin/soffice',
            '/usr/lib/libreoffice/program/soffice',
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        ]

        for soffice in soffice_candidates:
            try:
                result = subprocess.run(
                    [soffice, '--headless', '--convert-to', 'png', '--outdir', emf_dir, emf_path],
                    capture_output=True,
                    timeout=30,
                )
                lo_output = self._normalize_path(os.path.join(emf_dir, emf_stem + '.png'))
                if result.returncode == 0 and os.path.exists(lo_output):
                    if lo_output != png_path:
                        os.rename(lo_output, png_path)
                    print(f"   🔄 EMF→PNG（LibreOffice）：{os.path.basename(emf_path)}")
                    return png_path
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue
            except Exception as e:
                print(f"   ⚠️ LibreOffice 转换异常：{e}")
                continue

        try:
            from spire.doc.common import Image as SpireImage
            img = SpireImage.FromFile(emf_path)
            img.Save(png_path)
            if os.path.exists(png_path):
                print(f"   🔄 EMF→PNG（Spire.Image）：{os.path.basename(emf_path)}")
                return png_path
        except Exception as e:
            print(f"   ⚠️ Spire.Image 转换失败：{e}")

        print(f"   ⚠️ EMF→PNG 所有方案均失败，跳过：{os.path.basename(emf_path)}")
        return None

    def _image_to_base64(self, img_path):
        """【内部方法】将图片文件转为Base64编码（带MIME前缀）。
        对 EMF/WMF 格式自动先转为 PNG，再做 base64。
        """
        img_path = self._normalize_path(img_path)
        try:
            if not os.path.exists(img_path):
                print(f"⚠️ 图片文件不存在（绝对路径）：{img_path}")
                return ""

            img_ext = os.path.splitext(img_path)[1].lower()
            if img_ext in ('.emf', '.wmf'):
                converted = self._convert_emf_to_png(img_path)
                if converted and os.path.exists(converted):
                    img_path = converted
                    img_ext  = '.png'
                else:
                    return ""

            with open(img_path, 'rb') as f:
                img_data = f.read()

            mime_map = {
                '.png':  'image/png',
                '.jpg':  'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif':  'image/gif',
                '.bmp':  'image/bmp',
                '.webp': 'image/webp',
                '.svg':  'image/svg+xml',
            }
            mime_type = mime_map.get(img_ext, 'image/png')
            return f"data:{mime_type};base64,{base64.b64encode(img_data).decode('utf-8')}"

        except Exception as e:
            print(f"⚠️ 图片 {img_path} 转Base64失败：{e}")
            return ""

    def _build_spire_to_original_map(self, spire_img_names, image_display_order,
                                     original_img_dir, fallback_img_dir):
        """
        【修复问题1】构建 Spire生成图片名 → 原始图片绝对路径 的映射。

        原逻辑用位置索引（spire_img_names[i] → image_display_order[i]）映射，
        当 Spire 跳过、合并或重命名图片时，索引对不上导致 base64 错位。

        修复策略：
        1. 优先用文件名精确匹配（去扩展名后不区分大小写比较）。
           Spire 生成的文件名通常与原始文件名一致或只改了扩展名。
        2. 精确匹配失败时，按位置顺序降级（保留原兜底行为）。
        3. 位置索引也超出范围时，尝试在 fallback_img_dir 中按 spire 名查找。

        返回：{spire_name: abs_img_path}
        """
        result = {}

        # 建立原始文件名（去扩展名小写）→ 绝对路径 的查找表
        orig_stem_map = {}
        for orig_name in image_display_order:
            stem = os.path.splitext(orig_name)[0].lower()
            abs_path = self._normalize_path(os.path.join(original_img_dir, orig_name))
            if os.path.exists(abs_path):
                orig_stem_map[stem] = abs_path

        for idx, spire_name in enumerate(spire_img_names):
            spire_stem = os.path.splitext(spire_name)[0].lower()

            # 1. 精确匹配（去扩展名不区分大小写）
            if spire_stem in orig_stem_map:
                result[spire_name] = orig_stem_map[spire_stem]
                continue

            # 2. 位置索引降级
            if idx < len(image_display_order):
                candidate = self._normalize_path(
                    os.path.join(original_img_dir, image_display_order[idx])
                )
                if os.path.exists(candidate):
                    result[spire_name] = candidate
                    print(f"   ⚠️ {spire_name} 精确匹配失败，位置索引降级 → {image_display_order[idx]}")
                    continue

            # 3. fallback：在 spire 生成目录中按原名查找
            fallback = self._normalize_path(os.path.join(fallback_img_dir, spire_name))
            if os.path.exists(fallback):
                result[spire_name] = fallback
                print(f"   ⚠️ {spire_name} 降级到 Spire 生成目录")
            else:
                print(f"   ⚠️ {spire_name} 找不到对应原始图片，跳过")

        return result

    # ------------------------------------------------------------------ #
    #  DPI / 尺寸修正工具                                                   #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _extract_image_display_sizes(docx_path):
        """
        从 DOCX 的 document.xml 中提取每张图片的"显示尺寸"（即 Word 排版时的实际渲染宽高），
        返回 {原始图片文件名: (width_px, height_px)} 的映射，分辨率基准为 96 DPI。

        ── 为什么需要此方法 ──────────────────────────────────────────────────
        图片在 DOCX 中的渲染尺寸存储在 <wp:extent cx="..." cy="..."/>，单位为 EMU
        （English Metric Unit，914400 EMU = 1 inch）。这是 Word 排版的唯一权威来源，
        与图片文件本身的 DPI 元数据无关。

        Spire 把图片导出为 PNG 时，有时会把 PNG 的 DPI 元数据写成 96，但原始图片
        （如 Retina 截图）实际是 192 DPI，导致 HTML <img> 按像素渲染时显示为 2 倍大。
        同理，HTML→DOCX 方向，Spire 读 <img> 时如果没有显式 width/height，会用图片
        物理像素 ÷ 假定 DPI 算 EMU，同样产生 2 倍误差。

        通过从 <wp:extent> 读取 EMU 换算成 96 DPI 像素，可以强制锁定渲染尺寸，
        完全规避图片文件 DPI 元数据的干扰。

        ── 覆盖范围 ──────────────────────────────────────────────────────────
        扫描 word/document.xml（正文）。页眉/页脚/脚注中的图片也可扩展，
        但正文是绝大多数 2 倍问题的发生场景，优先处理。

        ── XML 结构 ──────────────────────────────────────────────────────────
        DrawingML（内联图片，最常见）：
            <w:drawing>
              <wp:inline>
                <wp:extent cx="914400" cy="685800"/>   ← 宽1in×0.75in
                <a:graphic>
                  <a:graphicData>
                    <pic:pic>
                      <pic:blipFill>
                        <a:blip r:embed="rId5"/>        ← 关联到 rId5 → image1.png
                      </pic:blipFill>
                    </pic:pic>
                  </a:graphicData>
                </a:graphic>
              </wp:inline>
            </w:drawing>

        DrawingML（浮动图片）：
            <wp:anchor>
              <wp:extent cx="..." cy="..."/>
              ...（同上）
            </wp:anchor>

        VML（旧格式，通常出现在兼容模式文档）：
            <v:shape style="width:72pt;height:54pt;...">
              <v:imagedata r:id="rId6" o:title="..."/>
            </v:shape>
            → 从 style 中解析 width/height

        返回：{img_filename: (width_px_int, height_px_int)}
              96 DPI 下：1 inch = 96 px，1 EMU = 96/914400 px
        """
        EMU_PER_INCH  = 914400
        PX_PER_INCH   = 96        # HTML 标准基准 DPI
        EMU_TO_PX     = PX_PER_INCH / EMU_PER_INCH

        WP_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        A_NS   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        R_NS   = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        VML_NS = 'urn:schemas-microsoft-com:vml'
        W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        size_map = {}   # {img_filename: (w_px, h_px)}

        try:
            with zipfile.ZipFile(docx_path, 'r') as zf:
                # 读取 rels，建立 rId → 图片文件名 的映射
                rels_xml = zf.read('word/_rels/document.xml.rels').decode('utf-8')
                rid_to_img = {}
                for m in re.finditer(
                    r'<Relationship\s+Id="(rId\d+)"\s+Type="[^"]*image[^"]*"\s+Target="([^"]+)"',
                    rels_xml
                ):
                    rid_to_img[m.group(1)] = os.path.basename(m.group(2))

                doc_xml = zf.read('word/document.xml').decode('utf-8')

            root = etree.fromstring(doc_xml.encode('utf-8'))

            # ── DrawingML：wp:inline / wp:anchor ──────────────────────
            for container_tag in (f'{{{WP_NS}}}inline', f'{{{WP_NS}}}anchor'):
                for container in root.iter(container_tag):
                    extent = container.find(f'{{{WP_NS}}}extent')
                    if extent is None:
                        continue
                    try:
                        cx = int(extent.get('cx', 0))
                        cy = int(extent.get('cy', 0))
                    except ValueError:
                        continue

                    w_px = round(cx * EMU_TO_PX)
                    h_px = round(cy * EMU_TO_PX)

                    # 找到对应的 r:embed（blip）
                    blip = container.find(
                        f'.//{{{A_NS}}}blip'
                    )
                    if blip is not None:
                        r_embed = blip.get(f'{{{R_NS}}}embed')
                        if r_embed and r_embed in rid_to_img:
                            fname = rid_to_img[r_embed]
                            if fname not in size_map:   # 同一文件多处引用取第一次出现
                                size_map[fname] = (w_px, h_px)

            # ── VML：v:shape + v:imagedata ────────────────────────────
            def _css_dim_to_px(val_str):
                """CSS 尺寸字符串 → 96 DPI px（整数）"""
                if not val_str:
                    return None
                val_str = val_str.strip().lower()
                try:
                    if val_str.endswith('pt'):
                        return round(float(val_str[:-2]) * PX_PER_INCH / 72)
                    if val_str.endswith('in'):
                        return round(float(val_str[:-2]) * PX_PER_INCH)
                    if val_str.endswith('cm'):
                        return round(float(val_str[:-2]) / 2.54 * PX_PER_INCH)
                    if val_str.endswith('mm'):
                        return round(float(val_str[:-2]) / 25.4 * PX_PER_INCH)
                    if val_str.endswith('px'):
                        return round(float(val_str[:-2]))
                except (ValueError, AttributeError):
                    pass
                return None

            for shape in root.iter(f'{{{VML_NS}}}shape'):
                imagedata = shape.find(f'{{{VML_NS}}}imagedata')
                if imagedata is None:
                    # v:imagedata 有时用 o:title 命名空间
                    for child in shape:
                        if child.tag.endswith('}imagedata') or child.tag == 'imagedata':
                            imagedata = child
                            break
                if imagedata is None:
                    continue

                r_id = imagedata.get(f'{{{R_NS}}}id')
                if not r_id or r_id not in rid_to_img:
                    continue

                fname = rid_to_img[r_id]
                if fname in size_map:
                    continue  # DrawingML 已处理，跳过

                style = shape.get('style', '')
                style_dict = {}
                for part in style.split(';'):
                    if ':' in part:
                        k, _, v = part.partition(':')
                        style_dict[k.strip().lower()] = v.strip()

                w_px = _css_dim_to_px(style_dict.get('width'))
                h_px = _css_dim_to_px(style_dict.get('height'))

                if w_px and h_px:
                    size_map[fname] = (w_px, h_px)

        except Exception as e:
            print(f"⚠️ 提取图片显示尺寸失败：{e}")

        print(f"📐 从 DOCX 提取到 {len(size_map)} 张图片的显示尺寸")
        return size_map

    def _fix_html_img_sizes(self, html_content: str, size_map: dict,
                             spire_img_names: list, image_display_order: list) -> str:
        """
        将 HTML 中每个 <img> 的 width/height 属性强制设置为从 DOCX <wp:extent> 读取的显示尺寸。

        ── 为什么只用 style="width:...;height:..." 不够 ─────────────────────
        HTML <img> 的实际渲染尺寸由以下优先级决定：
          1. style="width:Xpx; height:Ypx"   ← 最高优先级
          2. width="X" height="Y" 属性        ← 次优先级（像素）
          3. 图片文件本身的物理像素尺寸        ← 默认（受 DPI 影响！）

        Spire 生成的 HTML 有时只有 style，有时只有属性，有时两者都有但值不一致。
        本方法同时写入 style 内的 width/height 和 HTML 属性 width/height，
        确保所有浏览器和 Spire 反向读取时都使用正确的尺寸。

        ── 匹配策略 ──────────────────────────────────────────────────────────
        先按 spire_img_names[i] ↔ image_display_order[i] 的映射找到原始文件名，
        再从 size_map 中查对应的像素尺寸。
        """
        if not size_map:
            return html_content

        # 构建 spire生成名 → 原始文件名 的映射（与 _build_spire_to_original_map 逻辑一致）
        spire_to_orig = {}
        orig_stems = {
            os.path.splitext(n)[0].lower(): n
            for n in image_display_order
        }
        for idx, spire_name in enumerate(spire_img_names):
            spire_stem = os.path.splitext(spire_name)[0].lower()
            if spire_stem in orig_stems:
                spire_to_orig[spire_name] = orig_stems[spire_stem]
            elif idx < len(image_display_order):
                spire_to_orig[spire_name] = image_display_order[idx]

        def _replace_img_tag(m):
            tag = m.group(0)
            src_m = re.search(r'src="([^"]+)"', tag)
            if not src_m:
                return tag

            src_basename = os.path.basename(src_m.group(1))
            orig_name = spire_to_orig.get(src_basename)
            if not orig_name:
                return tag

            sizes = size_map.get(orig_name)
            if not sizes:
                # 尝试去扩展名匹配（Spire 可能改了扩展名）
                orig_stem = os.path.splitext(orig_name)[0].lower()
                for k, v in size_map.items():
                    if os.path.splitext(k)[0].lower() == orig_stem:
                        sizes = v
                        break
            if not sizes:
                return tag

            w_px, h_px = sizes

            # 1. 写入 / 替换 HTML 属性 width / height
            tag = re.sub(r'\s+width="[^"]*"', '', tag)
            tag = re.sub(r'\s+height="[^"]*"', '', tag)
            # 在 > 或第一个属性前插入
            tag = re.sub(r'(<img\b)', rf'\1 width="{w_px}" height="{h_px}"', tag)

            # 2. 写入 / 替换 style 内的 width / height
            style_m = re.search(r'style="([^"]*)"', tag)
            if style_m:
                style_str = style_m.group(1)
                style_str = re.sub(r'\bwidth\s*:[^;]+;?', '', style_str, flags=re.IGNORECASE)
                style_str = re.sub(r'\bheight\s*:[^;]+;?', '', style_str, flags=re.IGNORECASE)
                style_str = style_str.rstrip('; ')
                new_style  = f"{style_str}; width:{w_px}px; height:{h_px}px".lstrip('; ')
                tag = tag[:style_m.start()] + f'style="{new_style}"' + tag[style_m.end():]
            else:
                # 没有 style 属性，追加一个
                tag = re.sub(r'(<img\b)', rf'\1 style="width:{w_px}px; height:{h_px}px"', tag)
                # 上面已经 insert 了一次，避免重复，把多余的删掉
                tag = re.sub(r'(<img\b)(.*?)(<img\b)', r'\1\2', tag)  # 保险去重

            print(f"   📐 {src_basename} → {orig_name} 锁定尺寸 {w_px}×{h_px}px")
            return tag

        return re.sub(r'<img\b[^>]*>', _replace_img_tag, html_content, flags=re.IGNORECASE)

    # ------------------------------------------------------------------ #
    #  表格宽度修正（DOCX→HTML 方向）                                      #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _fix_html_table_widths(html_content: str,
                                content_width_pt: float = 467.0) -> str:
        """
        将 Spire 导出 HTML 中超出版心宽度的表格等比缩放至版心宽度以内。

        ── 问题原因 ─────────────────────────────────────────────────────────
        Spire 把 DOCX 表格宽度原样转为 HTML pt 值（如 width:659.65pt），
        A4 页面版心约 467pt（165mm，左右各 2.5cm 页边距），超出部分被截断。

        ── 处理范围 ─────────────────────────────────────────────────────────
        • <table> 标签的 style="width:Xpt" 或 width="X"
        • 同一表格内所有 <col width="Xpt">
        • 所有 <td>/<th> 的 style="width:Xpt" 或 width="X"
          （只按 table 级别的缩放比例等比缩减，不单独判断 td 是否超出）

        ── 单位支持 ─────────────────────────────────────────────────────────
        pt / px / in / cm / mm / 纯数字（默认 pt，与 Spire 输出一致）

        ── 缩放规则 ─────────────────────────────────────────────────────────
        • 表格宽度 ≤ content_width_pt → 不缩放
        • 表格宽度 > content_width_pt → ratio = content_width_pt / table_width
          所有尺寸 × ratio，结果保留 2 位小数，单位保持 pt

        ── content_width_pt 默认值说明 ──────────────────────────────────────
        A4 宽 595.28pt，左右页边距各 2.5cm（70.87pt），版心 = 595.28 - 141.74 ≈ 453pt。
        实践中 Spire 生成的 HTML 表格往往用文档设置宽度而非 A4 标准，
        默认取 467pt（约 165mm）作为安全上限，调用方可按实际文档覆盖。
        """

        # ── 单位换算：任意 CSS 宽度字符串 → pt ───────────────────────────
        def _to_pt(val_str: str) -> float | None:
            if not val_str:
                return None
            s = val_str.strip().lower()
            try:
                if s.endswith('pt'):  return float(s[:-2])
                if s.endswith('px'):  return float(s[:-2]) * 72 / 96
                if s.endswith('in'):  return float(s[:-2]) * 72
                if s.endswith('cm'):  return float(s[:-2]) / 2.54 * 72
                if s.endswith('mm'):  return float(s[:-2]) / 25.4 * 72
                return float(s)       # 纯数字默认 pt
            except ValueError:
                return None

        def _fmt(val_pt: float) -> str:
            """pt 值格式化为 2 位小数字符串"""
            return f"{val_pt:.2f}pt"

        # ── 从 style 字符串中提取 width 值（pt）──────────────────────────
        def _style_width_pt(style_str: str) -> float | None:
            m = re.search(r'\bwidth\s*:\s*([^;]+)', style_str, re.IGNORECASE)
            if m:
                return _to_pt(m.group(1).strip())
            return None

        # ── 替换 style 中的 width ─────────────────────────────────────────
        def _replace_style_width(style_str: str, new_pt: float) -> str:
            return re.sub(
                r'\bwidth\s*:\s*[^;]+',
                f'width:{_fmt(new_pt)}',
                style_str,
                flags=re.IGNORECASE
            )

        # ── 处理单个元素标签里的 width（style 优先，其次 HTML 属性）───────
        def _scale_tag_width(tag: str, ratio: float) -> str:
            """按 ratio 缩放 tag 里的 width，返回修改后的 tag 字符串。"""
            # 优先处理 style="...width:Xpt..."
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            if style_m:
                style_str = style_m.group(1)
                w_pt = _style_width_pt(style_str)
                if w_pt is not None:
                    new_style = _replace_style_width(style_str, w_pt * ratio)
                    tag = tag[:style_m.start()] + f'style="{new_style}"' + tag[style_m.end():]
                    return tag

            # 无 style width，处理 HTML width 属性
            attr_m = re.search(r'\bwidth="([^"]*)"', tag, re.IGNORECASE)
            if attr_m:
                w_pt = _to_pt(attr_m.group(1))
                if w_pt is not None:
                    new_val = _fmt(w_pt * ratio)
                    tag = tag[:attr_m.start()] + f'width="{new_val}"' + tag[attr_m.end():]

            return tag

        # ── 逐表格扫描并缩放 ─────────────────────────────────────────────
        # 策略：找到每个 <table...> 开标签，判断其宽度，超出则计算 ratio，
        # 然后在该 </table> 范围内缩放所有 <col>/<td>/<th> 的宽度。
        # 不使用 HTML 解析器（依赖已剥离），改用正则逐段处理。

        result_parts = []
        cursor = 0
        table_open_re  = re.compile(r'<table\b[^>]*>', re.IGNORECASE | re.DOTALL)
        table_close_re = re.compile(r'</table\s*>', re.IGNORECASE)
        cell_tag_re    = re.compile(r'<(?:col|td|th)\b[^>]*>', re.IGNORECASE | re.DOTALL)

        for tbl_open_m in table_open_re.finditer(html_content):
            # 写入 table 之前的内容
            result_parts.append(html_content[cursor:tbl_open_m.start()])

            tbl_open_tag = tbl_open_m.group(0)
            search_from  = tbl_open_m.end()

            # 找到对应的 </table>（用嵌套深度计数应对嵌套表格）
            depth      = 1
            pos        = search_from
            close_end  = len(html_content)

            while pos < len(html_content) and depth > 0:
                next_open  = table_open_re.search(html_content, pos)
                next_close = table_close_re.search(html_content, pos)

                if next_close is None:
                    break
                if next_open and next_open.start() < next_close.start():
                    depth += 1
                    pos    = next_open.end()
                else:
                    depth -= 1
                    if depth == 0:
                        close_end = next_close.end()
                    pos = next_close.end()

            table_inner = html_content[search_from:close_end - len('</table>') if close_end < len(html_content) else close_end]
            # 更精确地取 </table> 前的内容
            close_m = table_close_re.search(html_content, search_from)
            # 重新找到最外层 </table>
            depth2 = 1
            pos2   = search_from
            outer_close_m = None
            for m in re.finditer(r'<(/?)table\b[^>]*>', html_content[search_from:], re.IGNORECASE):
                if m.group(1) == '':   # 开标签
                    depth2 += 1
                else:                  # 闭标签
                    depth2 -= 1
                    if depth2 == 0:
                        abs_start = search_from + m.start()
                        abs_end   = search_from + m.end()
                        outer_close_m = (abs_start, abs_end)
                        break

            if outer_close_m:
                table_inner   = html_content[search_from : outer_close_m[0]]
                cursor_next   = outer_close_m[1]
            else:
                table_inner   = html_content[search_from:]
                cursor_next   = len(html_content)

            # 确定表格宽度
            tbl_width_pt = None
            style_m = re.search(r'style="([^"]*)"', tbl_open_tag, re.IGNORECASE)
            if style_m:
                tbl_width_pt = _style_width_pt(style_m.group(1))
            if tbl_width_pt is None:
                attr_m = re.search(r'\bwidth="([^"]*)"', tbl_open_tag, re.IGNORECASE)
                if attr_m:
                    tbl_width_pt = _to_pt(attr_m.group(1))

            if tbl_width_pt is None or tbl_width_pt <= content_width_pt:
                # 不需要缩放，原样输出
                result_parts.append(tbl_open_tag)
                result_parts.append(table_inner)
                result_parts.append('</table>')
                cursor = cursor_next
                continue

            # 需要缩放
            ratio = content_width_pt / tbl_width_pt
            print(f"   📏 表格宽度 {tbl_width_pt:.1f}pt → {content_width_pt:.1f}pt（ratio={ratio:.4f}）")

            # 缩放 table 开标签的宽度
            tbl_open_tag = _scale_tag_width(tbl_open_tag, ratio)

            # 缩放 table_inner 中所有 <col>/<td>/<th> 的宽度
            def _scale_cell(cm):
                return _scale_tag_width(cm.group(0), ratio)

            table_inner_scaled = cell_tag_re.sub(_scale_cell, table_inner)

            result_parts.append(tbl_open_tag)
            result_parts.append(table_inner_scaled)
            result_parts.append('</table>')
            cursor = cursor_next

        result_parts.append(html_content[cursor:])
        return ''.join(result_parts)

    def _fix_html_img_sizes_for_import(self, html_text: str,
                                        page_width_px: int = 794,
                                        content_width_px: int = 620) -> str:
        """
        HTML→DOCX 方向的图片尺寸修正。

        ── 问题根因 ──────────────────────────────────────────────────────────
        Spire 将 HTML 转 DOCX 时，图片最终 EMU 的计算链路为：
          1. 如果 <img> 有明确的 width/height 属性（纯数字 px）→ 直接用
          2. 如果只有 style="width:Xpx" → 能读到，但部分版本会忽略
          3. 如果没有任何尺寸信息（如本案：只有 max-width:100%; height:auto）
             → Spire 读取图片文件物理像素，假设 96 DPI 换算 EMU
             → 192 DPI 截图 → EMU 是正常的 2 倍 → 图片偏大 2 倍且超出 A4

        ── 处理策略（三路来源，优先级从高到低）──────────────────────────────
        A. img 标签已有明确 width/height（px/pt/in/cm/mm）
           → 直接换算为 px，做 A4 版心约束后写回
        B. img 标签没有明确尺寸，但 src 是 data:URI（base64 已解包为文件路径）
           → 用 PIL/struct 读取文件物理像素，做约束后写入
        C. img 标签没有明确尺寸，src 是 http/https URL
           → 下载图片（带超时、带缓存）→ 读物理像素 → 做约束后写入
           → 下载失败则跳过，不干预（Spire 自行处理，保持原有行为）

        ── A4 版心约束 ───────────────────────────────────────────────────────
        page_width_px:    A4 页面宽度，96 DPI 下约 794px（210mm）
        content_width_px: 版心宽度（扣除页边距），默认 620px（约 165mm，对应常见 2.5cm 页边距）
        约束规则：
          - 宽度超过 content_width_px 时，等比缩放使宽度 = content_width_px
          - 宽度未超出时，保持原始物理像素（不放大）

        ── style 中 max-width 的处理 ─────────────────────────────────────────
        HTML 里经常出现 style="max-width:100%; height:auto"，
        这对浏览器渲染有意义，但 Spire 解析 HTML 时往往忽略 max-width。
        修正后将 max-width/height:auto 替换为确定的 width:Xpx; height:Ypx，
        同时写入 HTML width/height 属性，双保险。
        """
        import struct
        import urllib.request
        import urllib.error

        MM_PER_INCH  = 25.4
        PX_PER_INCH  = 96.0   # HTML/CSS 标准基准 DPI

        # ── 单位换算 ──────────────────────────────────────────────────────
        def _css_val_to_px(val_str):
            if not val_str:
                return None
            s = val_str.strip().lower()
            try:
                if s.endswith('px'):  return float(s[:-2])
                if s.endswith('pt'):  return float(s[:-2]) * PX_PER_INCH / 72
                if s.endswith('in'):  return float(s[:-2]) * PX_PER_INCH
                if s.endswith('cm'):  return float(s[:-2]) / 2.54 * PX_PER_INCH
                if s.endswith('mm'):  return float(s[:-2]) / MM_PER_INCH * PX_PER_INCH
                if s.endswith('%'):   return None
                return float(s)
            except ValueError:
                return None

        # ── 从文件/bytes 读取物理像素（不依赖 PIL，纯 struct 解析主流格式）──
        def _read_image_size_from_bytes(data: bytes):
            """
            返回 (width, height) 整数像素，失败返回 (None, None)。
            支持 PNG / JPEG / GIF / BMP / WEBP。
            """
            try:
                # PNG: 8字节签名 + IHDR chunk(4长度+4类型+4w+4h)
                if data[:8] == b'\x89PNG\r\n\x1a\n':
                    w, h = struct.unpack('>II', data[16:24])
                    return w, h
                # JPEG: 扫描 SOFx marker
                if data[:2] == b'\xff\xd8':
                    i = 2
                    while i < len(data) - 8:
                        if data[i] != 0xff:
                            break
                        marker = data[i+1]
                        length = struct.unpack('>H', data[i+2:i+4])[0]
                        if marker in (0xc0, 0xc1, 0xc2, 0xc3,
                                      0xc5, 0xc6, 0xc7, 0xc9, 0xca, 0xcb):
                            h, w = struct.unpack('>HH', data[i+5:i+9])
                            return w, h
                        i += 2 + length
                # GIF
                if data[:6] in (b'GIF87a', b'GIF89a'):
                    w, h = struct.unpack('<HH', data[6:10])
                    return w, h
                # BMP
                if data[:2] == b'BM':
                    w, h = struct.unpack('<II', data[18:26])
                    return w, abs(h)
                # WEBP: RIFF....WEBP VP8 /VP8L/VP8X
                if data[:4] == b'RIFF' and data[8:12] == b'WEBP':
                    chunk = data[12:16]
                    if chunk == b'VP8 ':
                        w = struct.unpack('<H', data[26:28])[0] & 0x3fff
                        h = struct.unpack('<H', data[28:30])[0] & 0x3fff
                        return w, h
                    if chunk == b'VP8L':
                        bits = struct.unpack('<I', data[21:25])[0]
                        w = (bits & 0x3fff) + 1
                        h = ((bits >> 14) & 0x3fff) + 1
                        return w, h
                    if chunk == b'VP8X':
                        w = (struct.unpack('<I', data[24:27] + b'\x00')[0] & 0xffffff) + 1
                        h = (struct.unpack('<I', data[27:30] + b'\x00')[0] & 0xffffff) + 1
                        return w, h
            except Exception:
                pass
            return None, None

        def _read_image_size_from_file(path: str):
            try:
                with open(path, 'rb') as f:
                    data = f.read(512)   # 只需文件头
                return _read_image_size_from_bytes(data)
            except Exception:
                return None, None

        # ── HTTP 图片下载缓存（同一次调用内复用）────────────────────────
        _url_cache: dict[str, tuple] = {}   # url → (w, h) or (None, None)

        def _fetch_image_size_from_url(url: str):
            if url in _url_cache:
                return _url_cache[url]
            result = (None, None)
            for attempt, timeout in enumerate((30, 60), start=1):
                try:
                    req = urllib.request.Request(
                        url,
                        headers={'User-Agent': 'Mozilla/5.0', 'Range': 'bytes=0-4095'}
                    )
                    with urllib.request.urlopen(req, timeout=timeout) as resp:
                        data = resp.read(4096)
                    result = _read_image_size_from_bytes(data)
                    if result == (None, None):
                        # 服务器忽略 Range，完整下载
                        req2 = urllib.request.Request(
                            url, headers={'User-Agent': 'Mozilla/5.0'}
                        )
                        with urllib.request.urlopen(req2, timeout=timeout) as resp2:
                            full = resp2.read()
                        result = _read_image_size_from_bytes(full)
                    if result != (None, None):
                        break
                except Exception as e:
                    print(f"   ⚠️ 获取图片尺寸失败（第{attempt}次，timeout={timeout}s）{url[:60]}：{e}")
            _url_cache[url] = result
            return result

        # ── A4 版心等比约束 ──────────────────────────────────────────────
        def _constrain(w_raw: float, h_raw: float):
            """等比缩放使宽度不超过版心，返回 (w_px, h_px) 整数"""
            if w_raw > content_width_px:
                scale = content_width_px / w_raw
                return round(content_width_px), round(h_raw * scale)
            return round(w_raw), round(h_raw)

        # ── 写入 <img> 标签的 width/height ──────────────────────────────
        def _apply_sizes(tag: str, w_px: int, h_px: int) -> str:
            # 移除旧属性
            tag = re.sub(r'\s+width="[^"]*"',  '', tag, flags=re.IGNORECASE)
            tag = re.sub(r'\s+height="[^"]*"', '', tag, flags=re.IGNORECASE)
            # 插入新属性
            tag = re.sub(r'(<img\b)', rf'\1 width="{w_px}" height="{h_px}"', tag)

            # 更新 style：移除 max-width/height:auto/旧尺寸，写入确定值
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
            if style_m:
                s = style_m.group(1)
                s = re.sub(r'\bmax-width\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                s = re.sub(r'\bmax-height\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                s = re.sub(r'\bwidth\s*:[^;]+;?',  '', s, flags=re.IGNORECASE)
                s = re.sub(r'\bheight\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                s = s.rstrip('; ')
                new_style = f"{s}; width:{w_px}px; height:{h_px}px".lstrip('; ')
                tag = tag[:style_m.start()] + f'style="{new_style}"' + tag[style_m.end():]
            else:
                tag = re.sub(r'(<img\b)',
                             rf'\1 style="width:{w_px}px; height:{h_px}px"', tag)
            return tag

        # ── 辅助：从图片 src 获取物理像素 ─────────────────────────────
        def _get_phys_size(src: str, tag: str):
            """返回 (phys_w, phys_h) 或 (None, None)。"""
            if src.startswith('data:'):
                try:
                    b64_part = src.split(',', 1)[1] if ',' in src else ''
                    import base64 as _b64
                    raw_bytes = _b64.b64decode(b64_part[:4096] + '==')
                    return _read_image_size_from_bytes(raw_bytes)
                except Exception:
                    return None, None
            elif os.path.exists(src):
                return _read_image_size_from_file(src)
            elif src.startswith('http://') or src.startswith('https://'):
                return _fetch_image_size_from_url(src)
            return None, None

        # ── 主替换逻辑（支持 td 上下文宽度约束）──────────────────────
        def _process_img(tag: str, max_w_px: float) -> str:
            """
            处理单个 <img> 标签，max_w_px 为当前上下文的最大允许宽度（px）。
            - 表格内图片：max_w_px = td 可用宽度（td.width - padding*2）
            - 普通图片：max_w_px = content_width_px（全局版心）
            """
            src_m = re.search(r'src="([^"]+)"', tag, re.IGNORECASE)
            src   = src_m.group(1) if src_m else ''

            # 解析 style 字典（style 值内部可能有换行缩进，先压缩）
            style_dict = {}
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE | re.DOTALL)
            if style_m:
                style_val = re.sub(r'[\r\n]+\s*', ' ', style_m.group(1))
                for part in style_val.split(';'):
                    if ':' in part:
                        k, _, v = part.partition(':')
                        style_dict[k.strip().lower()] = v.strip()

            # 读取标签里的 width / height
            raw_w = _css_val_to_px(style_dict.get('width'))
            raw_h = _css_val_to_px(style_dict.get('height'))

            attr_w_m = re.search(r'\bwidth="([^"]*)"',  tag, re.IGNORECASE)
            attr_h_m = re.search(r'\bheight="([^"]*)"', tag, re.IGNORECASE)
            if raw_w is None and attr_w_m:
                raw_w = _css_val_to_px(attr_w_m.group(1))
            if raw_h is None and attr_h_m:
                raw_h = _css_val_to_px(attr_h_m.group(1))

            # height 是否为 "auto" / 缺失（需要从物理像素推算）
            h_val_str = style_dict.get('height', '').strip().lower()
            height_is_auto = (h_val_str in ('auto', '') or raw_h is None)

            # ── 路径 A：width 和 height 都明确 ─────────────────────────
            if raw_w and raw_h and raw_w > 0 and raw_h > 0:
                # 用 max_w_px（上下文宽度）约束，而不是全局 content_width_px
                if raw_w > max_w_px:
                    scale = max_w_px / raw_w
                    w_px, h_px = round(max_w_px), round(raw_h * scale)
                else:
                    w_px, h_px = round(raw_w), round(raw_h)
                print(f"   📐 [A] {round(raw_w)}×{round(raw_h)}px → {w_px}×{h_px}px（上限{round(max_w_px)}）")
                return _apply_sizes(tag, w_px, h_px)

            # ── 路径 A½：有明确 width 但 height=auto/缺失 ───────────────
            # 这是本次问题的核心场景：style="width:180pt; height:auto"
            # Spire 遇到 height:auto 会忽略 width 直接用物理像素，导致图片炸大
            # 修复：从图片物理像素推算正确的 height，然后写入明确的 height 值
            if raw_w and raw_w > 0 and height_is_auto:
                phys_w, phys_h = _get_phys_size(src, tag)
                if phys_w and phys_h and phys_w > 0:
                    # 按 raw_w 等比算出 height
                    computed_h = raw_w * phys_h / phys_w
                    # 再做上下文宽度约束
                    if raw_w > max_w_px:
                        scale = max_w_px / raw_w
                        w_px, h_px = round(max_w_px), round(computed_h * scale)
                    else:
                        w_px, h_px = round(raw_w), round(computed_h)
                    print(f"   📐 [A½] width={round(raw_w)}px + height:auto → 物理{phys_w}×{phys_h} → {w_px}×{h_px}px")
                    return _apply_sizes(tag, w_px, h_px)
                else:
                    # 无法获取物理像素：按 max_w_px 上限写入 width，清除 height:auto
                    # 用 4:3 兜底比例保证图片有合理高度（Spire 会从文件里读实际值）
                    final_w = round(min(raw_w, max_w_px))
                    # 不设置 height，让 Spire 自行决定——但必须移除 height:auto
                    tag2 = re.sub(r'\s+height="[^"]*"', '', tag, flags=re.IGNORECASE)
                    style_m2 = re.search(r'style="([^"]*)"', tag2, re.IGNORECASE)
                    if style_m2:
                        s = style_m2.group(1)
                        s = re.sub(r'\bheight\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                        s = re.sub(r'\bmax-width\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                        s = re.sub(r'\bwidth\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                        s = s.rstrip('; ')
                        s = f"{s}; width:{final_w}px".lstrip('; ')
                        tag2 = tag2[:style_m2.start()] + f'style="{s}"' + tag2[style_m2.end():]
                    tag2 = re.sub(r'\s+width="[^"]*"', '', tag2, flags=re.IGNORECASE)
                    tag2 = re.sub(r'(<img\b)', rf'\1 width="{final_w}"', tag2)
                    print(f"   ⚠️ [A½-fallback] 无法获取物理像素，width={final_w}px，height交由Spire决定")
                    return tag2

            # ── 路径 B/C：完全无尺寸信息，从物理像素推算 ───────────────
            phys_w, phys_h = _get_phys_size(src, tag)

            if not phys_w or not phys_h:
                if 'max-width' in style_dict or 'max-height' in style_dict:
                    tag = _apply_sizes(tag, round(max_w_px), round(max_w_px))
                    print(f"   ⚠️ 无法获取物理像素，回退到上下文宽度：{src[:60]}")
                return tag

            if phys_w:
                print(f"   📐 [B/C] 物理 {phys_w}×{phys_h}：{src[:60]}")

            # 用上下文宽度约束（而非全局 content_width_px）
            if float(phys_w) > max_w_px:
                scale = max_w_px / phys_w
                w_px, h_px = round(max_w_px), round(phys_h * scale)
            else:
                w_px, h_px = round(phys_w), round(phys_h)
            return _apply_sizes(tag, w_px, h_px)

        # ── 上下文感知替换：区分表格内/表格外图片 ───────────────────────
        # 表格内图片的最大宽度受 td 宽度（减去 padding）限制，
        # 而不是全局版心宽度，否则图片会溢出单元格挤在一起。
        PT_TO_PX = PX_PER_INCH / 72

        def _td_max_w_px(td_tag: str) -> float:
            """从 <td> 标签提取可用宽度（px），失败返回 content_width_px。"""
            style_dict_td = {}
            sm = re.search(r'style="([^"]*)"', td_tag, re.IGNORECASE)
            if sm:
                for part in sm.group(1).split(';'):
                    if ':' in part:
                        k, _, v = part.partition(':')
                        style_dict_td[k.strip().lower()] = v.strip()

            td_w_px = _css_val_to_px(style_dict_td.get('width'))
            if td_w_px is None:
                attr_m = re.search(r'\bwidth="([^"]*)"', td_tag, re.IGNORECASE)
                if attr_m:
                    td_w_px = _css_val_to_px(attr_m.group(1))

            if not td_w_px:
                return float(content_width_px)

            # 减去 padding（CSS padding 属性，支持单值/四值）
            padding_str = style_dict_td.get('padding', '').strip()
            padding_px  = 0.0
            if padding_str:
                parts = padding_str.split()
                # 取左右 padding（padding: top right bottom left 或 padding: all）
                if len(parts) == 1:
                    p = _css_val_to_px(parts[0]) or 0
                    padding_px = p * 2   # 左 + 右
                elif len(parts) == 2:
                    p = _css_val_to_px(parts[1]) or 0
                    padding_px = p * 2
                elif len(parts) >= 4:
                    pl = _css_val_to_px(parts[3]) or 0
                    pr = _css_val_to_px(parts[1]) or 0
                    padding_px = pl + pr
                elif len(parts) == 3:
                    p = _css_val_to_px(parts[1]) or 0
                    padding_px = p * 2

            avail = td_w_px - padding_px
            return max(avail, 40.0)   # 至少 40px，防止除零

        def _replace_img_global(m):
            """全局替换回调，对表格外图片使用全局版心宽度。"""
            return _process_img(m.group(0), float(content_width_px))

        # 先处理表格外图片，再逐表格处理表格内图片
        # 策略：把 HTML 按 <table>...</table> 块拆分，
        #        表格外段落用全局宽度，表格内按 td 逐列处理。
        IMG_RE    = re.compile(r'<img\b[^>]*>', re.IGNORECASE | re.DOTALL)
        TABLE_RE  = re.compile(r'<table\b[^>]*>.*?</table\s*>', re.IGNORECASE | re.DOTALL)
        TD_RE     = re.compile(r'(<t[dh]\b[^>]*>)(.*?)(?=<t[dh]\b|</tr|</table)', re.IGNORECASE | re.DOTALL)

        result_parts = []
        cursor = 0

        for tbl_m in TABLE_RE.finditer(html_text):
            # 表格外段落：用全局版心宽度
            before = html_text[cursor:tbl_m.start()]
            result_parts.append(IMG_RE.sub(_replace_img_global, before))

            # 表格内：逐 td 处理，每个 td 用自己的宽度约束
            tbl_html   = tbl_m.group(0)
            tbl_result = []
            td_cursor  = 0

            for td_m in TD_RE.finditer(tbl_html):
                # td 开标签前的内容（表格结构标签等）
                tbl_result.append(tbl_html[td_cursor:td_m.start()])
                td_open    = td_m.group(1)
                td_content = td_m.group(2)
                max_w      = _td_max_w_px(td_open)

                def _make_td_replacer(mw):
                    def _r(im):
                        return _process_img(im.group(0), mw)
                    return _r

                td_content_fixed = IMG_RE.sub(_make_td_replacer(max_w), td_content)
                tbl_result.append(td_open)
                tbl_result.append(td_content_fixed)
                td_cursor = td_m.end()

            tbl_result.append(tbl_html[td_cursor:])
            result_parts.append(''.join(tbl_result))
            cursor = tbl_m.end()

        # 最后一段表格外内容
        result_parts.append(IMG_RE.sub(_replace_img_global, html_text[cursor:]))
        return ''.join(result_parts)

    def _embed_images_to_html(self, html_path, image_display_order, original_img_dir):
        """
        【内部方法】对已生成的HTML文件做图片base64内嵌（in-place）
        在流式合并完成后统一调用，避免chunk阶段内存溢出

        修复问题1：改用 _build_spire_to_original_map 做名称匹配，不再纯靠位置索引。
        修复：替换正则使用完整文件名精确匹配（src=" ... "），防止子串误匹配。
        """
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        html_content = html_content.replace('\\', '/')

        img_pattern = re.compile(r'<img[^>]*src="([^"]+)"[^>]*>')
        spire_img_names = []
        for match in img_pattern.finditer(html_content):
            name = os.path.basename(self._normalize_path(match.group(1)))
            if name not in spire_img_names:
                spire_img_names.append(name)

        print(f"=== 待内嵌图片：{spire_img_names} ===")

        # 【修复问题1】用名称匹配构建映射，不依赖纯位置索引
        fallback_dir = os.path.dirname(html_path)
        name_map = self._build_spire_to_original_map(
            spire_img_names, image_display_order, original_img_dir, fallback_dir
        )

        for spire_name, img_path in name_map.items():
            base64_str = self._image_to_base64(img_path)
            if not base64_str:
                continue

            html_content = re.compile(
                r'src="[^"]*/?' + re.escape(spire_name) + r'"'
            ).sub(f'src="{base64_str}"', html_content)
            print(f"🔄 {spire_name} → {os.path.basename(img_path)} 已转为Base64")

        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print("✅ 图片内嵌完成")

    # ------------------------------------------------------------------ #
    #  分片工具                                                             #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _pydocx_collect_body_elements(docx_path):
        """
        用 python-docx 读取文档，返回 body 下所有顶层元素列表。
        每个元素是 lxml Element，tag 为 w:p 或 w:tbl。
        """
        doc = PythonDocx(docx_path)
        body = doc.element.body
        elements = [
            el for el in body
            if el.tag in (qn('w:p'), qn('w:tbl'))
        ]
        print(f"[python-docx 诊断] body 顶层元素总数: {len(elements)}")
        para_cnt  = sum(1 for e in elements if e.tag == qn('w:p'))
        table_cnt = sum(1 for e in elements if e.tag == qn('w:tbl'))
        print(f"[python-docx 诊断] 段落: {para_cnt}, 表格: {table_cnt}")
        return doc, elements

    @staticmethod
    def _pydocx_count_table_paras(tbl_el):
        """统计 w:tbl 元素内所有 w:p 的数量"""
        return len(tbl_el.findall('.//' + qn('w:p')))

    @staticmethod
    def _pydocx_copy_sectPr(src_doc, dst_doc):
        """
        将源文档的页面设置（w:sectPr）复制到目标文档，
        同时彻底移除页眉页脚相关内容，避免 spire 在 chunk HTML 里渲染页眉页脚。
        """
        HF_TAGS = {
            qn('w:headerReference'),
            qn('w:footerReference'),
            qn('w:titlePg'),
        }

        src_body = src_doc.element.body
        src_sectPr = src_body.find(qn('w:sectPr'))
        if src_sectPr is None:
            last_p = src_body.findall(qn('w:p'))
            if last_p:
                pPr = last_p[-1].find(qn('w:pPr'))
                if pPr is not None:
                    src_sectPr = pPr.find(qn('w:sectPr'))

        if src_sectPr is not None:
            sectPr_copy = copy.deepcopy(src_sectPr)

            for tag in HF_TAGS:
                for node in sectPr_copy.findall(tag):
                    sectPr_copy.remove(node)

            pgMar = sectPr_copy.find(qn('w:pgMar'))
            if pgMar is not None:
                pgMar.attrib.pop(qn('w:header'), None)
                pgMar.attrib.pop(qn('w:footer'), None)

            dst_body = dst_doc.element.body
            old = dst_body.find(qn('w:sectPr'))
            if old is not None:
                dst_body.remove(old)
            dst_body.append(sectPr_copy)

    @staticmethod
    def _pydocx_new_doc(src_doc):
        """
        创建一个继承源文档样式、编号和页面设置的空白 python-docx Document。
        """
        dst = PythonDocx()
        dst.element.body.clear()

        try:
            src_styles_xml = src_doc.part.styles._element
            dst.part.styles._element.clear()
            for child in src_styles_xml:
                dst.part.styles._element.append(copy.deepcopy(child))
        except Exception as e:
            print(f"⚠️ 样式复制失败（使用默认样式）：{e}")

        try:
            src_numbering_part = src_doc.part.numbering_part
            if src_numbering_part is not None:
                rel_type = (
                    'http://schemas.openxmlformats.org/officeDocument/2006/'
                    'relationships/numbering'
                )
                dst_part = dst.part
                existing_num_part = None
                for rel in dst_part.rels.values():
                    if rel.reltype == rel_type:
                        existing_num_part = rel.target_part
                        break
                if existing_num_part is not None:
                    existing_num_part._element.clear()
                    for child in src_numbering_part._element:
                        existing_num_part._element.append(copy.deepcopy(child))
        except Exception as e:
            print(f"⚠️ 编号定义复制失败（忽略）：{e}")

        DocxHtmlConverter._pydocx_copy_sectPr(src_doc, dst)
        return dst

    @staticmethod
    def _inject_resources_into_chunk(src_docx_path, chunk_docx_path,
                                     referenced_images=None):
        """
        通过 zipfile 将源文档的资源精确注入已保存的 chunk docx。

        【修复问题4】OLE预览图判断条件扩展：
        原条件 'ole' in entry or 'vml' in entry 会漏掉部分 OLE 预览图。
        改为同时检查 Type 属性是否包含 oleObject/vmlDrawing/vmldrawing，
        以及 Target 路径是否包含 vml/ole 等关键字，覆盖更全面。
        同时对 v:imagedata 引用的图片（即 vmlDrawing 内的图片 rId）无条件放行，
        避免 OLE 对象预览图因 referenced_images 过滤而丢失。
        """
        RELS_PATH = 'word/_rels/document.xml.rels'

        with zipfile.ZipFile(chunk_docx_path, 'r') as zf:
            chunk_doc_xml  = zf.read('word/document.xml').decode('utf-8')
            chunk_rels_xml = zf.read(RELS_PATH).decode('utf-8') if RELS_PATH in zf.namelist() else ''
            chunk_files    = set(zf.namelist())

        doc_rids = set(
            re.findall(r'r:(?:id|embed|link)="(rId\d+)"', chunk_doc_xml)
        ) | set(
            re.findall(r'relationships}(?:id|embed|link)="(rId\d+)"', chunk_doc_xml)
        )
        existing_rids = set(re.findall(r'Id="(rId\d+)"', chunk_rels_xml))
        missing_rids  = doc_rids - existing_rids

        if not missing_rids:
            return

        with zipfile.ZipFile(src_docx_path, 'r') as zf:
            src_rels_xml = zf.read(RELS_PATH).decode('utf-8') if RELS_PATH in zf.namelist() else ''
            src_namelist = zf.namelist()

            # 额外扫描 vmlDrawing 文件里引用的图片 rId，这些属于 OLE 预览图
            vml_image_rids = set()
            for fname in src_namelist:
                if re.match(r'word/vmlDrawing\d*\.vml', fname):
                    try:
                        vml_content = zf.read(fname).decode('utf-8', errors='replace')
                        vml_image_rids.update(
                            re.findall(r'r:(?:id|href)="(rId\d+)"', vml_content)
                        )
                    except Exception:
                        pass

            src_rels_map = {}
            for m in re.finditer(r'<Relationship[^>]*/>', src_rels_xml):
                entry  = m.group(0)
                rid_m  = re.search(r'Id="(rId\d+)"', entry)
                tgt_m  = re.search(r'Target="([^"]+)"', entry)
                if rid_m and tgt_m:
                    rid    = rid_m.group(1)
                    target = tgt_m.group(1)
                    zip_path = 'word/' + target if not target.startswith('/') else target.lstrip('/')
                    src_rels_map[rid] = (entry, zip_path)

            entries_to_add = []
            files_to_copy  = {}

            HF_REL_TYPES = ('header', 'footer')

            for rid in sorted(missing_rids):
                if rid not in src_rels_map:
                    print(f"   ⚠️ rId={rid} 在源文档 rels 中也找不到，跳过")
                    continue

                entry, zip_path = src_rels_map[rid]

                if any(hf in entry.lower() for hf in HF_REL_TYPES):
                    continue

                # 【修复问题4】OLE预览图判断：扩展检查范围
                if zip_path.startswith('word/media/'):
                    fname = os.path.basename(zip_path)
                    entry_lower  = entry.lower()
                    target_lower = zip_path.lower()
                    is_ole_or_vml = (
                        'oleobject'   in entry_lower or
                        'vmldrawing'  in entry_lower or
                        'ole'         in target_lower or
                        'vml'         in target_lower or
                        rid in vml_image_rids          # vmlDrawing 文件中直接引用的 rId
                    )
                    if (referenced_images is not None
                            and fname not in referenced_images
                            and not is_ole_or_vml):
                        continue

                entries_to_add.append(entry)

                if zip_path in src_namelist and zip_path not in chunk_files:
                    files_to_copy[zip_path] = zf.read(zip_path)

        if not entries_to_add and not files_to_copy:
            return

        tmp_path = chunk_docx_path + '.tmp'
        with zipfile.ZipFile(chunk_docx_path, 'r') as src_zip, \
             zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as dst_zip:

            for item in src_zip.infolist():
                data = src_zip.read(item.filename)

                if item.filename == RELS_PATH and entries_to_add:
                    updated = data.decode('utf-8').replace(
                        '</Relationships>',
                        '\n'.join(entries_to_add) + '\n</Relationships>'
                    )
                    dst_zip.writestr(item, updated.encode('utf-8'))

                elif item.filename == 'word/settings.xml':
                    settings = data.decode('utf-8')
                    settings = re.sub(r'<w:evenAndOddHeaders\s*/>', '', settings)
                    dst_zip.writestr(item, settings.encode('utf-8'))

                else:
                    dst_zip.writestr(item, data)

            for zip_path, file_data in files_to_copy.items():
                dst_zip.writestr(zip_path, file_data)

        os.replace(tmp_path, chunk_docx_path)
        print(f"   💉 注入 {len(entries_to_add)} 条 rels + {len(files_to_copy)} 个资源文件")

    @staticmethod
    def _pydocx_append_element(dst_doc, el, src_doc=None):
        """将元素深拷贝追加到目标文档 body（sectPr 前）。"""
        dst_body = dst_doc.element.body
        sect_pr  = dst_body.find(qn('w:sectPr'))
        cloned   = copy.deepcopy(el)
        if sect_pr is not None:
            sect_pr.addprevious(cloned)
        else:
            dst_body.append(cloned)

    @staticmethod
    def _pydocx_make_marker_para(group_id):
        """
        生成一个携带 TABLE_SPLIT_MARKER 的段落 lxml Element。
        字体 1pt、白色，尽量不影响视觉排版。
        """
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'),  '0')
        pPr.append(spacing)
        p.append(pPr)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '2')
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '2')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FFFFFF')
        rPr.extend([sz, szCs, color])
        r.append(rPr)

        t = OxmlElement('w:t')
        t.text = f'TABLE_SPLIT_MARKER::{group_id}'
        r.append(t)
        p.append(r)
        return p

    def _needs_chunking(self, docx_path):
        """
        【内部方法】用 python-docx 统计段落/表格数，判断是否需要分片。
        """
        _, elements = self._pydocx_collect_body_elements(docx_path)
        total_paras  = sum(1 for e in elements if e.tag == qn('w:p'))
        total_tables = sum(1 for e in elements if e.tag == qn('w:tbl'))
        needs = total_paras > self.MAX_PARAGRAPHS or total_tables > self.MAX_TABLES
        print(f"📊 文档规模：{total_paras} 段落，{total_tables} 表格，{'需要' if needs else '无需'}分片")
        return needs

    @staticmethod
    def _pydocx_sanitize_element(el, src_rels_rids):
        """
        处理元素内的 w:object（OLE 嵌入对象）：
        - 保留 v:shape（含预览图 v:imagedata），视觉上图片仍显示
        - 移除 o:OLEObject 本体节点（spire 处理此节点时崩溃）
        - 全面同步尺寸与位置，确保移除 OLEObject 后渲染结果与原文档一致

        ── OLE 对象典型 XML 结构 ──────────────────────────────────────────
            <w:pPr>
                <w:framePr w:w="..." w:h="..." w:x="..." w:y="..."
                           w:hAnchor="..." w:vAnchor="..." w:wrap="..."/>
                           <!-- 浮动对象时存在；内联时无此节点 -->
            </w:pPr>
            <w:object w:dxaOrig="2160" w:dyaOrig="1440">
                <!-- w:dxaOrig/w:dyaOrig：原始宽高，单位 twip（1/20 pt） -->
                <v:shape id="..." style="position:absolute;
                                         width:113.25pt;height:75.75pt;
                                         margin-left:9pt;margin-top:3.75pt;
                                         mso-position-horizontal-relative:text;
                                         mso-position-vertical-relative:line"
                         coordsize="21600,21600"
                         o:ole="" filled="f" stroked="f">
                    <v:fill o:detectmouseclick="t"/>
                    <v:path v:extrusionok="f" gradientshapeok="t" o:connecttype="rect"/>
                    <v:imagedata r:id="rIdN" o:title=""/>
                    <!-- r:id → 预览图资源 rId，必须保留 -->
                </v:shape>
                <o:OLEObject Type="Embed" ProgID="..." r:id="rIdM" .../>
                <!-- r:id → OLE 对象本体 rId，移除此节点 -->
            </w:object>

        ── 尺寸信息来源及优先级 ────────────────────────────────────────────
        尺寸分布在三处，存在冗余，处理时以"最可靠"者为基准对其余两处做对齐：

        A. v:shape style width/height（最可靠，Spire/Word 渲染直接依赖）
           → 首选基准。单位通常为 pt，也可能为 in（需换算）。

        B. w:object w:dxaOrig/w:dyaOrig（原始创建尺寸，单位 twip = 1/20 pt）
           → A 缺失时作为备用基准；A 存在时反向同步回此属性保持一致。

        C. v:shape coordsize（VML 内部坐标系大小，默认 "21600,21600"）
           → 通常固定值，不代表实际像素尺寸，不修改。

        ── 位置信息来源 ────────────────────────────────────────────────────
        D. v:shape style margin-left/margin-top（内联偏移 / 浮动偏移）
           → 原样保留，不修改。

        E. v:shape style position/mso-position-* 属性（绝对/相对定位模式）
           → 原样保留，不修改。

        F. w:pPr/w:framePr（浮动框架定位，w:x/w:y/w:w/w:h/w:hAnchor/w:vAnchor）
           → 与 A/B 的尺寸对齐（w:w/w:h 同步），锚点/环绕属性不修改。

        ── 同步逻辑 ─────────────────────────────────────────────────────────
        1. 解析 v:shape style，提取 width/height（换算为 twip 和 pt 两种格式）
        2. 若 v:shape style 无 width/height → 从 w:dxaOrig/w:dyaOrig 换算补入
        3. 将最终确定的尺寸（twip）回写到 w:object 的 w:dxaOrig/w:dyaOrig
        4. 若段落有 w:framePr，将 w:w/w:h 与最终尺寸对齐
        5. 移除 o:OLEObject 本体
        """
        O_NS     = 'urn:schemas-microsoft-com:office:office'
        W_NS     = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        VML_NS   = 'urn:schemas-microsoft-com:vml'
        R_NS     = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        RID_ATTR = f'{{{R_NS}}}id'

        # ── 单位换算工具 ─────────────────────────────────────────────────
        def _css_dim_to_twip(val_str):
            """
            将 CSS 尺寸字符串（pt / in / cm / mm / px）换算为 twip（整数）。
            twip = 1/20 pt；1in = 72pt；1cm ≈ 28.35pt；1mm ≈ 2.835pt；1px ≈ 0.75pt
            返回整数 twip，无法解析时返回 None。
            """
            if not val_str:
                return None
            val_str = val_str.strip().lower()
            try:
                if val_str.endswith('pt'):
                    return round(float(val_str[:-2]) * 20)
                if val_str.endswith('in'):
                    return round(float(val_str[:-2]) * 72 * 20)
                if val_str.endswith('cm'):
                    return round(float(val_str[:-2]) * 28.3465 * 20)
                if val_str.endswith('mm'):
                    return round(float(val_str[:-2]) * 2.83465 * 20)
                if val_str.endswith('px'):
                    return round(float(val_str[:-2]) * 0.75 * 20)
            except (ValueError, AttributeError):
                pass
            return None

        def _twip_to_pt_str(twip):
            """twip（整数）→ 'X.Xpt' 字符串"""
            return f"{twip / 20:.1f}pt"

        def _parse_style(style_str):
            """解析 CSS style 字符串 → (有序key列表, {key: value}) """
            order, d = [], {}
            for part in style_str.split(';'):
                part = part.strip()
                if not part:
                    continue
                if ':' in part:
                    k, _, v = part.partition(':')
                    k, v = k.strip(), v.strip()
                    d[k] = v
                    order.append(k)
                else:
                    d[part] = None
                    order.append(part)
            return order, d

        def _serialize_style(order, d):
            """(有序key列表, dict) → CSS style 字符串"""
            parts = []
            for k in order:
                parts.append(f"{k}:{d[k]}" if d[k] is not None else k)
            return ';'.join(parts)

        cloned = copy.deepcopy(el)

        for obj in cloned.findall('.//' + qn('w:object')):
            ole_node = obj.find(f'{{{O_NS}}}OLEObject')
            if ole_node is None:
                continue

            ole_rid = ole_node.get(RID_ATTR)

            # ── A. 读取 v:shape style 中的 width/height ──────────────────
            v_shape = obj.find(f'{{{VML_NS}}}shape')
            shape_w_twip = None   # 最终确定的宽（twip）
            shape_h_twip = None   # 最终确定的高（twip）
            style_order, style_dict = [], {}

            if v_shape is not None:
                style_order, style_dict = _parse_style(v_shape.get('style', ''))
                shape_w_twip = _css_dim_to_twip(style_dict.get('width'))
                shape_h_twip = _css_dim_to_twip(style_dict.get('height'))

            # ── B. 读取 w:dxaOrig / w:dyaOrig 作为备用 ──────────────────
            dxa_attr = f'{{{W_NS}}}dxaOrig'
            dya_attr = f'{{{W_NS}}}dyaOrig'
            dxa_orig = obj.get(dxa_attr)
            dya_orig = obj.get(dya_attr)

            try:
                dxa_twip = int(dxa_orig) if dxa_orig else None
            except ValueError:
                dxa_twip = None
            try:
                dya_twip = int(dya_orig) if dya_orig else None
            except ValueError:
                dya_twip = None

            # ── 确定最终尺寸（A优先，A缺失用B，B也缺失则不改动）────────
            if shape_w_twip is None:
                shape_w_twip = dxa_twip          # 用 B 补 A
            if shape_h_twip is None:
                shape_h_twip = dya_twip

            log_parts = []

            # ── 步骤1：v:shape style width/height 补全/修正 ──────────────
            if v_shape is not None and (shape_w_twip or shape_h_twip):
                modified = False

                if shape_w_twip and (not style_dict.get('width') or
                                     style_dict['width'] in ('0', '0pt', 'auto')):
                    style_dict['width'] = _twip_to_pt_str(shape_w_twip)
                    if 'width' not in style_order:
                        style_order.insert(0, 'width')
                    modified = True

                if shape_h_twip and (not style_dict.get('height') or
                                     style_dict['height'] in ('0', '0pt', 'auto')):
                    style_dict['height'] = _twip_to_pt_str(shape_h_twip)
                    if 'height' not in style_order:
                        idx = (style_order.index('width') + 1
                               if 'width' in style_order else 0)
                        style_order.insert(idx, 'height')
                    modified = True

                if modified:
                    v_shape.set('style', _serialize_style(style_order, style_dict))
                    log_parts.append(
                        f"style补全 {style_dict.get('width')}×{style_dict.get('height')}"
                    )

            # ── 步骤2：回写 w:dxaOrig/w:dyaOrig（与 v:shape style 对齐）──
            # 以 v:shape style 为权威来源，将最终尺寸同步回 w:object 属性，
            # 防止 Spire 读取 w:dxaOrig/w:dyaOrig 时用旧值覆盖渲染尺寸。
            if shape_w_twip and str(shape_w_twip) != (dxa_orig or ''):
                obj.set(dxa_attr, str(shape_w_twip))
                log_parts.append(f"dxaOrig→{shape_w_twip}")
            if shape_h_twip and str(shape_h_twip) != (dya_orig or ''):
                obj.set(dya_attr, str(shape_h_twip))
                log_parts.append(f"dyaOrig→{shape_h_twip}")

            # ── 步骤3：同步 w:framePr（浮动框架宽高）────────────────────
            # w:framePr 存在于 w:pPr 下，w:w/w:h 记录浮动框的显示尺寸（twip）。
            # 若与最终确定的尺寸不一致，同步修正，避免浮动框尺寸与图片不符。
            para_el = obj.getparent()
            while para_el is not None and para_el.tag != f'{{{W_NS}}}p':
                para_el = para_el.getparent()

            if para_el is not None:
                pPr = para_el.find(f'{{{W_NS}}}pPr')
                if pPr is not None:
                    frame_pr = pPr.find(f'{{{W_NS}}}framePr')
                    if frame_pr is not None:
                        fw_attr = f'{{{W_NS}}}w'
                        fh_attr = f'{{{W_NS}}}h'

                        if shape_w_twip:
                            old_fw = frame_pr.get(fw_attr)
                            if old_fw != str(shape_w_twip):
                                frame_pr.set(fw_attr, str(shape_w_twip))
                                log_parts.append(f"framePr.w {old_fw}→{shape_w_twip}")

                        if shape_h_twip:
                            old_fh = frame_pr.get(fh_attr)
                            if old_fh != str(shape_h_twip):
                                frame_pr.set(fh_attr, str(shape_h_twip))
                                log_parts.append(f"framePr.h {old_fh}→{shape_h_twip}")

            # ── 步骤4：移除 OLEObject 本体 ───────────────────────────────
            obj.remove(ole_node)

            summary = f"（{', '.join(log_parts)}）" if log_parts else "（尺寸无变化）"
            print(f"   ✂️ 移除 OLEObject 本体（r:id={ole_rid}）{summary}")

        return cloned

    def _split_docx_to_chunks(self, docx_path, chunk_dir, image_display_order=None):
        """
        【内部方法】用 python-docx + lxml 将大文档拆分为多个子 DOCX。

        【修复问题3】大表格拆分完毕后，后续普通段落的计数从正确起点开始。
        原bug：拆分大表格后 i += 1 跳过了表格本身，但此后立即进入
               下一次循环取 elements[i]（此时 i 已指向表格后第一个元素），
               如果该元素是段落且 para_count 刚重置为0，逻辑上没问题，
               但如果大表格后紧跟着另一张大表格，chunk 会多保存一个空 chunk。
               修复：大表格拆分后不额外 i+=1（因为循环体末尾没有统一的 i+=1），
               改为在 while 顶部用明确的 continue 控制流跳过已处理位置。
        """
        os.makedirs(chunk_dir, exist_ok=True)

        src_doc, elements = self._pydocx_collect_body_elements(docx_path)
        total_elements = len(elements)

        src_rels_rids = set()
        try:
            with zipfile.ZipFile(docx_path, 'r') as zf:
                if 'word/_rels/document.xml.rels' in zf.namelist():
                    rels_xml = zf.read('word/_rels/document.xml.rels').decode('utf-8')
                    for m in re.finditer(r'<Relationship[^>]*/>', rels_xml):
                        entry = m.group(0)
                        if 'oleObject' in entry:
                            rid_m = re.search(r'Id="(rId\d+)"', entry)
                            if rid_m:
                                src_rels_rids.add(rid_m.group(1))
        except Exception as e:
            print(f"⚠️ 读取源文档 rels 失败：{e}")

        chunk_paths = []
        chunk_idx   = 0
        para_count  = 0
        table_count = 0
        dst_doc     = self._pydocx_new_doc(src_doc)

        def _save_chunk(doc, idx, p_cnt, t_cnt):
            path = self._normalize_path(
                os.path.join(chunk_dir, f"chunk_{idx:04d}.docx")
            )
            doc.save(path)
            ref_imgs = set(image_display_order) if image_display_order else None
            self._inject_resources_into_chunk(docx_path, path, ref_imgs)
            print(f"✅ 切片 chunk_{idx:04d}：{p_cnt} 段落，{t_cnt} 表格")
            return path

        def _flush_current_chunk():
            """保存当前非空 chunk，重置计数器，返回新 doc。"""
            nonlocal chunk_idx, para_count, table_count, dst_doc
            if para_count > 0 or table_count > 0:
                chunk_paths.append(_save_chunk(dst_doc, chunk_idx, para_count, table_count))
                chunk_idx  += 1
            dst_doc     = self._pydocx_new_doc(src_doc)
            para_count  = 0
            table_count = 0

        i = 0
        while i < total_elements:
            el = elements[i]

            # ── 表格 ──────────────────────────────────────────────────
            if el.tag == qn('w:tbl'):
                inner_paras = self._pydocx_count_table_paras(el)
                rows        = el.findall(qn('w:tr'))
                total_rows  = len(rows)

                if (para_count + inner_paras <= self.MAX_PARAGRAPHS and
                        table_count + 1 <= self.MAX_TABLES):
                    # 整张表格放入当前 chunk
                    clean_tbl = self._pydocx_sanitize_element(el, src_rels_rids)
                    self._pydocx_append_element(dst_doc, clean_tbl)
                    para_count  += inner_paras
                    table_count += 1
                    i += 1
                else:
                    # 先保存当前非空 chunk
                    # 【修复问题3】统一用 _flush_current_chunk，避免重置逻辑遗漏
                    _flush_current_chunk()

                    # 按行拆分大表格
                    split_group_id = uuid.uuid4().hex[:8]
                    row_cursor     = 0

                    tbl_pr   = el.find(qn('w:tblPr'))
                    tbl_grid = el.find(qn('w:tblGrid'))

                    while row_cursor < total_rows:
                        # 【修复问题3】每个表格分片独立创建新 doc，
                        # 不复用外层 dst_doc（外层 doc 已在 _flush_current_chunk 重置）
                        split_doc = self._pydocx_new_doc(src_doc)

                        split_doc.element.body.insert(
                            0, self._pydocx_make_marker_para(split_group_id)
                        )

                        sub_tbl = OxmlElement('w:tbl')
                        if tbl_pr is not None:
                            sub_tbl.append(copy.deepcopy(tbl_pr))
                        if tbl_grid is not None:
                            sub_tbl.append(copy.deepcopy(tbl_grid))

                        rows_in_chunk     = 0
                        chunk_inner_paras = 0

                        while row_cursor < total_rows:
                            row = rows[row_cursor]
                            row_paras = len(row.findall('.//' + qn('w:p')))
                            if (chunk_inner_paras + row_paras > self.MAX_PARAGRAPHS
                                    and rows_in_chunk > 0):
                                break
                            clean_row = self._pydocx_sanitize_element(row, src_rels_rids)
                            sub_tbl.append(clean_row)
                            chunk_inner_paras += row_paras
                            rows_in_chunk     += 1
                            row_cursor        += 1

                        self._pydocx_append_element(split_doc, sub_tbl, src_doc)

                        split_path = self._normalize_path(
                            os.path.join(chunk_dir, f"chunk_{chunk_idx:04d}.docx")
                        )
                        split_doc.save(split_path)
                        ref_imgs = set(image_display_order) if image_display_order else None
                        self._inject_resources_into_chunk(docx_path, split_path, ref_imgs)
                        print(f"✅ 切片 chunk_{chunk_idx:04d}：（表格分片 {split_group_id}，{rows_in_chunk} 行）")
                        chunk_paths.append(split_path)
                        chunk_idx += 1

                    # 【修复问题3】大表格拆分完毕后，外层 dst_doc/para_count/table_count
                    # 已由 _flush_current_chunk 重置，此处只需推进 i，不再重复重置
                    i += 1

            # ── 普通段落 ──────────────────────────────────────────────
            else:
                if para_count + 1 > self.MAX_PARAGRAPHS and (para_count > 0 or table_count > 0):
                    _flush_current_chunk()

                clean_el = self._pydocx_sanitize_element(el, src_rels_rids)
                self._pydocx_append_element(dst_doc, clean_el)
                para_count += 1
                i += 1

        # 保存最后一个非空 chunk
        if para_count > 0 or table_count > 0:
            chunk_paths.append(_save_chunk(dst_doc, chunk_idx, para_count, table_count))

        print(f"📦 共切分为 {len(chunk_paths)} 个子文档")
        return chunk_paths

    # ------------------------------------------------------------------ #
    #  HTML 合并工具                                                        #
    # ------------------------------------------------------------------ #

    def _clean_header_footer(self, html_content):
        """
        【修复问题6】去除页眉页脚相关元素。

        原正则用 .*?（非贪婪）加 re.DOTALL，会把嵌套的 </div> 之前的所有内容
        全部吃掉，导致正文内容被误删。

        修复：改用基于嵌套深度计数的方式精确提取匹配的 </div>，
        而不是依赖正则来匹配嵌套结构。
        同时统一匹配 Spire 实际使用的 -spr-headerfooter-type 属性。
        """
        # Spire 生成的页眉页脚 div 统一用这个属性标记，优先按此匹配
        spire_hf_pattern = re.compile(
            r'<div[^>]*-spr-headerfooter-type[^>]*>',
            re.IGNORECASE
        )
        # 通用 class/id 包含 header/footer 的 div（兜底）
        generic_hf_pattern = re.compile(
            r'<div[^>]*(?:class|id)\s*=\s*["\'][^"\']*(?:header|footer)[^"\']*["\'][^>]*>',
            re.IGNORECASE
        )

        def _remove_div_block(content, pattern):
            """找到 pattern 匹配的开标签，然后按嵌套深度找到对应的 </div> 删除整块。"""
            result = []
            pos = 0
            for m in pattern.finditer(content):
                result.append(content[pos:m.start()])
                # 从开标签结束位置开始，用深度计数找到匹配的 </div>
                depth  = 1
                cursor = m.end()
                while cursor < len(content) and depth > 0:
                    open_m  = re.search(r'<div[^>]*>', content[cursor:], re.IGNORECASE)
                    close_m = re.search(r'</div\s*>', content[cursor:], re.IGNORECASE)
                    if close_m is None:
                        break
                    if open_m and open_m.start() < close_m.start():
                        depth  += 1
                        cursor += open_m.end()
                    else:
                        depth  -= 1
                        cursor += close_m.end()
                pos = cursor  # 跳过整个 div 块（包含内容和闭合标签）
            result.append(content[pos:])
            return ''.join(result)

        html_content = _remove_div_block(html_content, spire_hf_pattern)
        html_content = _remove_div_block(html_content, generic_hf_pattern)
        return html_content

    def _docx_to_html_no_embed(self, docx_path, html_path):
        """
        【内部方法】DOCX转HTML，图片保留为文件引用，不做base64内嵌
        专供分片流程使用，避免chunk阶段大文件内存问题
        """
        docx_path = self._normalize_path(docx_path)
        html_path = self._normalize_path(html_path)
        html_dir = os.path.dirname(html_path)

        spire_img_dir = self._normalize_path(
            os.path.join(html_dir, f"img_{uuid.uuid4().hex[:8]}")
        )

        document = Document()
        try:
            document.LoadFromFile(docx_path)
            document.HtmlExportOptions.ImageEmbedded = False
            document.HtmlExportOptions.ImagesPath = spire_img_dir
            document.HtmlExportOptions.ImageFormat = self.default_image_format
            document.SaveToFile(html_path, FileFormat.Html)
        except Exception as e:
            print(f"❌ Spire转换HTML失败：{e}")
            return False
        finally:
            document.Close()
            del document

        css_file_path = self._normalize_path(os.path.splitext(html_path)[0] + '_styles.css')
        if os.path.exists(css_file_path):
            with open(css_file_path, 'r', encoding='utf-8') as f:
                css_content = f.read()
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            html_content = re.sub(r'<link[^>]*href="[^"]+\.css"[^>]*>', '', html_content)
            html_content = html_content.replace(
                '</head>',
                f'<style type="text/css">\n{css_content}\n</style>\n</head>'
            )
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            os.remove(css_file_path)

        print(f"✅ chunk转换完成（图片未内嵌）：{os.path.basename(html_path)}")
        return True

    def _merge_html_files_to_disk(self, chunk_html_paths, output_path):
        """
        【内部方法】流式合并多个chunk HTML为一个完整HTML文件

        【修复问题2】游标推进逻辑重构：
        原代码用 body[marker_m.end():].index(after_table) 做子串定位，
        当 after_table 内容在 body 中多次出现时会定位到错误位置。
        修复：直接在正则匹配结果上用 re.search 从 marker_m.end() 开始搜索，
        取得表格的绝对偏移量，完全避免子串多次匹配问题。

        【修复问题8】最后一个 chunk 如果 body_m 匹配失败走了 continue，
        _flush_pending_table 会被跳过，导致最后一段分片表格丢失。
        修复：在循环结束后、写 </body></html> 之前，无论如何都 flush 一次。

        【修复问题6】_clean_header_footer 改用嵌套深度匹配，
        避免误删正文 div。
        """
        if not chunk_html_paths:
            return

        marker_re = re.compile(
            r'<p[^>]*>\s*<span[^>]*>TABLE_SPLIT_MARKER::([a-f0-9]{8})</span>\s*</p>',
            re.IGNORECASE | re.DOTALL
        )

        pending_table_group = None
        pending_table_open  = None
        pending_trs         = []

        def _flush_pending_table(out_f):
            nonlocal pending_table_group, pending_table_open, pending_trs
            if pending_table_group:
                out_f.write(f"{pending_table_open}\n")
                for trs in pending_trs:
                    out_f.write(trs)
                out_f.write("</table>\n")
                pending_table_group = None
                pending_table_open  = None
                pending_trs         = []

        def _extract_trs(text, search_start=0):
            """
            【修复问题2】接受 search_start 参数，在指定偏移后搜索表格，
            返回 (tbl_open_tag, trs_content, after_content, tbl_close_abs_end)。
            tbl_close_abs_end 是 </table> 在原始 body 字符串中的结束偏移，
            供调用方直接设置游标，完全避免子串重复匹配。
            """
            text_stripped = re.sub(r'^\s*<div[^>]*>\s*', '', text, flags=re.IGNORECASE)
            tbl_open  = re.search(r'<table[^>]*>', text_stripped, re.IGNORECASE)
            tbl_close = re.search(r'</table>', text_stripped, re.IGNORECASE)
            if tbl_open and tbl_close:
                trs   = text_stripped[tbl_open.end():tbl_close.start()]
                after = text_stripped[tbl_close.end():]
                after = re.sub(r'^\s*</div>\s*', '', after, flags=re.IGNORECASE)
                after = re.sub(r'\s*</div>\s*$', '', after, flags=re.IGNORECASE)
                return tbl_open.group(0), trs, after
            return None, None, text

        def _clean_div_wrapper(text):
            text = re.sub(r'^\s*<div[^>]*>\s*', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\s*</div>\s*$', '', text, flags=re.IGNORECASE)
            return text

        def _fix_border_top(trs_content):
            """续接片第一行所有 td/th 加 border-top:none，消除接缝线。"""
            tr_end = re.search(r'</tr>', trs_content, re.IGNORECASE)
            if not tr_end:
                return trs_content
            first_row = trs_content[:tr_end.end()]
            rest      = trs_content[tr_end.end():]

            def fix_cell(m):
                tag = m.group(0)
                style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE)
                if style_m:
                    s = re.sub(r'border-top-style\s*:[^;]+;?', '', style_m.group(1), flags=re.IGNORECASE)
                    s = re.sub(r'border-top-width\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = re.sub(r'border-top-color\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = s.rstrip('; ') + '; border-top-style:none;'
                    tag = tag[:style_m.start()] + f'style="{s}"' + tag[style_m.end():]
                return tag

            return re.sub(r'<t[dh][^>]*>', fix_cell, first_row, flags=re.IGNORECASE) + rest

        with open(output_path, 'w', encoding='utf-8') as out_f:
            for file_idx, chunk_path in enumerate(chunk_html_paths):
                print(f"🔗 合并 chunk {file_idx}: {os.path.basename(chunk_path)}")

                with open(chunk_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                if file_idx == 0:
                    head_m = re.search(r'^(.*?<body[^>]*>)', content, re.DOTALL | re.IGNORECASE)
                    if head_m:
                        out_f.write(head_m.group(1) + '\n')

                body_m = re.search(r'<body[^>]*>(.*?)</body>', content, re.DOTALL | re.IGNORECASE)
                if not body_m:
                    print(f"   ⚠️ chunk {file_idx} 未找到 body，跳过")
                    continue

                body = body_m.group(1)

                # 【修复问题6】用改进后的 _clean_header_footer 清除页眉页脚
                body = self._clean_header_footer(body)

                # 【修复问题2】重构游标推进：全程使用正则匹配的绝对偏移量
                cursor = 0
                while cursor <= len(body):
                    marker_m = marker_re.search(body, cursor)

                    if not marker_m:
                        # 没有更多 marker，处理剩余内容
                        remaining = _clean_div_wrapper(body[cursor:])
                        if remaining.strip():
                            if pending_table_group:
                                tbl_open, trs, after = _extract_trs(remaining)
                                if trs is not None:
                                    pending_trs.append(_fix_border_top(trs))
                                    if after.strip():
                                        _flush_pending_table(out_f)
                                        out_f.write(after)
                            else:
                                out_f.write(remaining)
                        break

                    # 处理 marker 之前的普通内容
                    before = _clean_div_wrapper(body[cursor:marker_m.start()])
                    group  = marker_m.group(1)

                    if before.strip():
                        if pending_table_group and pending_table_group != group:
                            _flush_pending_table(out_f)
                        out_f.write(before)

                    # 【修复问题2】从 marker 结束位置开始搜索表格，
                    # 用正则直接在 body 上取绝对偏移，不做子串 index 查找
                    search_from = marker_m.end()
                    tbl_open_re  = re.search(r'<table[^>]*>',  body, re.IGNORECASE | re.DOTALL, )
                    # 重新在正确位置搜索
                    tbl_open_m  = re.search(r'<table[^>]*>',  body[search_from:], re.IGNORECASE)
                    tbl_close_m = re.search(r'</table>',       body[search_from:], re.IGNORECASE)

                    if tbl_open_m and tbl_close_m:
                        abs_open_end  = search_from + tbl_open_m.end()
                        abs_close_end = search_from + tbl_close_m.end()
                        trs_content   = body[abs_open_end : search_from + tbl_close_m.start()]
                        after_content = body[abs_close_end:]
                        after_content = re.sub(r'^\s*</div>\s*', '', after_content, flags=re.IGNORECASE)
                        after_content = re.sub(r'\s*</div>\s*$', '', after_content, flags=re.IGNORECASE)

                        if pending_table_group == group:
                            pending_trs.append(_fix_border_top(trs_content))
                        else:
                            _flush_pending_table(out_f)
                            pending_table_group = group
                            pending_table_open  = tbl_open_m.group(0)
                            pending_trs         = [trs_content]

                        # 【修复问题2】游标直接设为 </table> 的绝对结束位置
                        cursor = abs_close_end
                    else:
                        # marker 后没有找到表格，写出剩余内容
                        remaining = _clean_div_wrapper(body[search_from:])
                        if remaining.strip():
                            out_f.write(remaining)
                        break

                out_f.write('\n')
                print(f"   ✅ chunk {file_idx} 合并完成")

            # 【修复问题8】循环结束后无论如何都 flush，防止最后一个 chunk
            # body_m 匹配失败走 continue 时遗漏 pending table
            _flush_pending_table(out_f)
            out_f.write("</body>\n</html>\n")

        print(f"✅ 流式合并完成：{output_path}")

    # ------------------------------------------------------------------ #
    #  分片转换主流程                                                        #
    # ------------------------------------------------------------------ #

    def _chunked_docx_to_html(self, docx_path, html_path, temp_dir_prefix):
        """
        【内部方法】分片转换主流程
        步骤：拆分文档 → chunk各自转HTML（不内嵌图片）→ 流式合并 → 统一内嵌图片

        【修复问题5】temp_dir_prefix 由调用方传入，每次调用均唯一。
        """
        html_dir = os.path.dirname(html_path)
        chunk_dir = self._normalize_path(
            os.path.join(html_dir, f"{temp_dir_prefix}_chunks")
        )
        original_img_dir = self._normalize_path(
            os.path.join(chunk_dir, "original_images")
        )

        image_display_order = self._get_image_order_from_docx(docx_path)
        self._extract_original_images(docx_path, original_img_dir)

        # 提前读取显示尺寸映射（用于修正 2 倍大问题）
        img_size_map = self._extract_image_display_sizes(docx_path)

        try:
            chunk_paths = self._split_docx_to_chunks(
                docx_path, chunk_dir, image_display_order
            )

            chunk_html_paths = []
            for idx, chunk_path in enumerate(chunk_paths):
                chunk_html_path = self._normalize_path(
                    os.path.join(chunk_dir, f"chunk_{idx:04d}.html")
                )
                ok = self._docx_to_html_no_embed(chunk_path, chunk_html_path)
                if ok:
                    chunk_html_paths.append(chunk_html_path)
                else:
                    print(f"⚠️ chunk_{idx:04d} 转换失败，跳过")

            print(f"📋 收集到 {len(chunk_html_paths)} 个chunk HTML文件，开始流式合并...")

            self._merge_html_files_to_disk(chunk_html_paths, html_path)

            print("🖼️ 开始统一内嵌图片...")
            self._embed_images_to_html(html_path, image_display_order, original_img_dir)

            # 修正图片显示尺寸（DPI 2 倍问题）
            if img_size_map:
                print("📐 修正图片显示尺寸...")
                with open(html_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                # 提取 spire_img_names（此时已是 base64，从 src="data:..." 无法取文件名，
                # 改从内嵌前的 Spire 文件名列表构建；_embed_images_to_html 已完成替换，
                # 这里直接用 image_display_order 构建 identity 映射）
                html_content = self._fix_html_img_sizes(
                    html_content, img_size_map,
                    image_display_order,   # spire_img_names 与 order 同名
                    image_display_order
                )
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)

            # 修正表格宽度（超出 A4 版心时等比缩放）
            print("📏 修正表格宽度...")
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            html_content = self._fix_html_table_widths(html_content)
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            print(f"🎉 分片转换完成：{html_path}")

        except Exception as e:
            print(f"❌ 分片转换异常：{e}")
            import traceback
            traceback.print_exc()

        finally:
            if os.path.exists(chunk_dir):
                shutil.rmtree(chunk_dir, ignore_errors=True)
                print(f"🗑️ 清理chunk目录：{chunk_dir}")

        if os.path.exists(html_path):
            with open(html_path, 'r', encoding='utf-8') as f:
                return f.read()
        return ""

    # ------------------------------------------------------------------ #
    #  公开方法                                                             #
    # ------------------------------------------------------------------ #

    def docx_to_single_html(self, docx_path, html_path):
        """
        公开方法：DOCX转单文件HTML
        特性：图片Base64内嵌 | CSS内嵌 | 图片无压缩 | 顺序对齐
              超限自动分片 | 大表格按行拆分后还原 | 去除页眉页脚

        【修复问题5】每次调用生成独立的 temp_dir_prefix，
        避免多次调用或多线程并发时临时目录互相覆盖。

        :param docx_path: 输入DOCX文件路径（支持相对/绝对）
        :param html_path: 输出HTML文件路径（支持相对/绝对）
        :return: 生成的HTML文本内容，失败返回空字符串
        """
        # 1. 路径校验与标准化
        docx_path = self._normalize_path(docx_path)
        html_path = self._normalize_path(html_path)

        if not os.path.exists(docx_path):
            print(f"❌ 输入DOCX文件不存在（绝对路径）：{docx_path}")
            return ""

        html_dir = os.path.dirname(html_path)
        os.makedirs(html_dir, exist_ok=True)

        # 【修复问题5】每次调用生成新前缀
        temp_dir_prefix = self._make_temp_dir_prefix()

        # 2. 检测文档规模，超限走分片流程
        if self._needs_chunking(docx_path):
            return self._chunked_docx_to_html(docx_path, html_path, temp_dir_prefix)

        # 3. 创建唯一临时目录
        spire_temp_dir   = self._normalize_path(os.path.join(html_dir, temp_dir_prefix))
        original_img_dir = self._normalize_path(os.path.join(spire_temp_dir, "original_images"))
        spire_img_dir    = self._normalize_path(os.path.join(spire_temp_dir, "images"))

        # 4. 解析图片顺序 + 提取原始图片 + 提取显示尺寸
        image_display_order = self._get_image_order_from_docx(docx_path)
        extracted_imgs = self._extract_original_images(docx_path, original_img_dir)
        img_size_map   = self._extract_image_display_sizes(docx_path)

        if not image_display_order and extracted_imgs:
            image_display_order = sorted(extracted_imgs)
            print(f"⚠️ 顺序解析为空，兜底使用：{image_display_order}")

        # 5. Spire转换生成临时HTML
        document = Document()
        try:
            document.LoadFromFile(docx_path)
            document.HtmlExportOptions.ImageEmbedded = False
            document.HtmlExportOptions.ImagesPath = spire_img_dir
            document.HtmlExportOptions.ImageFormat = self.default_image_format
            document.SaveToFile(html_path, FileFormat.Html)
        except Exception as e:
            print(f"❌ Spire转换HTML失败：{e}")
            return ""
        finally:
            document.Close()
            del document

        # 6. 读取HTML，统一路径分隔符
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        html_content = html_content.replace('\\', '/')

        # 7. 内嵌CSS（赋值回 html_content，后续操作在内存中继续）
        css_file_path = self._normalize_path(os.path.splitext(html_path)[0] + '_styles.css')
        if os.path.exists(css_file_path):
            with open(css_file_path, 'r', encoding='utf-8') as f:
                css_content = f.read()
            html_content = re.sub(r'<link[^>]*href="[^"]+\.css"[^>]*>', '', html_content)
            html_content = html_content.replace(
                '</head>',
                f'<style type="text/css">\n{css_content}\n</style>\n</head>'
            )
            print("✅ 已内嵌CSS样式")

        # 8. 提取Spire生成的图片文件名列表
        img_pattern = re.compile(r'<img[^>]*src="([^"]+)"[^>]*>')
        spire_img_names = []
        for match in img_pattern.finditer(html_content):
            spire_img_name = os.path.basename(self._normalize_path(match.group(1)))
            if spire_img_name not in spire_img_names:
                spire_img_names.append(spire_img_name)
        print(f"=== Spire 图片列表：{spire_img_names} ===")

        # 9. 【修复问题1】用名称匹配构建映射，不依赖纯位置索引
        actual_spire_img_dir = self._find_actual_img_dir(spire_img_dir)
        name_map = self._build_spire_to_original_map(
            spire_img_names, image_display_order, original_img_dir, actual_spire_img_dir
        )

        for spire_name, img_path in name_map.items():
            base64_str = self._image_to_base64(img_path)
            if not base64_str:
                continue

            html_content = re.compile(
                r'src="[^"]*/?' + re.escape(spire_name) + r'"'
            ).sub(f'src="{base64_str}"', html_content)
            print(f"🔄 {spire_name} → {os.path.basename(img_path)} 已转为Base64")

        # 9b. 修正图片显示尺寸（DPI 2 倍问题）
        # 必须在 base64 替换后执行：此时 src 已是 data URI，
        # _fix_html_img_sizes 通过 spire_img_names/image_display_order 映射定位 img 标签
        if img_size_map:
            print("📐 修正图片显示尺寸...")
            html_content = self._fix_html_img_sizes(
                html_content, img_size_map, spire_img_names, image_display_order
            )

        # 9c. 修正表格宽度（超出 A4 版心时等比缩放）
        print("📏 修正表格宽度...")
        html_content = self._fix_html_table_widths(html_content)

        # 10. 保存最终HTML
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        # 11. 清理临时文件
        for temp_path in [spire_temp_dir, css_file_path]:
            if os.path.exists(temp_path):
                try:
                    if os.path.isdir(temp_path):
                        shutil.rmtree(temp_path, ignore_errors=True)
                    else:
                        os.remove(temp_path)
                    print(f"🗑️ 清理临时文件：{temp_path}")
                except Exception as e:
                    print(f"⚠️ 清理临时文件失败 {temp_path}：{e}")

        print(f"\n🎉 DOCX转HTML完成！")
        print(f"📄 最终文件绝对路径：{html_path}")
        print(f"✅ 特性：图片Base64内嵌 | CSS内嵌 | 图片无压缩 | 顺序对齐")
        return html_content

    # ------------------------------------------------------------------ #
    #  HTML → DOCX 分片工具                                                #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _html_count_paragraphs(html_text: str) -> int:
        """
        快速估算 HTML 中的段落数，用于判断是否需要分片。
        统计所有块级段落标签：<p>、<li>、<h1>~<h6>、<td>、<th>、<caption>、<dt>、<dd>。
        表格内的 <td>/<th> 也计入，与 Spire 的段落计数逻辑保持一致。
        """
        return len(re.findall(
            r'<(?:p|li|h[1-6]|td|th|caption|dt|dd)[\s>]',
            html_text,
            re.IGNORECASE
        ))

    def _split_html_to_chunks(self, html_text: str) -> list[str]:
        """
        将大 HTML 按段落数切分为若干 chunk，每个 chunk 均为完整的 HTML 文档。

        切分策略（只在顶层块级标签边界处切割，绝不在标签内部截断）：
        1. 提取 <head> 及 <body> 开标签，每个 chunk 都复用相同的 head（保留样式）。
        2. 扫描 <body> 内的顶层元素，按 _html_count_paragraphs 累计段落估算值。
        3. 累计值达到 MAX_PARAGRAPHS 时在当前顶层元素结束处切割，开启新 chunk。
        4. 表格（<table>）整体作为一个顶层元素，不在 <tr> 内部切割，
           避免 Spire 读到不完整表格结构。
           但如果单张表格本身段落数就超限，则按 <tr> 行粒度再做子切分。

        返回：chunk HTML 字符串列表，每项均为 <html>...<body>...</body></html>。
        """
        # ── 1. 提取 head + body 开标签 ──────────────────────────────────
        head_match = re.search(r'(<html[^>]*>.*?<body[^>]*>)', html_text,
                               re.DOTALL | re.IGNORECASE)
        if head_match:
            preamble = head_match.group(1)
            body_start = head_match.end()
        else:
            preamble = '<html><body>'
            body_start = 0

        body_end_match = re.search(r'</body\s*>', html_text, re.IGNORECASE)
        body_end = body_end_match.start() if body_end_match else len(html_text)
        body_content = html_text[body_start:body_end]

        # ── 2. 解析顶层块元素（用嵌套深度计数，不依赖 HTML 解析器）──────
        TOP_LEVEL_TAGS = re.compile(
            r'<(table|p|ul|ol|dl|h[1-6]|div|blockquote|pre|figure|section|article|header|footer|aside|nav|main)[\s>]',
            re.IGNORECASE
        )
        OPEN_TAG  = re.compile(r'<([a-zA-Z][a-zA-Z0-9]*)[\s/>]')
        CLOSE_TAG = re.compile(r'</([a-zA-Z][a-zA-Z0-9]*)\s*>')
        VOID_TAGS = {'br','hr','img','input','meta','link','area','base',
                     'col','embed','param','source','track','wbr'}

        def _find_top_level_blocks(text):
            """
            在 text 中按顺序找出所有顶层块元素的 (start, end) 区间。
            两个区间之间的纯文本/空白也作为独立片段返回，tag=None。
            返回：[(start, end, tag_name_or_None), ...]
            """
            blocks  = []
            cursor  = 0
            n       = len(text)

            while cursor < n:
                m = TOP_LEVEL_TAGS.search(text, cursor)
                if not m:
                    # 剩余纯文本
                    if cursor < n:
                        blocks.append((cursor, n, None))
                    break

                # 两个块之间的纯文本片段
                if m.start() > cursor:
                    blocks.append((cursor, m.start(), None))

                tag_name = m.group(1).lower()
                open_pos = m.start()

                if tag_name in VOID_TAGS:
                    blocks.append((open_pos, m.end(), tag_name))
                    cursor = m.end()
                    continue

                # 用深度计数找到对应的闭合标签
                depth  = 0
                pos    = open_pos
                end    = open_pos

                for tm in re.finditer(r'</?[a-zA-Z][a-zA-Z0-9]*[\s/>]?', text[open_pos:]):
                    raw = tm.group(0)
                    abs_start = open_pos + tm.start()
                    abs_end   = open_pos + tm.end()

                    # 自闭合或 void 标签不计深度
                    inner_tag = re.match(r'</?([a-zA-Z][a-zA-Z0-9]*)', raw)
                    if not inner_tag:
                        continue
                    inner_name = inner_tag.group(1).lower()

                    if raw.startswith('</'):
                        if inner_name == tag_name:
                            depth -= 1
                            if depth == 0:
                                # 找到闭合标签，扫到 > 结束
                                close_end = text.find('>', abs_end - 1)
                                end = (close_end + 1) if close_end != -1 else abs_end
                                break
                    elif inner_name not in VOID_TAGS and not raw.endswith('/>'):
                        if inner_name == tag_name or abs_start == open_pos:
                            depth += 1

                if end <= open_pos:
                    # 未找到匹配闭合标签，取到文本末尾
                    end = n

                blocks.append((open_pos, end, tag_name))
                cursor = end

            return blocks

        blocks = _find_top_level_blocks(body_content)

        # ── 3. 按段落数切分 blocks → chunks ─────────────────────────────
        chunks_html = []
        current_parts = []
        current_para_count = 0

        def _flush_chunk():
            nonlocal current_parts, current_para_count
            if current_parts:
                body_inner = ''.join(current_parts)
                chunks_html.append(f'{preamble}\n{body_inner}\n</body></html>')
                current_parts = []
                current_para_count = 0

        for start, end, tag in blocks:
            fragment = body_content[start:end]
            frag_paras = self._html_count_paragraphs(fragment)

            # 单张表格本身段落数超限 → 按 <tr> 行粒度子切分
            if tag == 'table' and frag_paras > self.MAX_PARAGRAPHS:
                _flush_chunk()
                sub_chunks = self._split_html_table_rows(fragment, preamble)
                chunks_html.extend(sub_chunks)
                continue

            # 普通块：超限时先 flush，再放入新 chunk
            if current_para_count + frag_paras > self.MAX_PARAGRAPHS and current_parts:
                _flush_chunk()

            current_parts.append(fragment)
            current_para_count += frag_paras

        _flush_chunk()

        print(f"📦 HTML 切分为 {len(chunks_html)} 个 chunk（总估算段落：{self._html_count_paragraphs(body_content)}）")
        return chunks_html

    def _split_html_table_rows(self, table_html: str, preamble: str) -> list[str]:
        """
        将单张超大 HTML 表格按 <tr> 行粒度切分为多个 chunk。
        每个 chunk 包含完整的 <table>...</table> 结构（含 <thead>/<tbody>/<tfoot>）。

        策略：
        - 如果存在 <thead>，将其作为每个 chunk 的固定表头（不重复计入段落数）。
        - <tbody>/<tfoot> 中的 <tr> 按 MAX_PARAGRAPHS 累计切割。
        - 每个 chunk 单独包在完整 HTML 文档中。
        """
        # 提取 <table ...> 开标签和 </table>
        tbl_open_m = re.match(r'<table[^>]*>', table_html, re.IGNORECASE)
        tbl_open   = tbl_open_m.group(0) if tbl_open_m else '<table>'

        # 提取 thead（作为每个 chunk 的固定表头）
        thead_m = re.search(r'<thead[\s>].*?</thead\s*>', table_html,
                            re.IGNORECASE | re.DOTALL)
        thead_html = thead_m.group(0) if thead_m else ''

        # 提取所有 <tr>（跳过 thead 内部的行）
        body_area = table_html
        if thead_m:
            body_area = table_html[:thead_m.start()] + table_html[thead_m.end():]

        tr_pattern = re.compile(r'<tr[\s>].*?</tr\s*>', re.IGNORECASE | re.DOTALL)
        all_trs = tr_pattern.findall(body_area)

        chunks_html = []
        current_trs = []
        current_para_count = 0

        def _flush_table_chunk():
            nonlocal current_trs, current_para_count
            if current_trs:
                body_inner = (
                    f'{tbl_open}\n'
                    f'{thead_html}\n'
                    f'<tbody>\n{"".join(current_trs)}\n</tbody>\n'
                    f'</table>'
                )
                chunks_html.append(f'{preamble}\n{body_inner}\n</body></html>')
                current_trs = []
                current_para_count = 0

        for tr in all_trs:
            tr_paras = self._html_count_paragraphs(tr)
            if current_para_count + tr_paras > self.MAX_PARAGRAPHS and current_trs:
                _flush_table_chunk()
            current_trs.append(tr)
            current_para_count += tr_paras

        _flush_table_chunk()

        print(f"   📊 超大表格按行切分为 {len(chunks_html)} 个 chunk（{len(all_trs)} 行）")
        return chunks_html

    def _html_chunk_to_docx(self, html_chunk: str, output_path: str,
                             temp_img_dir: str | None) -> bool:
        """
        将单个 HTML chunk 转为 DOCX（通过 Spire）。
        temp_img_dir 是图片解包目录，临时 HTML 写到同一目录，
        使 Spire 通过相对路径加载图片，避免中文/空格路径解析失败。
        """
        document       = None
        temp_html_path = None
        try:
            # 临时 HTML 与图片放在同一目录，让 Spire 用相对路径找图片
            html_dir = (temp_img_dir if temp_img_dir and os.path.isdir(temp_img_dir)
                        else os.path.dirname(output_path))
            os.makedirs(html_dir, exist_ok=True)

            temp_html_path = self._normalize_path(
                os.path.join(html_dir, f"_chunk_{uuid.uuid4().hex[:8]}.html")
            )

            # 将 src 绝对路径转为相对路径，降低 Spire 路径解析失败的概率
            html_to_write = html_chunk
            if temp_img_dir and os.path.isdir(temp_img_dir):
                def _to_rel_src(m):
                    tag   = m.group(0)
                    src_m = re.search(r'src="([^"]+)"', tag)
                    if not src_m:
                        return tag
                    src = src_m.group(1)
                    if src.startswith('http') or src.startswith('data:'):
                        return tag
                    try:
                        rel = os.path.relpath(src, html_dir).replace('\\', '/')
                        return tag[:src_m.start()] + f'src="{rel}"' + tag[src_m.end():]
                    except ValueError:
                        return tag
                html_to_write = re.sub(
                    r'<img\b[^>]*>', _to_rel_src, html_chunk, flags=re.IGNORECASE
                )

            with open(temp_html_path, 'w', encoding='utf-8') as f:
                f.write(html_to_write)

            document = Document()
            document.LoadFromFile(temp_html_path, FileFormat.Html, self.html_validation_type)
            document.SaveToFile(output_path, FileFormat.Docx2016)
            return True

        except Exception as e:
            print(f"   ❌ chunk 转 DOCX 失败（{os.path.basename(output_path)}）：{e}")
            return False
        finally:
            if document:
                document.Close()
                del document
            if temp_html_path and os.path.exists(temp_html_path):
                try:
                    os.remove(temp_html_path)
                except Exception:
                    pass

    def _merge_docx_chunks(self, chunk_docx_paths: list[str], output_docx_path: str) -> bool:
        """
        将多个 chunk DOCX 文件用 python-docx + lxml 合并为一个完整 DOCX。

        合并策略：
        - 以第一个 chunk 的样式/编号/页面设置为基础文档。
        - 逐个追加后续 chunk 的 body 顶层元素（w:p / w:tbl），跳过 w:sectPr。
        - 媒体资源（word/media/*）从各 chunk 读取并写入最终文档，
          文件名冲突时自动重命名（加 _cN 后缀），同时修正 document.xml 中的引用路径。
        - 图片 rId 在各 chunk 中独立编号，合并时对后续 chunk 做 rId 重映射，
          避免最终文档内 rId 冲突导致图片错乱。

        注意：此方法不依赖 Spire，纯 python-docx + zipfile 操作，不受段落限制。
        """
        if not chunk_docx_paths:
            return False
        if len(chunk_docx_paths) == 1:
            shutil.copy2(chunk_docx_paths[0], output_docx_path)
            return True

        # ── 第一步：以 chunk_0 为基础，收集其所有媒体文件名 ─────────────
        # 用 zipfile 直接操作，避免 python-docx 的关系系统限制

        # 读取所有 chunk 的原始 zip 内容
        def _read_zip(path):
            files = {}
            with zipfile.ZipFile(path, 'r') as zf:
                for item in zf.infolist():
                    files[item.filename] = zf.read(item.filename)
            return files

        base_files = _read_zip(chunk_docx_paths[0])

        # 当前已用的媒体文件名集合
        used_media = {
            os.path.basename(k)
            for k in base_files
            if k.startswith('word/media/')
        }

        # 解析 base 的 document.xml 和 rels
        RELS_PATH = 'word/_rels/document.xml.rels'
        base_doc_xml  = base_files.get('word/document.xml', b'').decode('utf-8')
        base_rels_xml = base_files.get(RELS_PATH, b'').decode('utf-8')

        # 解析 base rels，找到最大 rId 编号，后续 chunk 从此续编
        existing_rids = re.findall(r'Id="(rId\d+)"', base_rels_xml)
        max_rid = max((int(r[3:]) for r in existing_rids), default=0)

        # 提取 base body 内容（</body> 之前的部分，不含 sectPr）
        # 我们将在 base_doc_xml 中的 </body> 前插入所有后续 chunk 的内容
        # 用 lxml 操作更安全
        base_doc_tree = etree.fromstring(base_doc_xml.encode('utf-8'))
        base_body = base_doc_tree.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')

        # 移除 base body 末尾的 sectPr（合并后统一放在最末）
        W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        sectPr_tag = f'{{{W_NS}}}sectPr'
        base_sectPr = base_body.find(sectPr_tag)
        if base_sectPr is not None:
            base_body.remove(base_sectPr)

        # 新增的 rels 条目和媒体文件
        extra_rels  = []
        extra_media = {}  # zip_path → bytes

        # ── 第二步：逐个处理后续 chunk ──────────────────────────────────
        for chunk_idx, chunk_path in enumerate(chunk_docx_paths[1:], start=1):
            chunk_files = _read_zip(chunk_path)
            chunk_doc_xml  = chunk_files.get('word/document.xml', b'').decode('utf-8')
            chunk_rels_xml = chunk_files.get(RELS_PATH, b'').decode('utf-8')

            # 解析 chunk rels：建立 rId → (type, target) 映射
            chunk_rels_map = {}
            for m in re.finditer(r'<Relationship\s+Id="(rId\d+)"\s+Type="([^"]+)"\s+Target="([^"]+)"[^/]*/>', chunk_rels_xml):
                chunk_rels_map[m.group(1)] = (m.group(2), m.group(3))

            # 为 chunk 中的每个 image/hyperlink rId 分配新 rId，并注册到 base rels
            rid_remap = {}  # 旧 rId → 新 rId

            for old_rid, (rel_type, target) in chunk_rels_map.items():
                # 只迁移 image 和 hyperlink 关系，其他（header/footer/styles等）跳过
                rel_type_lower = rel_type.lower()
                if not any(k in rel_type_lower for k in ('image', 'hyperlink', 'oleobject')):
                    continue

                max_rid += 1
                new_rid = f'rId{max_rid}'
                rid_remap[old_rid] = new_rid

                if 'image' in rel_type_lower or 'oleobject' in rel_type_lower:
                    # 媒体文件：处理命名冲突
                    orig_fname  = os.path.basename(target)
                    new_fname   = orig_fname
                    fname_stem  = os.path.splitext(orig_fname)[0]
                    fname_ext   = os.path.splitext(orig_fname)[1]

                    if new_fname in used_media:
                        new_fname = f'{fname_stem}_c{chunk_idx}{fname_ext}'
                        # 极端情况：还是冲突，加 uuid
                        if new_fname in used_media:
                            new_fname = f'{fname_stem}_{uuid.uuid4().hex[:6]}{fname_ext}'

                    used_media.add(new_fname)
                    new_target  = f'media/{new_fname}'
                    src_zip_path = f'word/media/{orig_fname}'

                    if src_zip_path in chunk_files:
                        extra_media[f'word/{new_target}'] = chunk_files[src_zip_path]

                    extra_rels.append(
                        f'<Relationship Id="{new_rid}" Type="{rel_type}" Target="{new_target}"/>'
                    )
                else:
                    # hyperlink 等外部关系
                    extra_rels.append(
                        f'<Relationship Id="{new_rid}" Type="{rel_type}" Target="{target}" TargetMode="External"/>'
                    )

            # 对 chunk document.xml 做 rId 替换
            chunk_doc_patched = chunk_doc_xml
            # 按 rId 编号从大到小替换，避免 rId9 误替换 rId99 的子串
            for old_rid in sorted(rid_remap, key=lambda r: int(r[3:]), reverse=True):
                new_rid = rid_remap[old_rid]
                chunk_doc_patched = chunk_doc_patched.replace(
                    f'"{old_rid}"', f'"{new_rid}"'
                )

            # 解析 chunk body，提取顶层 w:p / w:tbl（跳过 w:sectPr）
            try:
                chunk_tree = etree.fromstring(chunk_doc_patched.encode('utf-8'))
                chunk_body = chunk_tree.find(f'{{{W_NS}}}body')
                if chunk_body is None:
                    print(f"   ⚠️ chunk {chunk_idx} 无 body，跳过")
                    continue
                for el in list(chunk_body):
                    if el.tag == sectPr_tag:
                        continue
                    base_body.append(copy.deepcopy(el))
            except Exception as e:
                print(f"   ⚠️ chunk {chunk_idx} XML 解析失败：{e}，跳过")
                continue

            print(f"   ✅ chunk {chunk_idx} 合并完成（rId 重映射 {len(rid_remap)} 条，媒体 {len(extra_media)} 个）")

        # ── 第三步：把 base_sectPr 追加回 body 末尾 ─────────────────────
        if base_sectPr is not None:
            base_body.append(base_sectPr)

        # ── 第四步：重组最终 DOCX zip ────────────────────────────────────
        new_doc_xml = etree.tostring(base_doc_tree, xml_declaration=True,
                                     encoding='UTF-8', standalone=True)

        # 更新 rels
        new_rels_xml = base_rels_xml.replace(
            '</Relationships>',
            '\n'.join(extra_rels) + '\n</Relationships>'
        )

        tmp_path = output_docx_path + '.mergetmp'
        with zipfile.ZipFile(chunk_docx_paths[0], 'r') as src_zip, \
             zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as dst_zip:

            for item in src_zip.infolist():
                if item.filename == 'word/document.xml':
                    dst_zip.writestr(item, new_doc_xml)
                elif item.filename == RELS_PATH:
                    dst_zip.writestr(item, new_rels_xml.encode('utf-8'))
                else:
                    dst_zip.writestr(item, src_zip.read(item.filename))

            for zip_path, data in extra_media.items():
                dst_zip.writestr(zip_path, data)

        os.replace(tmp_path, output_docx_path)
        print(f"✅ DOCX 合并完成：{output_docx_path}")
        return True

    # ------------------------------------------------------------------ #
    #  公开方法（续）                                                        #
    # ------------------------------------------------------------------ #

    def html_text_to_docx(self, html_text: str, output_docx_path: str):
        """
        公开方法：HTML文本转DOCX

        支持超大 HTML（段落 > MAX_PARAGRAPHS）自动切片转换，绕过 Spire 免费版限制。

        流程：
        1. 图片预处理：
           a. 修正 MIME 类型声明与实际格式不匹配问题（如声明 jpeg 实际是 webp）
           b. 修正 height:auto → 根据物理宽高比等比计算实际高度（px）
           c. base64 data URI 解包为临时图片文件（避免 Spire 因 data URI 过长崩溃）
           d. 外链 URL 图片下载后解包（带超时缓存）
           e. A4 版心尺寸约束
        2. 估算段落数，小文档直接转，大文档走分片流程
        3. 分片流程：HTML 切分 → 各片独立转 DOCX → python-docx 合并
        4. 清理所有临时文件

        :param html_text: 输入HTML字符串
        :param output_docx_path: 输出DOCX文件路径（支持相对/绝对）
        :return: 成功返回True，失败返回False
        """
        output_docx_path = self._normalize_path(output_docx_path)

        if not html_text.strip():
            print("❌ HTML文本为空，无法转换")
            return False

        output_dir   = os.path.dirname(output_docx_path)
        os.makedirs(output_dir, exist_ok=True)

        temp_img_dir  = None
        chunk_dir     = None

        try:
            # ── 步骤1a：修正 MIME 错误 + height:auto + base64/URL 图片解包 ──
            # 必须在 Spire 加载 HTML 前完成，三件事一次遍历全部 <img> 标签处理
            html_text, temp_img_dir = self._extract_base64_images(html_text, output_dir)

            # ── 步骤1b：A4 版心尺寸约束（已有明确 px 尺寸时做等比缩放）──────
            html_text = self._fix_html_img_sizes_for_import(html_text)

            # ── 步骤2：判断是否需要分片 ──────────────────────────────────
            para_count = self._html_count_paragraphs(html_text)
            print(f"📊 HTML 段落估算：{para_count}，阈值：{self.MAX_PARAGRAPHS}")

            if para_count <= self.MAX_PARAGRAPHS:
                print("✅ 无需分片，直接转换")
                return self._html_chunk_to_docx(html_text, output_docx_path, temp_img_dir)

            # ── 步骤3：分片流程 ──────────────────────────────────────────
            print(f"⚡ 触发 HTML 分片转换（段落估算 {para_count} > {self.MAX_PARAGRAPHS}）")

            chunk_dir = self._normalize_path(
                os.path.join(output_dir, f"html2docx_{uuid.uuid4().hex[:8]}")
            )
            os.makedirs(chunk_dir, exist_ok=True)

            html_chunks = self._split_html_to_chunks(html_text)

            chunk_docx_paths = []
            for idx, chunk_html in enumerate(html_chunks):
                chunk_docx_path = self._normalize_path(
                    os.path.join(chunk_dir, f"chunk_{idx:04d}.docx")
                )
                ok = self._html_chunk_to_docx(chunk_html, chunk_docx_path, temp_img_dir)
                if ok:
                    chunk_docx_paths.append(chunk_docx_path)
                    print(f"   ✅ chunk_{idx:04d} 转换完成")
                else:
                    print(f"   ⚠️ chunk_{idx:04d} 转换失败，跳过")

            if not chunk_docx_paths:
                print("❌ 所有 chunk 均转换失败")
                return False

            # ── 步骤4：合并所有 chunk DOCX ──────────────────────────────
            print(f"🔗 开始合并 {len(chunk_docx_paths)} 个 chunk DOCX...")
            return self._merge_docx_chunks(chunk_docx_paths, output_docx_path)

        except Exception as e:
            print(f"❌ HTML转DOCX失败：{str(e)}")
            import traceback
            traceback.print_exc()
            return False

        finally:
            if temp_img_dir and os.path.exists(temp_img_dir):
                try:
                    shutil.rmtree(temp_img_dir, ignore_errors=True)
                    print(f"🗑️ 清理图片临时目录：{temp_img_dir}")
                except Exception as e:
                    print(f"⚠️ 清理图片临时目录失败：{e}")
            if chunk_dir and os.path.exists(chunk_dir):
                try:
                    shutil.rmtree(chunk_dir, ignore_errors=True)
                    print(f"🗑️ 清理chunk目录：{chunk_dir}")
                except Exception as e:
                    print(f"⚠️ 清理chunk目录失败：{e}")


    # ------------------------------------------------------------------ #
    #  图片预处理工具（HTML→DOCX 方向）                                     #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _decode_b64_safe(b64_raw: str) -> bytes | None:
        """容忍换行/空白/残缺 padding 的 base64 解码，失败返回 None。"""
        try:
            clean   = re.sub(r'\s', '', b64_raw)
            padding = (4 - len(clean) % 4) % 4
            return base64.b64decode(clean + '=' * padding, validate=False)
        except Exception:
            return None

    @staticmethod
    def _read_image_wh(data: bytes):
        """
        从图片二进制头解析物理像素（w, h），支持 PNG/JPEG/GIF/BMP/WEBP。
        失败返回 (None, None)。
        """
        import struct
        try:
            if data[:8] == b'\x89PNG\r\n\x1a\n':
                return struct.unpack('>II', data[16:24])
            if data[:2] == b'\xff\xd8':
                i = 2
                while i < len(data) - 8:
                    if data[i] != 0xff: break
                    marker = data[i + 1]
                    if i + 3 >= len(data): break
                    length = struct.unpack('>H', data[i + 2:i + 4])[0]
                    if marker in (0xc0, 0xc1, 0xc2, 0xc3, 0xc5, 0xc6, 0xc7, 0xc9, 0xca, 0xcb):
                        h, w = struct.unpack('>HH', data[i + 5:i + 9])
                        return w, h
                    i += 2 + length
            if data[:6] in (b'GIF87a', b'GIF89a'):
                return struct.unpack('<HH', data[6:10])
            if data[:2] == b'BM':
                w, h = struct.unpack('<II', data[18:26])
                return w, abs(h)
            # WEBP：支持 VP8 / VP8L / VP8X 三种子格式
            if data[:4] == b'RIFF' and len(data) >= 30 and data[8:12] == b'WEBP':
                chunk = data[12:16]
                if chunk == b'VP8 ':
                    w = struct.unpack('<H', data[26:28])[0] & 0x3fff
                    h = struct.unpack('<H', data[28:30])[0] & 0x3fff
                    return w, h
                if chunk == b'VP8L' and len(data) >= 25:
                    bits = struct.unpack('<I', data[21:25])[0]
                    return (bits & 0x3fff) + 1, ((bits >> 14) & 0x3fff) + 1
                if chunk == b'VP8X' and len(data) >= 30:
                    w = (int.from_bytes(data[24:27], 'little') & 0xffffff) + 1
                    h = (int.from_bytes(data[27:30], 'little') & 0xffffff) + 1
                    return w, h
        except Exception:
            pass
        return None, None

    @staticmethod
    def _guess_mime(data: bytes) -> str:
        """从文件头推断真实 MIME 类型，默认返回 image/png。"""
        if data[:8]  == b'\x89PNG\r\n\x1a\n': return 'image/png'
        if data[:2]  == b'\xff\xd8':            return 'image/jpeg'
        if data[:6]  in (b'GIF87a', b'GIF89a'):  return 'image/gif'
        if data[:2]  == b'BM':                    return 'image/bmp'
        if data[:4]  == b'RIFF' and data[8:12] == b'WEBP': return 'image/webp'
        return 'image/png'

    def _extract_base64_images(self, html_text: str, base_dir: str):
        """
        HTML→DOCX 方向的图片预处理，返回 (modified_html, temp_img_dir)。

        一次遍历全部 <img> 标签，完成三件事：

        1. MIME 类型校正
           data URI 中声明的 MIME（如 image/jpeg）与实际二进制格式（如 WEBP）
           可能不一致，Spire 会按声明的 MIME 解析，遇到不符时静默丢弃图片。
           修复：解码后用文件头重新判断真实格式，写临时文件时用正确扩展名，
           同时把 HTML 里的 src data URI 替换为正确的本地文件路径。

        2. height:auto 等比推算
           CSS height:auto 语义是"按宽度等比缩放"，Spire 不理解此语义，
           会直接用 0 或物理像素高度，导致图片变形。
           修复：解码图片读出物理像素宽高，用 style width 推算正确高度（px），
           同时写入 HTML 属性 width/height 供 Spire 读取。

        3. base64 data URI 解包为本地文件
           Spire 对超长 data URI 字符串可能解析超时或静默丢图。
           修复：统一解包为临时文件，HTML src 替换为本地路径。

        外链 http/https URL 图片同样支持：下载后写临时文件，同步做 MIME 校正。
        没有任何图片时 temp_img_dir 返回 None。
        """
        import urllib.request

        MIME_TO_EXT = {
            'image/png':     '.png',
            'image/jpeg':    '.jpg',
            'image/gif':     '.gif',
            'image/bmp':     '.bmp',
            'image/webp':    '.webp',
            'image/svg+xml': '.svg',
            'image/tiff':    '.tif',
        }

        b64_re  = re.compile(
            r'data:(image/[a-zA-Z0-9+\-]+);base64,([A-Za-z0-9+/=\s]+)',
            re.DOTALL
        )
        img_tag_re = re.compile(r'<img\b[^>]*>', re.IGNORECASE | re.DOTALL)
        src_re     = re.compile(r'src="([^"]*)"', re.IGNORECASE)

        # ── 预处理：还原 JSON 序列化导致的 \" 转义引号 ────────────────────
        # HTML 被作为 JSON 字符串值传输时，属性引号变为 \"，
        # 导致 src="..." 正则完全匹配不到 src=\"...\" 形式的 src。
        # 统一先还原，后续正则统一用普通双引号处理。
        if '\\"' in html_text:
            html_text = html_text.replace('\\"', '"')
            print("   🔧 检测到 JSON 转义引号，已还原 \\\" → \"")

        # ── 收集所有需要处理的 img 标签 ──────────────────────────────────
        matches = list(img_tag_re.finditer(html_text))
        if not matches:
            return html_text, None

        temp_img_dir = self._normalize_path(
            os.path.join(base_dir, f"b64tmp_{uuid.uuid4().hex[:8]}")
        )
        os.makedirs(temp_img_dir, exist_ok=True)

        url_cache  = {}   # url → bytes，同一 URL 只下载一次
        patch_list = []   # [(start, end, new_tag_str)]
        has_any    = False

        for m in matches:
            tag    = m.group(0)
            src_m  = src_re.search(tag)
            src    = src_m.group(1).strip() if src_m else ''

            img_bytes = None

            # ── base64 data URI ──────────────────────────────────────────
            b64_m = b64_re.match(src) if src else None
            if b64_m:
                img_bytes = self._decode_b64_safe(b64_m.group(2))

            # ── 外链 URL ─────────────────────────────────────────────────
            elif src.startswith('http://') or src.startswith('https://'):
                if src in url_cache:
                    img_bytes = url_cache[src]
                else:
                    # 重试三次，每次超时递增（30s / 60s / 90s），
                    # 适应内网大文件或服务器响应较慢的场景
                    for attempt, timeout in enumerate((30, 60, 90), start=1):
                        try:
                            req = urllib.request.Request(
                                src, headers={'User-Agent': 'Mozilla/5.0'}
                            )
                            with urllib.request.urlopen(req, timeout=timeout) as resp:
                                img_bytes = resp.read()
                            url_cache[src] = img_bytes
                            print(f"   🌐 下载图片（第{attempt}次）：{src[:60]}  {len(img_bytes):,}B")
                            break
                        except Exception as e:
                            print(f"   ⚠️ 下载失败（第{attempt}次，timeout={timeout}s）{src[:60]}：{e}")
                            if attempt == 3:
                                url_cache[src] = None  # 标记为已尝试过，避免重复下载
                    if not img_bytes:
                        img_bytes = url_cache.get(src)

            # ── 本地文件（已由上游解包）──────────────────────────────────
            elif src and os.path.exists(src):
                try:
                    with open(src, 'rb') as f:
                        img_bytes = f.read()
                except Exception as e:
                    print(f"   ⚠️ 读取本地图片失败：{e}")

            if not img_bytes:
                # 无法获取图片内容，保留原标签不改动
                continue

            has_any = True

            # 1. 校正真实 MIME
            real_mime = self._guess_mime(img_bytes[:16])

            # Spire HTML 导入不支持 WEBP / SVG / TIFF / GIF（部分版本）等格式，
            # 遇到不支持的格式会静默退化为占位符图片（小方块/破图）。
            # 统一转换为 JPEG（有损但体积小），PNG 作为兜底（无损）。
            # 转换优先用 Pillow；Pillow 不可用时保留原始格式（Spire 自行处理）。
            SPIRE_UNSUPPORTED = {'image/webp', 'image/svg+xml', 'image/tiff',
                                 'image/bmp', 'image/gif'}
            if real_mime in SPIRE_UNSUPPORTED:
                converted = False
                try:
                    from PIL import Image as _PILImage
                    import io as _io
                    pil_img = _PILImage.open(_io.BytesIO(img_bytes)).convert('RGB')
                    buf = _io.BytesIO()
                    pil_img.save(buf, format='JPEG', quality=92)
                    img_bytes = buf.getvalue()
                    real_mime = 'image/jpeg'
                    converted = True
                    print(f"   🔄 {real_mime} 不受 Spire 支持，已用 Pillow 转换为 JPEG")
                except Exception as e:
                    print(f"   ⚠️ Pillow 转换失败（{real_mime}→JPEG）：{e}，保留原格式")

            ext   = MIME_TO_EXT.get(real_mime, '.png')
            fname = f"img_{len(patch_list):04d}{ext}"
            fpath = self._normalize_path(os.path.join(temp_img_dir, fname))
            with open(fpath, 'wb') as f:
                f.write(img_bytes)

            declared_mime = b64_m.group(1).lower() if b64_m else ''
            if declared_mime and declared_mime != real_mime:
                print(f"   🔧 MIME 校正：{declared_mime} → {real_mime}（{fname}）")

            # 2. 读物理像素，推算 height:auto
            phys_w, phys_h = self._read_image_wh(img_bytes)

            # 从 style 读 width 值（含单位换算）
            # style 属性值内部可能有换行缩进，先把换行/多余空白压缩为单空格
            style_m = re.search(r'style="([^"]*)"', tag, re.IGNORECASE | re.DOTALL)
            style_w_px = None
            if style_m:
                style_val = re.sub(r'[\r\n]+\s*', ' ', style_m.group(1))
                w_m = re.search(
                    r'\bwidth\s*:\s*([\d.]+)\s*(px|pt|in|cm|mm)?',
                    style_val, re.IGNORECASE
                )
                if w_m:
                    val  = float(w_m.group(1))
                    unit = (w_m.group(2) or 'px').lower()
                    if unit == 'px':  style_w_px = round(val)
                    elif unit == 'pt': style_w_px = round(val * 96 / 72)
                    elif unit == 'in': style_w_px = round(val * 96)
                    elif unit == 'cm': style_w_px = round(val / 2.54 * 96)
                    elif unit == 'mm': style_w_px = round(val / 25.4 * 96)

            # 确定最终 w/h（px）
            w_px = style_w_px or phys_w
            if w_px and phys_w and phys_h and phys_w > 0:
                h_px = round(w_px * phys_h / phys_w)
            else:
                h_px = phys_h

            # 3. 重写 img 标签：src 替换为本地路径，写入 width/height 属性
            new_tag = tag
            # 更新 src
            if src_m:
                new_tag = (new_tag[:src_m.start()]
                           + f'src="{fpath}"'
                           + new_tag[src_m.end():])

            # 写入 width/height 属性（移除旧的，插入新的）
            if w_px and h_px:
                new_tag = re.sub(r'\s+width="[^"]*"', '', new_tag, flags=re.IGNORECASE)
                new_tag = re.sub(r'\s+height="[^"]*"', '', new_tag, flags=re.IGNORECASE)
                new_tag = re.sub(r'(<img\b)', rf'\1 width="{w_px}" height="{h_px}"', new_tag)

                # 同步更新 style 中的 width/height（移除 max-width/height:auto）
                style_m2 = re.search(r'style="([^"]*)"', new_tag, re.IGNORECASE)
                if style_m2:
                    s = style_m2.group(1)
                    s = re.sub(r'\bmax-width\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = re.sub(r'\bmax-height\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = re.sub(r'\bwidth\s*:[^;]+;?',  '', s, flags=re.IGNORECASE)
                    s = re.sub(r'\bheight\s*:[^;]+;?', '', s, flags=re.IGNORECASE)
                    s = s.rstrip('; ')
                    new_style = f"{s}; width:{w_px}px; height:{h_px}px".lstrip('; ')
                    new_tag = (new_tag[:style_m2.start()]
                               + f'style="{new_style}"'
                               + new_tag[style_m2.end():])

            patch_list.append((m.start(), m.end(), new_tag))
            print(f"   📤 {fname}（{real_mime}，{w_px}×{h_px}px）")

        if not has_any:
            # 没有成功处理任何图片，清理空目录
            shutil.rmtree(temp_img_dir, ignore_errors=True)
            return html_text, None

        # ── 从后往前替换，避免位移 ───────────────────────────────────────
        result = list(html_text)
        for start, end, new_tag in reversed(patch_list):
            result[start:end] = list(new_tag)
        html_text = ''.join(result)

        print(f"   📦 共处理 {len(patch_list)} 张图片 → {temp_img_dir}")
        return html_text, temp_img_dir



# ------------------------------ 调用示例 ------------------------------
if __name__ == "__main__":
    converter = DocxHtmlConverter()

    # 示例1：DOCX转单文件HTML（自动判断是否需要分片）
    input_docx = r"input_langwithtable.docx"
    output_html = r"C:\Users\you62\Desktop\index.html"
    # html_content = converter.docx_to_single_html(input_docx, output_html)

    # 示例2：HTML文本转DOCX
    if os.path.exists(output_html):
        with open(output_html, 'r', encoding='utf-8') as f:
            sample_html = f.read()
        converter.html_text_to_docx(sample_html, "output.docx")
    else:
        print(f"错误：未找到HTML文件 {output_html}")