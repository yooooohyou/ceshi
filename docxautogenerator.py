from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.oxml.shared import qn as qn_shared
from docx.oxml import parse_xml
import os
from PIL import Image


def generate_report_doc(
        title_text,
        second_row_img_path,
        other_img_paths=None,
        save_path='年度报告.docx',
        columns_per_row=2  # 新增参数：指定每行显示的图片数量，默认2列
):
    other_img_paths = other_img_paths or []

    # 校验参数合法性
    if not isinstance(columns_per_row, int) or columns_per_row < 1:
        raise ValueError("每行图片数量必须是大于等于1的整数")

    # 校验图片文件是否存在
    if not os.path.exists(second_row_img_path):
        raise FileNotFoundError(f"图片不存在：{second_row_img_path}")
    for p in other_img_paths:
        if not os.path.exists(p):
            raise FileNotFoundError(f"图片不存在：{p}")

    # 创建干净文档
    doc = Document()

    # 字体设置函数
    def set_font(run, font_name, size, color=None, bold=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        run.bold = bold

    # 重新计算表格行数和列数
    # 总行数 = 标题行(1) + 大图行(1) + 其他图片需要的行数
    # 其他图片行数 = 向上取整(图片数量 / 每行列数)
    other_img_rows = (len(other_img_paths) + columns_per_row - 1) // columns_per_row
    total_rows = 2 + other_img_rows
    table = doc.add_table(rows=total_rows, cols=columns_per_row)  # 列数改为自定义值
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 蓝色边框设置函数
    def set_cell_border(cell, color):
        for name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:color'), color)
            border.set(qn('w:sz'), '6')
            cell._element.tcPr.append(border)

    blue = '4472C4'
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, blue)
            cell.vertical_alignment = 1  # WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 强制前两行不分页
    for row in table.rows[:2]:
        trPr_list = row._element.xpath('.//w:trPr')
        if not trPr_list:
            trPr = OxmlElement('w:trPr')
            row._element.insert(0, trPr)
        else:
            trPr = trPr_list[0]

        kl = OxmlElement('w:keepLines')
        kl.set(qn('w:val'), 'true')
        trPr.append(kl)

    # 第一行：标题（合并所有列）
    hdr_cells = table.rows[0].cells
    # 合并当前行的所有列
    for cell in hdr_cells[1:]:
        hdr_cells[0].merge(cell)
    para = hdr_cells[0].paragraphs[0]
    run = para.add_run(title_text)
    set_font(run, '宋体', 12, RGBColor(255, 255, 255), bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), blue)
    hdr_cells[0]._element.tcPr.append(shd)

    # 第二行：大图（合并所有列）
    row2_cells = table.rows[1].cells
    # 合并当前行的所有列
    for cell in row2_cells[1:]:
        row2_cells[0].merge(cell)
    para = row2_cells[0].paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(second_row_img_path, width=Inches(5.0))

    # 后面图片：按自定义列数排版
    # 动态计算每张图片的宽度（适配不同列数）
    img_width = Inches((5.0 * 0.9) / columns_per_row)  # 总宽度5英寸，留10%间距
    for i, img_path in enumerate(other_img_paths):
        r = 2 + i // columns_per_row  # 计算当前图片所在行
        c = i % columns_per_row  # 计算当前图片所在列
        cell = table.rows[r].cells[c]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(img_path, width=img_width)

    doc.save(save_path)
    return save_path


def calculate_table_height(table):
    total_height = 0.0
    for row in table.rows:
        row_height = row.height if row.height else 0
        if row_height == 0:
            for cell in row.cells:
                cell_height = 0.0
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.element.xpath('.//a:blip'):
                            cell_height = Inches(2.8).inches
                            break
                    if cell_height > 0:
                        break
                if cell_height > 0:
                    row_height = cell_height
                    break
            if row_height == 0:
                row_height = 0.2
        total_height += row_height / 914400
    return total_height


def set_table_no_page_break(table):
    """最强防跨页：整表强制在同一页，Word 必生效（无FutureWarning）"""
    for row in table.rows:
        tr = row._element
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)

        # 禁止行被分页拆开
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)

        # 整行保持在一页
        keepLines = OxmlElement('w:keepLines')
        keepLines.set(qn('w:val'), 'true')
        trPr.append(keepLines)

        # 所有行连在一起
        keepNext = OxmlElement('w:keepNext')
        keepNext.set(qn('w:val'), 'true')
        trPr.append(keepNext)

    # 最后一行不要 keepNext，否则会把后面内容也拽到同一页
    if table.rows:
        last_tr = table.rows[-1]._element
        last_trPr = last_tr.find(qn('w:trPr'))
        if last_trPr is not None:
            keep_next_elems = [elem for elem in last_trPr if elem.tag.endswith('keepNext')]
            for elem in keep_next_elems:
                last_trPr.remove(elem)


def generate_fully_centered_patent_doc(
        patent_data,
        cert_img_paths=None,
        save_path='专利信息全居中文档.docx',
        fill_empty_space=True,
        last_img_display_mode=1
):
    if last_img_display_mode not in (1, 2):
        raise ValueError("last_img_display_mode 只能是 1 或 2")

    cert_img_paths = cert_img_paths or []
    valid_img_paths = [p for p in cert_img_paths if os.path.exists(p)]
    print(f"📸 有效图片数量：{len(valid_img_paths)}")  # 新增：打印有效图片数，方便排查

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    page_h = sec.page_height.inches - 2
    half_page = page_h / 2

    # 字体：宋体 小四（12磅）
    def set_font(run, color):
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.font.color.rgb = color

    def set_head_font(run):
        set_font(run, RGBColor(255, 255, 255))
        run.bold = True

    def set_border(cell, color='4472C4', sz='2'):
        for k in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{k}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:color'), color)
            b.set(qn('w:sz'), sz)
            cell._element.tcPr.append(b)

    # 专利表格
    table1 = doc.add_table(rows=len(patent_data) + 1, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = True

    heads = ['序号', '专利类型', '专利名称', '专利号', '专利权人', '授权公告日']
    for i, (cell, txt) in enumerate(zip(table1.rows[0].cells, heads)):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        set_head_font(r)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '4472C4')
        cell._element.tcPr.append(shd)
        set_border(cell, '4472C4')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for i, row in enumerate(patent_data, 1):
        for j, (cell, txt) in enumerate(zip(table1.rows[i].cells, row)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(txt))
            set_font(r, RGBColor(0, 0, 0))
            set_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 专利表格禁止跨页
    # set_table_no_page_break(table1)

    table_h = calculate_table_height(table1)

    # 分隔空行
    doc.add_paragraph()

    # ========== 核心修正：补图逻辑（优先保留所有原始图片，不截断） ==========
    final_imgs = valid_img_paths.copy()
    if fill_empty_space and table_h < half_page and valid_img_paths:
        # 计算需要的总高度对应的图片行数（每2张占2.8英寸）
        needed_height = half_page - table_h
        needed_rows = max(1, int(needed_height / 2.8) + 1)
        needed_imgs_min = needed_rows * 2  # 填满页面一半需要的最少图片数

        # 仅当原始图片数量 < 最少需要数时，才补充图片（不再截断原始图片）
        if len(final_imgs) < needed_imgs_min:
            print(f"🔍 原始图片{len(final_imgs)}张不足，需补充到{needed_imgs_min}张")
            # 补充图片（重复原始图片）
            while len(final_imgs) < needed_imgs_min:
                final_imgs.extend(valid_img_paths)
            # 最多补充到需要的数量（不超）
            final_imgs = final_imgs[:needed_imgs_min]
        else:
            print(f"✅ 原始图片{len(final_imgs)}张足够，无需补充/截断")

    print(f"📸 最终插入图片数量：{len(final_imgs)}")  # 新增：打印最终插入数

    # 图片表格（索引逻辑已修复）
    if final_imgs:
        n = len(final_imgs)
        img_table = doc.add_table(rows=(n + 1) // 2, cols=2)
        img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        img_table.autofit = True

        blue = '4472C4'
        for row in img_table.rows:
            for cell in row.cells:
                set_border(cell, blue, '6')
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.width = Inches(2.8)

        # 奇数最后一行处理
        if n % 2 == 1:
            last_row = img_table.rows[-1]
            if last_img_display_mode == 1:
                # mode1：删除最后一行第二个单元格
                last_row.cells[1]._element.getparent().remove(last_row.cells[1]._element)
            elif last_img_display_mode == 2:
                # mode2：合并最后一行两个单元格（不删除）
                last_row.cells[0].merge(last_row.cells[1])
                last_row.cells[0].width = Inches(5.6)

        # 图片插入逻辑（确保全部插入）
        inserted_count = 0  # 新增：统计实际插入数
        for i in range(n):  # 直接遍历图片数量，而非enumerate
            img_path = final_imgs[i]
            r = i // 2  # 行索引
            # 修正列索引逻辑：mode2且最后一行时，列索引强制为0；否则正常取模
            if n % 2 == 1 and r == len(img_table.rows) - 1 and last_img_display_mode == 2:
                c = 0
            else:
                c = i % 2

            # 仅判断行是否越界（列索引已通过逻辑控制，不会越界）
            if r >= len(img_table.rows):
                continue

            # 获取单元格并插入图片
            cell = img_table.rows[r].cells[c]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(img_path, width=Inches(2.8))
            inserted_count += 1

        print(f"📸 实际插入图片数量：{inserted_count}")  # 新增：验证插入数

        # 图片表格禁止跨页
        set_table_no_page_break(img_table)

    doc.save(save_path)
    print(f"✅ 文档生成成功：{os.path.abspath(save_path)}")
    return save_path


def generate_car_info_doc(car_data, save_path='公司车辆信息.docx', table_title='公司车辆信息'):
    """
    1:1复刻原文档+匹配指定字体/边框样式
    :param car_data: 车辆数据，二维列表[[序号,车牌号,行驶证图片路径,车辆图片路径], ...]
    :param save_path: 文档保存路径
    :param table_title: 表格顶部大标题
    """
    # 初始化文档
    doc = Document()
    # 页面边距设置
    sec = doc.sections[0]
    sec.top_margin = Cm(1.0)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.0)
    sec.right_margin = Cm(1.0)

    # 核心配色（按指定函数定义：蓝色4472C4、白色、黑色）
    COLOR_BLUE = '4472C4'       # 边框/表头底纹色
    COLOR_WHITE = RGBColor(255, 255, 255)  # 表头文字色
    COLOR_BLACK = RGBColor(0, 0, 0)        # 正文文字色
    # 列宽配置（适配A4，匹配原文档比例）
    COL_WIDTHS = [Cm(1.5), Cm(2.5), Cm(8.0), Cm(5.0)]
    # 图片插入宽度（适配列宽无变形）
    IMG_WIDTH_DRIVE = Cm(7.5)
    IMG_WIDTH_CAR = Cm(4.5)

    # ===================== 按指令要求定义的字体/边框函数 =====================
    def set_font(run, color):
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.font.color.rgb = color

    def set_head_font(run):
        set_font(run, RGBColor(255, 255, 255))
        run.bold = True

    def set_border(cell, color='4472C4', sz='2'):
        for k in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{k}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:color'), color)
            b.set(qn('w:sz'), sz)
            cell._element.tcPr.append(b)
    # ==========================================================================

    # 自定义工具函数（基于指定函数扩展，保证样式统一）
    def set_cell_bg(cell, bg_color=COLOR_BLUE):
        """设置单元格底纹（表头蓝色底纹）"""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg_color)
        cell._element.tcPr.append(shd)

    def add_text_to_cell(cell, text, is_header=False):
        """向单元格添加文字，区分表头/正文样式"""
        # 清空原有内容
        for para in cell.paragraphs:
            for run in para.runs:
                para._element.remove(run._element)
        # 文字居中排版
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # 插入文字并应用样式
        run = para.add_run(str(text))
        if is_header:
            set_head_font(run)  # 表头：白色+加粗+12号宋体
        else:
            set_font(run, COLOR_BLACK)  # 正文：黑色+12号宋体
        # 单元格基础样式：垂直居中+蓝色边框
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_border(cell)
        # 表头添加蓝色底纹
        if is_header:
            set_cell_bg(cell)

    def add_img_to_cell(cell, img_path, img_width):
        """向单元格插入图片，带蓝色边框+居中"""
        # 清空原有内容
        for para in cell.paragraphs:
            for run in para.runs:
                para._element.remove(run._element)
        # 图片居中
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 插入图片/缺失提示
        if os.path.exists(img_path):
            para.add_run().add_picture(img_path, width=img_width)
        else:
            # run = para.add_run(f'图片缺失：{os.path.basename(img_path)}')
            # set_font(run, RGBColor(255, 0, 0))  # 红色提示
            run = para.add_run(f'')
            # set_font(run, RGBColor(255, 0, 0))  # 红色提示
        # 单元格样式：垂直居中+蓝色边框
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_border(cell)

    # 数据校验
    if not isinstance(car_data, list) or len(car_data) == 0:
        raise ValueError("car_data必须是非空的二维列表")
    for idx, row in enumerate(car_data):
        if len(row) != 4:
            raise ValueError(f"car_data第{idx+1}行必须包含[序号,车牌号,行驶证路径,车辆图片路径]4个元素")

    # 表格创建：2行表头 + 数据行
    total_rows = 2 + len(car_data)
    total_cols = 4
    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体页面居中
    table.autofit = False  # 关闭自动适应，固定列宽

    # 设置固定列宽（修复index错误，正确遍历）
    for row in table.rows:
        for col_idx in range(total_cols):
            row.cells[col_idx].width = COL_WIDTHS[col_idx]

    # 第一行：合并4列+大标题（表头样式：白字+蓝底+蓝色边框）
    first_row_cells = table.rows[0].cells
    main_title_cell = first_row_cells[0]
    for cell in first_row_cells[1:]:
        main_title_cell.merge(cell)
    add_text_to_cell(main_title_cell, table_title, is_header=True)

    # 第二行：二级表头（序号/车牌号/行驶证/车辆图片，表头样式）
    second_headers = ['序号', '车牌号', '行驶证', '车辆图片']
    second_row_cells = table.rows[1].cells
    for col_idx, header in enumerate(second_headers):
        add_text_to_cell(second_row_cells[col_idx], header, is_header=True)

    # 数据行：填充车辆信息+插入图片（正文样式）
    for data_idx, car_row in enumerate(car_data, start=2):
        seq, plate_num, drive_img, car_img = car_row
        current_cells = table.rows[data_idx].cells
        # 序号、车牌号（文字）
        add_text_to_cell(current_cells[0], seq, is_header=False)
        add_text_to_cell(current_cells[1], plate_num, is_header=False)
        # 行驶证、车辆图片（图片）
        add_img_to_cell(current_cells[2], drive_img, IMG_WIDTH_DRIVE)
        add_img_to_cell(current_cells[3], car_img, IMG_WIDTH_CAR)

    # 保存文档
    doc.save(save_path)
    print(f"✅ 文档生成成功！路径：{os.path.abspath(save_path)}")
    return save_path

if __name__ == "__main__":
    # 示例参数
    # title = "2024年年度报告"  # 第一行标题
    # main_img = "./pic/图片1.jpg"  # 第二行主图片路径
    # other_imgs = ["./pic/图片2.jpg", "./pic/图片3.jpg", "./pic/图片4.jpg", "./pic/图片5.jpg",  "./pic/图片3.jpg", "./pic/图片4.jpg", "./pic/图片5.jpg"]
    # save_file = "./2024年度报告.docx"  # 保存路径
    #
    # # 调用函数生成文档
    # generate_report_doc(
    #     title_text=title,
    #     second_row_img_path=main_img,
    #     other_img_paths=other_imgs,
    #     save_path=save_file,
    #     columns_per_row=4
    # )
    # patent_data = [
    #     ['1', '发明专利', '一种核岛用防爆电话电缆', 'ZL201210300077.1', '远程电缆股份有限公司', '2015-07-29'],
    #     ['2', '发明专利', '一种电缆保护夹座安装方法', 'ZL201110454443.4', '远程电缆股份有限公司', '2014-06-25'],
    #     ['3', '发明专利', '硅烷交联聚乙烯绝缘电缆料及其制造方法', 'ZL200710022494.3', '远程电缆股份有限公司',
    #      '2010-12-08'],
    #     ['4', '发明专利', '多色架空绝缘电缆及其制造方法', 'ZL200710019537.2', '远程电缆股份有限公司', '2010-01-13'],
    #     ['5', '发明专利', '一种用于5G传输技术的配电专用线及其生产工艺', 'ZL201911360899.7', '远程电缆股份有限公司',
    #      '2020-09-25'],
    #     ['6', '发明专利', '一种核岛用防爆电话电缆', 'ZL201210300077.1', '远程电缆股份有限公司', '2015-07-29'],
    #     ['7', '发明专利', '一种电缆保护夹座安装方法', 'ZL201110454443.4', '远程电缆股份有限公司', '2014-06-25'],
    #     ['8', '发明专利', '硅烷交联聚乙烯绝缘电缆料及其制造方法', 'ZL200710022494.3', '远程电缆股份有限公司',
    #      '2010-12-08'],
    #     ['9', '发明专利', '多色架空绝缘电缆及其制造方法', 'ZL200710019537.2', '远程电缆股份有限公司', '2010-01-13'],
    #     ['10', '发明专利', '一种用于5G传输技术的配电专用线及其生产工艺', 'ZL201911360899.7', '远程电缆股份有限公司',
    #      '2020-09-25']
    # ]
    #
    # cert_img_paths = ["./pic/图片2.jpg", "./pic/图片3.jpg", "./pic/图片4.jpg", "./pic/图片5.jpg", "./pic/图片6.jpg", "./pic/图片7.jpg", "./pic/图片8.jpg", "./pic/图片3.jpg", "./pic/图片4.jpg", "./pic/图片5.jpg", "./pic/图片6.jpg", "./pic/图片7.jpg", "./pic/图片8.jpg"]
    #
    # save_path = "./专利报告.docx"
    # generate_fully_centered_patent_doc(patent_data, cert_img_paths, save_path, last_img_display_mode=1)
    car_info_data = [
        [1, '苏BSG126', './pic/图片13.jpg', './pic/图片14.jpg'],
        [2, '苏B5706L', './pic/图片15.jpg', './pic/图片16.jpg'],
        [3, '苏B950UF', '', './pic/图片18.jpg'],
        [4, '苏BJ0106', './pic/图片19.jpg', './pic/图片20.jpg'],
        # [5, '苏BU533V', './pic/图片21.jpg', './pic/图片22.jpg'],
    ]

    # 调用函数生成文档
    generate_car_info_doc(
        car_data=car_info_data,
        save_path='公司车辆信息.docx',
        table_title='公司车辆信息'
    )