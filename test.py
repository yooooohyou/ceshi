import time

from fastapi import FastAPI, HTTPException, Query, Body, UploadFile, File, Request
from fastapi.responses import HTMLResponse, FileResponse
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.oxml.shared import qn as qn_shared
from docx.oxml import parse_xml
import os
from PIL import Image
import shutil
import uuid
from typing import List, Optional, Any
import aiohttp
import asyncio
import tempfile
import zipfile
import re
import base64

# ========== 导入Spire.Doc相关模块 ==========
try:
    from spire.doc import *
    from spire.doc.common import *
except ImportError:
    print("⚠️ 未找到Spire.Doc库，请安装：pip install Spire.Doc")
    raise

# 创建FastAPI应用
app = FastAPI(
    title="Word文档生成API",
    description="生成包含图片的Word文档，并转换为单文件HTML返回（图片Base64内嵌）",
    version="1.0.0"
)

# 临时文件目录（用于存储下载的图片和生成的文档）
TEMP_DIR = "./uploads"
os.makedirs(TEMP_DIR, exist_ok=True)


# ===================== DOCX转HTML工具类 =====================
class DocxHtmlConverter:
    """
    DOCX与HTML互转工具类（修复外部引用路径问题）
    核心改进：
    1. 所有路径强制使用绝对路径，脱离工作目录依赖
    2. 临时目录使用唯一ID命名，避免多线程/多调用冲突
    3. 图片路径匹配改为全量绝对路径匹配
    4. 增强路径校验和异常处理
    """

    def __init__(self):
        """初始化转换器，可自定义全局配置"""
        self.default_image_format = 0  # 0=PNG，1=JPG，2=BMP，3=GIF
        self.html_validation_type = XHTMLValidationType.none
        # 生成唯一临时目录前缀，避免冲突
        self.temp_dir_prefix = f"spire_temp_{uuid.uuid4().hex[:8]}"

    def _normalize_path(self, path):
        """【内部方法】统一路径格式并转为绝对路径"""
        if not path:
            return ""
        # 先转为绝对路径，再统一分隔符
        abs_path = os.path.abspath(path)
        return abs_path.replace('\\', '/').replace('//', '/')

    def _get_image_order_from_docx(self, docx_path):
        """【内部方法】解析DOCX，提取图片在文档中的显示顺序"""
        image_order = []
        try:
            # 1. 解析关系文件，建立ID→图片名映射
            rels_mapping = {}
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                if 'word/_rels/document.xml.rels' in zip_file.namelist():
                    rels_content = zip_file.read('word/_rels/document.xml.rels').decode('utf-8')
                    rel_pattern = re.compile(
                        r'<Relationship\s+Id="(rId\d+)"\s+Type="[^"]*image[^"]*"\s+Target="([^"]+)"')
                    for match in rel_pattern.finditer(rels_content):
                        r_id = match.group(1)
                        target = match.group(2)
                        img_name = os.path.basename(target)
                        rels_mapping[r_id] = img_name

            # 2. 解析文档内容，提取图片ID出现顺序
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                if 'word/document.xml' in zip_file.namelist():
                    xml_content = zip_file.read('word/document.xml').decode('utf-8')
                    id_pattern = re.compile(r'(embed|link|r:id|id)="(rId\d+)"')
                    blip_ids = []
                    for match in id_pattern.finditer(xml_content):
                        r_id = match.group(2)
                        if r_id not in blip_ids:
                            blip_ids.append(r_id)

            # 3. 转换ID为图片名
            for r_id in blip_ids:
                if r_id in rels_mapping and rels_mapping[r_id] not in image_order:
                    image_order.append(rels_mapping[r_id])

        except Exception as e:
            print(f"⚠️ 解析图片顺序失败：{e}，将使用文件名排序")
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                img_names = []
                for file_info in zip_file.infolist():
                    if file_info.filename.startswith('word/media/') and not file_info.is_dir():
                        img_names.append(os.path.basename(file_info.filename))
                image_order = sorted(img_names)

        return image_order

    def _extract_original_images(self, docx_path, output_img_dir):
        """【内部方法】从DOCX中提取原始无压缩图片（强制绝对路径）"""
        # 确保输出目录是绝对路径
        output_img_dir = self._normalize_path(output_img_dir)

        # 清理并重建目录
        if os.path.exists(output_img_dir):
            shutil.rmtree(output_img_dir, ignore_errors=True)
        os.makedirs(output_img_dir, exist_ok=True)

        # 提取图片
        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            for file_info in zip_file.infolist():
                if file_info.filename.startswith('word/media/') and not file_info.is_dir():
                    img_filename = os.path.basename(file_info.filename)
                    # 强制使用绝对路径保存
                    save_path = self._normalize_path(os.path.join(output_img_dir, img_filename))
                    with open(save_path, 'wb') as f:
                        f.write(zip_file.read(file_info.filename))

        return [f for f in os.listdir(output_img_dir) if os.path.isfile(os.path.join(output_img_dir, f))]

    def _image_to_base64(self, img_path):
        """【内部方法】将图片文件转为Base64编码（带MIME前缀）"""
        # 强制转为绝对路径
        img_path = self._normalize_path(img_path)

        try:
            if not os.path.exists(img_path):
                print(f"⚠️ 图片文件不存在（绝对路径）：{img_path}")
                return ""

            with open(img_path, 'rb') as f:
                img_data = f.read()

            # 根据后缀匹配MIME类型
            img_ext = os.path.splitext(img_path)[1].lower()
            mime_map = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.bmp': 'image/bmp'
            }
            mime_type = mime_map.get(img_ext, 'image/png')
            base64_str = f"data:{mime_type};base64,{base64.b64encode(img_data).decode('utf-8')}"
            return base64_str
        except Exception as e:
            print(f"⚠️ 图片 {img_path} 转Base64失败：{e}")
            return ""

    def docx_to_single_html(self, docx_path, html_path):
        """
        公开方法：DOCX转单文件HTML（修复外部引用路径问题）
        :param docx_path: 输入DOCX文件路径（支持相对/绝对）
        :param html_path: 输出HTML文件路径（支持相对/绝对）
        :return: 生成的HTML文本内容
        """
        # 1. 基础校验
        docx_path = self._normalize_path(docx_path)
        html_path = self._normalize_path(html_path)

        if not os.path.exists(docx_path):
            print(f"❌ 输入DOCX文件不存在（绝对路径）：{docx_path}")
            return ""

        # 确保输出目录存在
        html_dir = os.path.dirname(html_path)
        os.makedirs(html_dir, exist_ok=True)

        # 2. 创建唯一临时目录（基于HTML输出目录，绝对路径）
        spire_temp_dir = self._normalize_path(os.path.join(html_dir, self.temp_dir_prefix))
        original_img_dir = self._normalize_path(os.path.join(spire_temp_dir, "original_images"))
        spire_img_dir = self._normalize_path(os.path.join(spire_temp_dir, "images"))  # Spire生成图片的目录

        # 3. 获取图片顺序+提取原始图片
        image_display_order = self._get_image_order_from_docx(docx_path)
        extracted_imgs = self._extract_original_images(docx_path, original_img_dir)
        if not image_display_order and extracted_imgs:
            image_display_order = sorted(extracted_imgs)

        # 4. Spire转换生成临时HTML（带外部资源）
        document = Document()
        try:
            document.LoadFromFile(docx_path)
            # 配置导出选项（强制绝对路径）
            document.HtmlExportOptions.ImageEmbedded = False
            document.HtmlExportOptions.ImagesPath = spire_img_dir  # 绝对路径
            document.HtmlExportOptions.ImageFormat = self.default_image_format
            document.SaveToFile(html_path, FileFormat.Html)
        except Exception as e:
            print(f"❌ Spire转换HTML失败：{e}")
            return ""
        finally:
            document.Close()
            del document

        # 5. 读取并处理HTML内容
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # 统一HTML中的路径分隔符
        html_content = html_content.replace('\\', '/')

        # 6. 提取Spire生成的图片列表（绝对路径匹配）
        spire_img_names = []
        img_pattern = re.compile(r'<img[^>]*src="([^"]+)"[^>]*>')
        for match in img_pattern.finditer(html_content):
            img_src = match.group(1)
            # 不管是相对还是绝对路径，只取文件名
            spire_img_name = os.path.basename(self._normalize_path(img_src))
            if spire_img_name not in spire_img_names:
                spire_img_names.append(spire_img_name)

        # 7. 内嵌CSS
        css_file_path = os.path.splitext(html_path)[0] + '_styles.css'
        css_file_path = self._normalize_path(css_file_path)
        if os.path.exists(css_file_path):
            with open(css_file_path, 'r', encoding='utf-8') as f:
                css_content = f.read()
            # 移除外部CSS引用，添加内联style
            html_content = re.sub(r'<link[^>]*href="[^"]+\.css"[^>]*>', '', html_content)
            html_content = html_content.replace(
                '</head>',
                f'<style type="text/css">\n{css_content}\n</style>\n</head>'
            )

        # 8. 内嵌图片（核心修复：绝对路径匹配+替换）
        for idx, spire_name in enumerate(spire_img_names):
            if idx < len(image_display_order):
                original_img_name = image_display_order[idx]
                # 强制使用绝对路径查找原始图片
                original_img_path = self._normalize_path(os.path.join(original_img_dir, original_img_name))
                base64_str = self._image_to_base64(original_img_path)

                if base64_str:
                    # 修复匹配规则：兼容相对/绝对路径的src属性
                    # 匹配所有包含该图片名的src属性（无论路径前缀）
                    spire_img_pattern = re.compile(f'src="[^"]*{re.escape(spire_name)}"')
                    html_content = spire_img_pattern.sub(f'src="{base64_str}"', html_content)
                else:
                    print(f"⚠️ 图片 {original_img_name} Base64转换失败，跳过替换")

        # 9. 保存最终HTML（强制覆盖）
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        # 10. 清理临时文件（忽略删除失败的情况）
        temp_paths = [spire_temp_dir, css_file_path]
        for temp_path in temp_paths:
            if os.path.exists(temp_path):
                try:
                    if os.path.isdir(temp_path):
                        shutil.rmtree(temp_path, ignore_errors=True)
                    else:
                        os.remove(temp_path)
                except Exception as e:
                    print(f"⚠️ 清理临时文件失败 {temp_path}：{e}")

        return html_content

    def html_text_to_docx(self, html_text: str, output_docx_path: str):
        """
        公开方法：HTML文本转DOCX（兼容所有Spire.Doc版本，无临时文件残留）
        """
        # 强制转为绝对路径
        output_docx_path = self._normalize_path(output_docx_path)

        if not html_text.strip():
            print(f"❌ HTML文本为空，无法转换")
            return False

        document = None
        temp_html_file = None
        try:
            # 确保输出目录存在
            output_dir = os.path.dirname(output_docx_path)
            os.makedirs(output_dir, exist_ok=True)

            # 1. 创建唯一临时HTML文件（绝对路径）
            temp_html_file = tempfile.NamedTemporaryFile(
                mode='w',
                suffix='.html',
                delete=False,
                encoding='utf-8'
            )
            temp_html_path = self._normalize_path(temp_html_file.name)

            # 2. 写入HTML文本并关闭句柄
            temp_html_file.write(html_text)
            temp_html_file.close()

            # 3. 加载并转换
            document = Document()
            document.LoadFromFile(
                temp_html_path,
                FileFormat.Html,
                self.html_validation_type
            )
            document.SaveToFile(output_docx_path, FileFormat.Docx2016)

            return True

        except Exception as e:
            print(f"❌ HTML转DOCX失败：{str(e)}")
            return False
        finally:
            # 强制清理资源
            if document:
                document.Close()
                del document

            if temp_html_file and os.path.exists(temp_html_file.name):
                try:
                    os.remove(temp_html_file.name)
                except Exception as e:
                    print(f"⚠️ 清理临时HTML文件失败：{e}")


# ===================== 异步下载图片函数 =====================
async def download_image(url: str, save_dir: str = TEMP_DIR) -> str:
    """
    异步下载网络图片到本地临时目录
    :param url: 图片网络地址
    :param save_dir: 保存目录
    :return: 本地文件路径
    """
    try:
        # 生成唯一文件名
        file_ext = url.split(".")[-1].split("?")[0]  # 处理URL带参数的情况
        if len(file_ext) > 5:  # 如果扩展名过长，使用默认
            file_ext = "jpg"
        file_name = f"{uuid.uuid4()}.{file_ext}"
        file_path = os.path.join(save_dir, file_name)

        # 异步下载图片
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=aiohttp.ClientTimeout(total=30)) as response:
                if response.status != 200:
                    raise HTTPException(status_code=400, detail=f"图片下载失败：{url}，状态码：{response.status}")

                # 保存图片到本地
                with open(file_path, "wb") as f:
                    f.write(await response.read())

        # 验证图片文件有效性
        try:
            with Image.open(file_path) as img:
                img.verify()  # 验证图片完整性
        except Exception as e:
            os.remove(file_path)
            raise HTTPException(status_code=400, detail=f"图片文件无效：{url}，错误：{str(e)}")

        return file_path

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"下载图片失败：{url}，错误：{str(e)}")


async def download_images(urls: List[str]) -> List[str]:
    """
    批量下载图片
    :param urls: 图片URL列表
    :return: 本地文件路径列表
    """
    if not urls:
        return []

    # 并发下载所有图片
    tasks = [download_image(url) for url in urls]
    local_paths = await asyncio.gather(*tasks)

    return local_paths


# ===================== 原有文档生成函数 =====================
def generate_report_doc(
        title_text,
        second_row_img_path,
        other_img_paths=None,
        save_path='年度报告.docx',
        columns_per_row=2
):
    other_img_paths = other_img_paths or []

    # 校验参数合法性
    if not isinstance(columns_per_row, int) or columns_per_row < 1:
        raise ValueError("每行图片数量必须是大于等于1的整数")

    # 校验图片文件是否存在
    if not os.path.exists(second_row_img_path):
        raise FileNotFoundError(f"图片不存在：{second_row_img_path}")
    for p in other_img_paths:
        if not os.path.exists(p):
            raise FileNotFoundError(f"图片不存在：{p}")

    # 创建干净文档
    doc = DocxDocument()

    # 字体设置函数
    def set_font(run, font_name, size, color=None, bold=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        run.bold = bold

    # 重新计算表格行数和列数
    other_img_rows = (len(other_img_paths) + columns_per_row - 1) // columns_per_row
    total_rows = 2 + other_img_rows
    table = doc.add_table(rows=total_rows, cols=columns_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 蓝色边框设置函数
    def set_cell_border(cell, color):
        for name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:color'), color)
            border.set(qn('w:sz'), '6')
            cell._element.tcPr.append(border)

    blue = '4472C4'
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, blue)
            cell.vertical_alignment = 1  # WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 强制前两行不分页
    for row in table.rows[:2]:
        trPr_list = row._element.xpath('.//w:trPr')
        if not trPr_list:
            trPr = OxmlElement('w:trPr')
            row._element.insert(0, trPr)
        else:
            trPr = trPr_list[0]

        kl = OxmlElement('w:keepLines')
        kl.set(qn('w:val'), 'true')
        trPr.append(kl)

    # 第一行：标题（合并所有列）
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells[1:]:
        hdr_cells[0].merge(cell)
    para = hdr_cells[0].paragraphs[0]
    run = para.add_run(title_text)
    set_font(run, '宋体', 12, RGBColor(255, 255, 255), bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), blue)
    hdr_cells[0]._element.tcPr.append(shd)

    # 第二行：大图（合并所有列）
    row2_cells = table.rows[1].cells
    for cell in row2_cells[1:]:
        row2_cells[0].merge(cell)
    para = row2_cells[0].paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(second_row_img_path, width=Inches(5.0))

    # 后面图片：按自定义列数排版
    img_width = Inches((5.0 * 0.9) / columns_per_row)
    for i, img_path in enumerate(other_img_paths):
        r = 2 + i // columns_per_row
        c = i % columns_per_row
        cell = table.rows[r].cells[c]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(img_path, width=img_width)

    doc.save(save_path)
    return save_path


def calculate_table_height(table):
    total_height = 0.0
    for row in table.rows:
        row_height = row.height if row.height else 0
        if row_height == 0:
            for cell in row.cells:
                cell_height = 0.0
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.element.xpath('.//a:blip'):
                            cell_height = Inches(2.8).inches
                            break
                    if cell_height > 0:
                        break
                if cell_height > 0:
                    row_height = cell_height
                    break
            if row_height == 0:
                row_height = 0.2
        total_height += row_height / 914400
    return total_height


def set_table_no_page_break(table):
    """最强防跨页：整表强制在同一页"""
    for row in table.rows:
        tr = row._element
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)

        # 禁止行被分页拆开
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)

        # 整行保持在一页
        keepLines = OxmlElement('w:keepLines')
        keepLines.set(qn('w:val'), 'true')
        trPr.append(keepLines)

        # 所有行连在一起
        keepNext = OxmlElement('w:keepNext')
        keepNext.set(qn('w:val'), 'true')
        trPr.append(keepNext)

    # 最后一行不要 keepNext
    if table.rows:
        last_tr = table.rows[-1]._element
        last_trPr = last_tr.find(qn('w:trPr'))
        if last_trPr is not None:
            keep_next_elems = [elem for elem in last_trPr if elem.tag.endswith('keepNext')]
            for elem in keep_next_elems:
                last_trPr.remove(elem)


def generate_fully_centered_patent_doc(
        patent_data,
        cert_img_paths=None,
        save_path='专利信息全居中文档.docx',
        fill_empty_space=True,
        last_img_display_mode=1
):
    if last_img_display_mode not in (1, 2):
        raise ValueError("last_img_display_mode 只能是 1 或 2")

    cert_img_paths = cert_img_paths or []
    valid_img_paths = [p for p in cert_img_paths if os.path.exists(p)]

    doc = DocxDocument()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    page_h = sec.page_height.inches - 2
    half_page = page_h / 2

    # 字体：宋体 小四（12磅）
    def set_font(run, color):
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.font.color.rgb = color

    def set_head_font(run):
        set_font(run, RGBColor(255, 255, 255))
        run.bold = True

    def set_border(cell, color='4472C4', sz='2'):
        for k in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{k}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:color'), color)
            b.set(qn('w:sz'), sz)
            cell._element.tcPr.append(b)

    # 专利表格
    table1 = doc.add_table(rows=len(patent_data) + 1, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = True

    heads = ['序号', '专利类型', '专利名称', '专利号', '专利权人', '授权公告日']
    for i, (cell, txt) in enumerate(zip(table1.rows[0].cells, heads)):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        set_head_font(r)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '4472C4')
        cell._element.tcPr.append(shd)
        set_border(cell, '4472C4')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for i, row in enumerate(patent_data, 1):
        for j, (cell, txt) in enumerate(zip(table1.rows[i].cells, row)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(txt))
            set_font(r, RGBColor(0, 0, 0))
            set_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    table_h = calculate_table_height(table1)

    # 分隔空行
    doc.add_paragraph()

    # 补图逻辑
    final_imgs = valid_img_paths.copy()
    if fill_empty_space and table_h < half_page and valid_img_paths:
        needed_height = half_page - table_h
        needed_rows = max(1, int(needed_height / 2.8) + 1)
        needed_imgs_min = needed_rows * 2

        if len(final_imgs) < needed_imgs_min:
            while len(final_imgs) < needed_imgs_min:
                final_imgs.extend(valid_img_paths)
            final_imgs = final_imgs[:needed_imgs_min]

    # 图片表格
    if final_imgs:
        n = len(final_imgs)
        img_table = doc.add_table(rows=(n + 1) // 2, cols=2)
        img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        img_table.autofit = True

        blue = '4472C4'
        for row in img_table.rows:
            for cell in row.cells:
                set_border(cell, blue, '6')
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.width = Inches(2.8)

        # 奇数最后一行处理
        if n % 2 == 1:
            last_row = img_table.rows[-1]
            if last_img_display_mode == 1:
                last_row.cells[1]._element.getparent().remove(last_row.cells[1]._element)
            elif last_img_display_mode == 2:
                last_row.cells[0].merge(last_row.cells[1])
                last_row.cells[0].width = Inches(5.6)

        # 图片插入逻辑
        inserted_count = 0
        for i in range(n):
            img_path = final_imgs[i]
            r = i // 2
            if n % 2 == 1 and r == len(img_table.rows) - 1 and last_img_display_mode == 2:
                c = 0
            else:
                c = i % 2

            if r >= len(img_table.rows):
                continue

            cell = img_table.rows[r].cells[c]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(img_path, width=Inches(2.8))
            inserted_count += 1

        set_table_no_page_break(img_table)

    doc.save(save_path)
    return save_path


def generate_car_info_doc(car_data, save_path='公司车辆信息.docx', table_title='公司车辆信息'):
    """生成车辆信息文档"""
    # 初始化文档
    doc = DocxDocument()
    # 页面边距设置
    sec = doc.sections[0]
    sec.top_margin = Cm(1.0)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.0)
    sec.right_margin = Cm(1.0)

    # 核心配色
    COLOR_BLUE = '4472C4'
    COLOR_WHITE = RGBColor(255, 255, 255)
    COLOR_BLACK = RGBColor(0, 0, 0)
    # 列宽配置
    COL_WIDTHS = [Cm(1.5), Cm(2.5), Cm(8.0), Cm(5.0)]
    # 图片插入宽度
    IMG_WIDTH_DRIVE = Cm(7.5)
    IMG_WIDTH_CAR = Cm(4.5)

    # 字体/边框函数
    def set_font(run, color):
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.font.color.rgb = color

    def set_head_font(run):
        set_font(run, RGBColor(255, 255, 255))
        run.bold = True

    def set_border(cell, color='4472C4', sz='2'):
        for k in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{k}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:color'), color)
            b.set(qn('w:sz'), sz)
            cell._element.tcPr.append(b)

    def set_cell_bg(cell, bg_color=COLOR_BLUE):
        """设置单元格底纹"""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg_color)
        cell._element.tcPr.append(shd)

    def add_text_to_cell(cell, text, is_header=False):
        """向单元格添加文字"""
        # 清空原有内容
        for para in cell.paragraphs:
            for run in para.runs:
                para._element.remove(run._element)
        # 文字居中排版
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # 插入文字并应用样式
        run = para.add_run(str(text))
        if is_header:
            set_head_font(run)
        else:
            set_font(run, COLOR_BLACK)
        # 单元格基础样式
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_border(cell)
        # 表头添加蓝色底纹
        if is_header:
            set_cell_bg(cell)

    def add_img_to_cell(cell, img_path, img_width):
        """向单元格插入图片"""
        # 清空原有内容
        for para in cell.paragraphs:
            for run in para.runs:
                para._element.remove(run._element)
        # 图片居中
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 插入图片/缺失提示
        if os.path.exists(img_path):
            para.add_run().add_picture(img_path, width=img_width)
        else:
            run = para.add_run(f'图片缺失：{os.path.basename(img_path)}')
            set_font(run, RGBColor(255, 0, 0))
        # 单元格样式
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_border(cell)

    # 数据校验
    if not isinstance(car_data, list) or len(car_data) == 0:
        raise ValueError("car_data必须是非空的二维列表")
    for idx, row in enumerate(car_data):
        if len(row) != 4:
            raise ValueError(f"car_data第{idx + 1}行必须包含[序号,车牌号,行驶证路径,车辆图片路径]4个元素")

    # 表格创建
    total_rows = 2 + len(car_data)
    total_cols = 4
    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 设置固定列宽
    for row in table.rows:
        for col_idx in range(total_cols):
            row.cells[col_idx].width = COL_WIDTHS[col_idx]

    # 第一行：合并4列+大标题
    first_row_cells = table.rows[0].cells
    main_title_cell = first_row_cells[0]
    for cell in first_row_cells[1:]:
        main_title_cell.merge(cell)
    add_text_to_cell(main_title_cell, table_title, is_header=True)

    # 第二行：二级表头
    second_headers = ['序号', '车牌号', '行驶证', '车辆图片']
    second_row_cells = table.rows[1].cells
    for col_idx, header in enumerate(second_headers):
        add_text_to_cell(second_row_cells[col_idx], header, is_header=True)

    # 数据行
    for data_idx, car_row in enumerate(car_data, start=2):
        seq, plate_num, drive_img, car_img = car_row
        current_cells = table.rows[data_idx].cells
        add_text_to_cell(current_cells[0], seq, is_header=False)
        add_text_to_cell(current_cells[1], plate_num, is_header=False)
        add_img_to_cell(current_cells[2], drive_img, IMG_WIDTH_DRIVE)
        add_img_to_cell(current_cells[3], car_img, IMG_WIDTH_CAR)

    # 保存文档
    doc.save(save_path)
    return save_path


# ===================== 工具函数 =====================
def generate_and_convert_to_html(generate_func, *args, **kwargs):
    """
    通用函数：生成DOCX文档并转换为HTML（修复文件占用问题）
    :param generate_func: 文档生成函数
    :return: HTML内容
    """
    # 1. 创建唯一的临时目录（避免文件冲突）
    temp_dir = tempfile.mkdtemp(prefix="docx_html_temp_")
    docx_path = os.path.join(temp_dir, f"temp_doc_{uuid.uuid4().hex[:8]}.docx")
    html_path = os.path.join(temp_dir, f"temp_html_{uuid.uuid4().hex[:8]}.html")

    try:
        # 2. 生成DOCX文档（使用普通文件，而非临时文件句柄）
        generate_func(*args, save_path=docx_path, **kwargs)

        # 3. 强制释放文件句柄（关键修复）
        import gc
        gc.collect()  # 强制垃圾回收，释放docx库的文件句柄
        time.sleep(0.1)  # 短暂等待，确保系统释放文件锁

        # 4. 转换为HTML（使用文件路径，而非句柄）
        converter = DocxHtmlConverter()
        html_content = converter.docx_to_single_html(docx_path, html_path)

        # 5. 读取HTML内容（确保内容完整）
        with open(html_path, 'r', encoding='utf-8') as f:
            final_html = f.read()

        return final_html

    finally:
        # 6. 强制清理所有临时文件（延迟清理，确保文件锁释放）
        def clean_temp_files():
            try:
                # 再次强制垃圾回收
                gc.collect()
                time.sleep(0.2)

                # 删除文件
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                if os.path.exists(html_path):
                    os.remove(html_path)
                # 删除目录
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception as e:
                print(f"⚠️ 清理临时文件警告：{e}")

        # 使用线程延迟清理（避免当前进程占用）
        import threading
        cleanup_thread = threading.Thread(target=clean_temp_files)
        cleanup_thread.daemon = True
        cleanup_thread.start()


# ===================== FastAPI接口定义（带详细入参说明） =====================

@app.post(
    "/generate-report-html",
    response_class=HTMLResponse,
    summary="生成年度报告文档并返回HTML",
    description="""
    生成包含标题和多张图片的年度报告Word文档，并转换为单文件HTML返回
    - 文档结构：表格布局，第一行是标题，第二行是主图，后续行是多列排版的其他图片
    - HTML特性：图片Base64内嵌，CSS内联，无外部资源依赖
    """
)
async def api_generate_report_html(
        title_text: str = Body(
            ...,
            description="报告标题文本（显示在表格第一行）",
            example="2024年度销售报告",
            min_length=1,
            max_length=100
        ),
        second_row_img_url: str = Body(
            ...,
            description="第二行主图片的网络URL（必填），支持HTTP/HTTPS",
            example="https://example.com/main_chart.jpg",
            regex=r'^https?://.*\.(jpg|jpeg|png|gif|bmp)$'
        ),
        other_img_urls: List[str] = Body(
            [],
            description="其他图片的网络URL列表（可选），会按列数排版在主图下方",
            example=["https://example.com/img1.jpg", "https://example.com/img2.png"],
            min_items=0
        ),
        columns_per_row: int = Body(
            2,
            description="每行显示的图片列数（主图除外）",
            example=2,
            ge=1,  # 大于等于1
            le=4  # 小于等于4
        )
):
    try:
        # 下载图片
        main_img_path = await download_image(second_row_img_url)
        other_img_paths = await download_images(other_img_urls)

        # 生成文档并转换为HTML
        html_content = generate_and_convert_to_html(
            generate_report_doc,
            title_text=title_text,
            second_row_img_path=main_img_path,
            other_img_paths=other_img_paths,
            columns_per_row=columns_per_row
        )

        # 清理下载的图片文件
        try:
            os.remove(main_img_path)
            for img_path in other_img_paths:
                os.remove(img_path)
        except:
            pass

        if not html_content:
            raise HTTPException(status_code=500, detail="HTML转换失败，生成的HTML内容为空")

        return HTMLResponse(content=html_content, status_code=200)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"生成报告失败：{str(e)}")


@app.post(
    "/generate-patent-html",
    response_class=HTMLResponse,
    summary="生成专利信息文档并返回HTML",
    description="""
    生成包含专利表格和证书图片的Word文档，并转换为单文件HTML
    - 表格列：序号、专利类型、专利名称、专利号、专利权人、授权公告日
    - 图片特性：自动补图填充空白区域，支持奇数图片的两种显示模式
    """
)
async def api_generate_patent_html(
        patent_data: List[List[str]] = Body(
            ...,
            description="专利数据二维列表（必填），每行对应一条专利信息，顺序：[序号,专利类型,专利名称,专利号,专利权人,授权公告日]",
            example=[
                ["1", "发明专利", "一种智能温控系统", "ZL202410000001.0", "某某科技有限公司", "2024-01-15"],
                ["2", "实用新型专利", "一种节能电机", "ZL202420000001.0", "某某科技有限公司", "2024-02-20"]
            ],
            min_items=1,  # 至少1条专利数据
            max_items=100  # 最多100条专利数据
        ),
        cert_img_urls: List[str] = Body(
            [],
            description="专利证书图片的网络URL列表（可选）",
            example=["https://example.com/patent1.jpg", "https://example.com/patent2.png"],
            min_items=0
        ),
        fill_empty_space: bool = Body(
            True,
            description="是否自动补图填充页面空白区域（重复使用现有图片）",
            example=True
        ),
        last_img_display_mode: int = Body(
            1,
            description="奇数图片时最后一行的显示模式：1=删除空单元格，2=合并单元格显示大图",
            example=1,
            ge=1,
            le=2
        )
):
    try:
        # 数据校验：确保每行都有6个字段
        for idx, row in enumerate(patent_data):
            if len(row) != 6:
                raise ValueError(
                    f"专利数据第{idx + 1}行必须包含6个字段（序号、专利类型、专利名称、专利号、专利权人、授权公告日）")

        # 下载图片
        cert_img_paths = await download_images(cert_img_urls)

        # 生成文档并转换为HTML
        html_content = generate_and_convert_to_html(
            generate_fully_centered_patent_doc,
            patent_data=patent_data,
            cert_img_paths=cert_img_paths,
            fill_empty_space=fill_empty_space,
            last_img_display_mode=last_img_display_mode
        )

        # 清理下载的图片文件
        try:
            for img_path in cert_img_paths:
                os.remove(img_path)
        except:
            pass

        if not html_content:
            raise HTTPException(status_code=500, detail="HTML转换失败，生成的HTML内容为空")

        return HTMLResponse(content=html_content, status_code=200)

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"生成专利文档失败：{str(e)}")


@app.post(
    "/generate-car-info-html",
    response_class=HTMLResponse,
    summary="生成车辆信息文档并返回HTML",
    description="""
    生成包含公司车辆信息的Word文档，并转换为单文件HTML
    - 表格列：序号、车牌号、行驶证图片、车辆图片
    - 图片特性：行驶证和车辆图片分别按固定宽度显示
    """
)
async def api_generate_car_info_html(
        car_data: List[List[Any]] = Body(
            ...,
            description="车辆数据二维列表（必填），每行顺序：[序号,车牌号,行驶证图片URL,车辆图片URL]",
            example=[
                ["1", "京A12345", "https://example.com/drive1.jpg", "https://example.com/car1.jpg"],
                ["2", "沪B67890", "https://example.com/drive2.jpg", "https://example.com/car2.jpg"]
            ],
            min_items=1,
            max_items=50
        ),
        table_title: str = Body(
            '公司车辆信息',
            description="表格标题（显示在表格第一行）",
            example="2024年公司运营车辆信息",
            min_length=1,
            max_length=50
        )
):
    try:
        # 数据校验
        for idx, row in enumerate(car_data):
            if len(row) != 4:
                raise ValueError(f"车辆数据第{idx + 1}行必须包含4个字段（序号、车牌号、行驶证URL、车辆图片URL）")

            # 校验URL格式
            seq, plate_num, drive_url, car_url = row
            url_pattern = r'^https?://.*\.(jpg|jpeg|png|gif|bmp)$'
            if not re.match(url_pattern, drive_url):
                raise ValueError(f"第{idx + 1}行行驶证URL格式错误：{drive_url}（必须是HTTP/HTTPS图片链接）")
            if not re.match(url_pattern, car_url):
                raise ValueError(f"第{idx + 1}行车辆图片URL格式错误：{car_url}（必须是HTTP/HTTPS图片链接）")

        # 下载所有车辆相关图片
        car_data_with_local_paths = []
        all_img_urls = []

        # 收集所有图片URL
        for car_row in car_data:
            seq, plate_num, drive_url, car_url = car_row
            all_img_urls.append(drive_url)
            all_img_urls.append(car_url)

        # 批量下载图片
        img_paths = await download_images(all_img_urls)
        img_path_map = dict(zip(all_img_urls, img_paths))

        # 替换URL为本地路径
        for car_row in car_data:
            seq, plate_num, drive_url, car_url = car_row
            drive_path = img_path_map.get(drive_url, "")
            car_path = img_path_map.get(car_url, "")
            car_data_with_local_paths.append([seq, plate_num, drive_path, car_path])

        # 生成文档并转换为HTML
        html_content = generate_and_convert_to_html(
            generate_car_info_doc,
            car_data=car_data_with_local_paths,
            table_title=table_title
        )

        # 清理下载的图片文件
        try:
            for img_path in img_paths:
                os.remove(img_path)
        except:
            pass

        if not html_content:
            raise HTTPException(status_code=500, detail="HTML转换失败，生成的HTML内容为空")

        return HTMLResponse(content=html_content, status_code=200)

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"生成车辆信息文档失败：{str(e)}")


@app.get(
    "/",
    summary="健康检查/接口说明",
    description="API服务状态检查，返回接口列表和基本说明"
)
async def root():
    return {
        "status": "running",
        "service": "Word文档生成API",
        "version": "1.0.0",
        "endpoints": [
            {
                "path": "/generate-report-html",
                "method": "POST",
                "description": "生成年度报告文档并返回HTML"
            },
            {
                "path": "/generate-patent-html",
                "method": "POST",
                "description": "生成专利信息文档并返回HTML"
            },
            {
                "path": "/generate-car-info-html",
                "method": "POST",
                "description": "生成车辆信息文档并返回HTML"
            }
        ],
        "notes": [
            "所有接口返回的HTML均为单文件，图片Base64内嵌，无外部资源依赖",
            "支持的图片格式：JPG/JPEG/PNG/GIF/BMP",
            "图片URL必须是可公开访问的HTTP/HTTPS链接"
        ]
    }


# 启动服务
if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8081,
        log_level="info"
    )